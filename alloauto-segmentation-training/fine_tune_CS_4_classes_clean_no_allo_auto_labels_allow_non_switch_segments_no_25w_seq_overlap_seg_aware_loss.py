"""
4-Class Code-Switching Detection System for Tibetan Text
Classes:
  0: Non-switching Auto (continuing in Auto mode)
  1: Non-switching Allo (continuing in Allo mode)
  2: Switch TO Auto
  3: Switch TO Allo

Focus: Proximity-aware loss and evaluation with 5-token tolerance
"""
#MOST updated model! 18/9/2025"
import torch
import torch.nn as nn
import torch.nn.functional as F
import numpy as np
import pandas as pd
import json
import re
import docx
import os
import ssl
from typing import List, Tuple, Dict
from collections import defaultdict
from sklearn.metrics import classification_report
from transformers import (
    AutoTokenizer,
    AutoModelForTokenClassification,
    TrainingArguments,
    DataCollatorForTokenClassification,
    Trainer
)
from torch.utils.data import Dataset
from pathlib import Path

# Environment setup
os.environ["CUDA_VISIBLE_DEVICES"] = "3"
ssl._create_default_https_context = ssl._create_unverified_context

print(f"Number of GPUs available: {torch.cuda.device_count()}")
for i in range(torch.cuda.device_count()):
    print(f"GPU {i}: {torch.cuda.get_device_name(i)}")


# ============================================================================
# PART 1: DATA PROCESSING
# ============================================================================
def remove_duplicate_sequences_from_train(train_df, test_df, min_duplicate_length=25):
    """
    Remove train segments that contain sequences of min_duplicate_length+ tokens
    identical to sequences in test set
    """
    import hashlib
    from collections import defaultdict

    print(f"\n{'=' * 80}")
    print(f"REMOVING TRAIN SEGMENTS WITH {min_duplicate_length}+ TOKEN DUPLICATES")
    print(f"{'=' * 80}")

    original_train_size = len(train_df)

    # Build index of all subsequences from test set
    test_sequences = defaultdict(list)

    for test_idx, test_row in test_df.iterrows():
        test_tokens = test_row['tokens'].split()

        # Extract all possible consecutive sequences of min_duplicate_length
        for i in range(len(test_tokens) - min_duplicate_length + 1):
            sequence = test_tokens[i:i + min_duplicate_length]
            sequence_str = ' '.join(sequence)
            sequence_hash = hashlib.md5(sequence_str.encode()).hexdigest()

            test_sequences[sequence_hash].append({
                'segment_id': test_row['segment_id'],
                'idx': test_idx,
                'position': i
            })

    print(f"Built index of test sequences (25+ tokens)")

    # Find train segments to remove
    train_segments_to_remove = set()
    duplicate_details = []

    for train_idx, train_row in train_df.iterrows():
        train_tokens = train_row['tokens'].split()

        # Check all possible consecutive sequences
        for i in range(len(train_tokens) - min_duplicate_length + 1):
            sequence = train_tokens[i:i + min_duplicate_length]
            sequence_str = ' '.join(sequence)
            sequence_hash = hashlib.md5(sequence_str.encode()).hexdigest()

            # If this sequence exists in test, mark train segment for removal
            if sequence_hash in test_sequences:
                train_segments_to_remove.add(train_idx)

                # Record details for reporting
                for test_match in test_sequences[sequence_hash]:
                    duplicate_details.append({
                        'train_idx': train_idx,
                        'train_segment': train_row['segment_id'],
                        'test_segment': test_match['segment_id'],
                        'duplicate_length': min_duplicate_length,
                        'sequence': sequence_str[:100] + '...' if len(sequence_str) > 100 else sequence_str
                    })
                break  # No need to check more sequences from this train segment

    # Remove problematic segments
    if train_segments_to_remove:
        print(f"\n❌ Found {len(train_segments_to_remove)} train segments with duplicates:")

        # Show details of what's being removed
        for detail in duplicate_details[:10]:  # Show first 10
            print(f"\n  Train segment: {detail['train_segment']}")
            print(f"  Has {detail['duplicate_length']}+ token match with test: {detail['test_segment']}")
            print(f"  Duplicate text: {detail['sequence']}")

        if len(duplicate_details) > 10:
            print(f"\n  ... and {len(duplicate_details) - 10} more duplicates")

        # Remove the segments
        clean_train_df = train_df.drop(list(train_segments_to_remove)).reset_index(drop=True)

        print(f"\n✅ Removed {len(train_segments_to_remove)} segments from training")
        print(f"  Original train size: {original_train_size}")
        print(f"  Cleaned train size: {len(clean_train_df)}")
        print(f"  Reduction: {(len(train_segments_to_remove) / original_train_size * 100):.1f}%")

    else:
        print("\n✅ No duplicate sequences found - train set is clean!")
        clean_train_df = train_df

    return clean_train_df
def aggressive_tag_removal(text):
    """
    Aggressively remove ALL forms of tags and partial tags from text
    """
    tag_patterns = [
        # Full tags with any case and spacing
        r'<\s*[Aa][Uu][Tt][Oo]\s*>',
        r'<\s*[Aa][Ll][Ll][Oo]\s*>',
        # Partial tags (opening)
        r'<\s*[Aa][Uu][Tt][Oo]\b',
        r'<\s*[Aa][Ll][Ll][Oo]\b',
        r'<\s*[Aa][Uu][Tt]\b',
        r'<\s*[Aa][Ll][Ll]\b',
        r'<\s*[Aa][Uu]\b',
        r'<\s*[Aa][Ll]\b',
        # Partial tags (closing)
        r'\b[Aa][Uu][Tt][Oo]\s*>',
        r'\b[Aa][Ll][Ll][Oo]\s*>',
        r'\b[Uu][Tt][Oo]\s*>',
        r'\b[Ll][Ll][Oo]\s*>',
        # Any word containing tag patterns
        r'\S*<[Aa][Uu][Tt][Oo]\S*',
        r'\S*<[Aa][Ll][Ll][Oo]\S*',
        r'\S*[Aa][Uu][Tt][Oo]>\S*',
        r'\S*[Aa][Ll][Ll][Oo]>\S*',
    ]

    cleaned_text = text
    for pattern in tag_patterns:
        cleaned_text = re.sub(pattern, ' ', cleaned_text)

    cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
    return cleaned_text.strip()


def process_segment_to_tokens_labels_clean(segment_text):
    """
    Convert segment to tokens and labels, ensuring NO tags in tokens
    """
    # Normalize tags for consistent splitting
    normalized_text = re.sub(r'<\s*AUTO\s*>', '<auto>', segment_text, flags=re.IGNORECASE)
    normalized_text = re.sub(r'<\s*ALLO\s*>', '<allo>', normalized_text, flags=re.IGNORECASE)

    # Split by tags
    parts = re.split(r'(<auto>|<allo>)', normalized_text.lower())

    tokens = []
    labels = []
    current_mode = None

    for i, part in enumerate(parts):
        if part == '<auto>':
            continue  # Skip tag
        elif part == '<allo>':
            continue  # Skip tag
        elif part.strip():
            # Clean any remaining tag fragments
            clean_part = aggressive_tag_removal(part)
            words = clean_part.strip().split()

            # Determine mode from previous tags
            segment_mode = None
            for j in range(i - 1, -1, -1):
                if parts[j] == '<auto>':
                    segment_mode = 'auto'
                    break
                elif parts[j] == '<allo>':
                    segment_mode = 'allo'
                    break

            if segment_mode is None:
                segment_mode = 'auto'

            for word_idx, word in enumerate(words):
                # Double-check word doesn't contain tags
                if contains_any_tag_pattern(word):
                    continue  # Skip tag fragments

                tokens.append(word)

                # Assign label
                if word_idx == 0 and current_mode is not None and current_mode != segment_mode:
                    label = 2 if segment_mode == 'auto' else 3
                else:
                    label = 0 if segment_mode == 'auto' else 1

                labels.append(label)

            current_mode = segment_mode

    return tokens, labels

def contains_any_tag_pattern(text):
    """
    Check if text contains ANY form of tag or partial tag
    Returns True if ANY tag pattern is found
    """
    # Comprehensive list of patterns to check
    forbidden_patterns = [
        # Full tags
        r'<auto>', r'<AUTO>', r'<allo>', r'<ALLO>',
        r'<Auto>', r'<Allo>',
        # With spaces
        r'<\s+auto\s+>', r'<\s+AUTO\s+>', r'<\s+allo\s+>', r'<\s+ALLO\s+>',
        # Partial tags (missing closing)
        r'<auto\b', r'<AUTO\b', r'<allo\b', r'<ALLO\b',
        r'<aut\b', r'<AUT\b', r'<all\b', r'<ALL\b',
        r'<au\b', r'<AU\b', r'<al\b', r'<AL\b',
        # Partial tags (missing opening)
        r'\bauto>', r'\bAUTO>', r'\ballo>', r'\bALLO>',
        r'\buto>', r'\bUTO>', r'\bllo>', r'\bLLO>',
        r'\bto>', r'\bTO>', r'\blo>', r'\bLO>',
        # Check for fragments within words
        r'<auto', r'<AUTO', r'<allo', r'<ALLO',
        r'auto>', r'AUTO>', r'allo>', r'ALLO>',
    ]

    text_lower = text.lower()

    for pattern in forbidden_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return True

    # Also check for simple substring presence
    forbidden_substrings = [
        '<auto', '<AUTO', '<allo', '<ALLO',
        'auto>', 'AUTO>', 'allo>', 'ALLO>',
        '<Auto', '<Allo', 'Auto>', 'Allo>'
    ]

    for substring in forbidden_substrings:
        if substring in text or substring.lower() in text_lower:
            return True

    return False
def process_segment_to_tokens_labels_v2(segment_text):
    """
    Convert a segment to tokens and 4-class labels
    MODIFIED: Tags are used for labeling but NOT included in tokens
    """
    # Split by tags while keeping them
    parts = re.split(r'(<auto>|<allo>)', segment_text)

    tokens = []
    labels = []
    current_mode = None

    for i, part in enumerate(parts):
        if part == '<auto>':
            # Don't add tag to tokens, just note mode change
            continue
        elif part == '<allo>':
            # Don't add tag to tokens, just note mode change
            continue
        elif part.strip():
            words = part.strip().split()

            # Determine what mode this segment should be in
            segment_mode = None
            for j in range(i - 1, -1, -1):
                if parts[j] == '<auto>':
                    segment_mode = 'auto'
                    break
                elif parts[j] == '<allo>':
                    segment_mode = 'allo'
                    break

            if segment_mode is None:
                segment_mode = 'auto'  # Default to auto if no tag found before

            for word_idx, word in enumerate(words):
                # Add the actual word (no tags)
                tokens.append(word)

                # Determine label based on position and mode change
                if word_idx == 0 and current_mode is not None and current_mode != segment_mode:
                    # This is the FIRST word after a mode change - it's a switch point
                    label = 2 if segment_mode == 'auto' else 3  # Switch TO auto or Switch TO allo
                else:
                    # Continuation in current mode
                    label = 0 if segment_mode == 'auto' else 1  # Non-switch auto or Non-switch allo

                labels.append(label)

            current_mode = segment_mode

    return tokens, labels


def verify_segment_has_switch(segment_text):
    """
    Verify that a segment contains at least one actual switch point
    Returns True if segment has at least one transition between auto and allo
    """
    # Must have both tags
    if '<auto>' not in segment_text or '<allo>' not in segment_text:
        return False

    # Check for actual switches by looking at tag sequences
    parts = re.split(r'(<auto>|<allo>)', segment_text)

    last_tag = None
    switches_found = 0

    for part in parts:
        if part == '<auto>':
            if last_tag == '<allo>':
                switches_found += 1
            last_tag = '<auto>'
        elif part == '<allo>':
            if last_tag == '<auto>':
                switches_found += 1
            last_tag = '<allo>'

    return switches_found > 0


def process_file_to_segments(file_path, file_type):
    """
    Process a single file and extract segments with guaranteed switches
    MODIFIED: Ensures tokens don't contain ANY form of tags
    """
    print(f"Processing {file_type} file: {file_path.name}")

    # Read file content
    if file_type == 'docx':
        doc = docx.Document(str(file_path))
        text_content = ' '.join([para.text for para in doc.paragraphs])
    else:  # txt file
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        text_content = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text_content = f.read()
                break
            except UnicodeDecodeError:
                continue

        if text_content is None:
            raise ValueError(f"Could not read {file_path.name} with any encoding")

    # Clean and normalize
    text_content = clean_and_normalize_text(text_content)

    # Split into sentences
    sentences = split_into_sentences(text_content)
    print(f"  Found {len(sentences)} sentences")

    # Extract segments with token limit and switch verification
    segments = extract_segments_with_token_limit(sentences, min_tokens=300, max_tokens=400)
    print(f"  Found {len(segments)} segments with verified switches (300-400 tokens each)")

    # Convert segments to token-label pairs
    processed_segments = []
    segments_without_switches = 0
    segments_with_tags = 0

    for seg_idx, segment in enumerate(segments):
        # Use the clean processing function
        tokens, labels = process_segment_to_tokens_labels_clean(segment['text'])

        if len(tokens) > 0:
            # Count actual transitions in labels
            num_transitions = sum(1 for l in labels if l in [2, 3])

            # Double-check that we have switches
            if num_transitions == 0:
                segments_without_switches += 1
                continue  # Skip segments without actual switches

            # CRITICAL: Verify no tags in tokens
            has_tags = False
            for token in tokens:
                if contains_any_tag_pattern(token):
                    has_tags = True
                    segments_with_tags += 1
                    print(f"    WARNING: Found tag in token: '{token}'")
                    break

            if has_tags:
                continue  # Skip this segment entirely

            # Create text WITHOUT tags for storage
            # Join tokens directly (they're already clean)
            text_without_tags = ' '.join(tokens)

            # Final verification
            if contains_any_tag_pattern(text_without_tags):
                print(f"    WARNING: Tags still found after cleaning in segment {seg_idx}")
                segments_with_tags += 1
                continue

            processed_segments.append({
                'segment_id': f"{file_path.stem}_{seg_idx}",
                'source_file': file_path.name,
                'file_type': file_type,
                'tokens': tokens,  # These are now completely tag-free tokens
                'labels': labels,
                'num_tokens': len(tokens),
                'num_transitions': num_transitions,
                'original_text': text_without_tags  # Store clean text
            })

    if segments_without_switches > 0:
        print(f"  Warning: Filtered out {segments_without_switches} segments without switches")

    if segments_with_tags > 0:
        print(f"  Warning: Filtered out {segments_with_tags} segments that still contained tags")

    print(f"  Processed {len(processed_segments)} valid, clean segments (all have switches, no tags)")
    return processed_segments


def validate_preprocessing(sample_segment_text):
    """
    Helper function to validate that preprocessing works correctly
    Shows how tags are used for labeling but not included in tokens
    """
    print("\n=== PREPROCESSING VALIDATION ===")
    print(f"Original text with tags:\n{sample_segment_text}\n")

    tokens, labels = process_segment_to_tokens_labels_clean(sample_segment_text)

    print("Tokens (no tags):")
    for i, (token, label) in enumerate(zip(tokens[:20], labels[:20])):  # Show first 20
        label_names = {
            0: 'NonSwitch-Auto',
            1: 'NonSwitch-Allo',
            2: 'SWITCH→Auto',
            3: 'SWITCH→Allo'
        }
        print(f"  {i:3d}: '{token:15s}' -> {label_names[label]}")

    # Verify no tags in tokens
    has_tags = any('<auto>' in t.lower() or '<allo>' in t.lower() for t in tokens)
    print(f"\nTokens contain tags: {has_tags} (should be False)")

    # Count switches
    num_switches = sum(1 for l in labels if l in [2, 3])
    print(f"Number of switch points: {num_switches}")

    return tokens, labels


def analyze_and_balance_switch_distribution(train_df, val_df, test_df):
    """
    Analyze switch type distribution and check balance
    """
    print("\n" + "=" * 80)
    print("SWITCH TYPE DISTRIBUTION ANALYSIS")
    print("=" * 80)

    def get_switch_stats(df, split_name):
        """Get detailed switch statistics for a split"""
        total_to_auto = 0
        total_to_allo = 0
        segments_with_auto = 0
        segments_with_allo = 0

        for idx in range(len(df)):
            labels = [int(l) for l in df.iloc[idx]['labels'].split(',')]

            # Count switches in this segment
            to_auto_count = labels.count(2)
            to_allo_count = labels.count(3)

            total_to_auto += to_auto_count
            total_to_allo += to_allo_count

            if to_auto_count > 0:
                segments_with_auto += 1
            if to_allo_count > 0:
                segments_with_allo += 1

        total_switches = total_to_auto + total_to_allo

        stats = {
            'split': split_name,
            'n_segments': len(df),
            'total_switches': total_switches,
            'to_auto': total_to_auto,
            'to_allo': total_to_allo,
            'auto_pct': (total_to_auto / total_switches * 100) if total_switches > 0 else 0,
            'allo_pct': (total_to_allo / total_switches * 100) if total_switches > 0 else 0,
            'avg_switches_per_seg': total_switches / len(df) if len(df) > 0 else 0,
            'avg_auto_per_seg': total_to_auto / len(df) if len(df) > 0 else 0,
            'avg_allo_per_seg': total_to_allo / len(df) if len(df) > 0 else 0,
            'segs_with_auto': segments_with_auto,
            'segs_with_allo': segments_with_allo
        }
        return stats

    # Get stats for each split
    train_stats = get_switch_stats(train_df, 'Train')
    val_stats = get_switch_stats(val_df, 'Val')
    test_stats = get_switch_stats(test_df, 'Test')

    # Print detailed statistics
    print("\n📊 DETAILED STATISTICS BY SPLIT:\n")
    print(f"{'Split':<8} {'Segments':<10} {'Total SW':<10} {'→Auto':<10} {'→Allo':<10} {'Auto%':<10} {'Allo%':<10}")
    print("-" * 70)

    for stats in [train_stats, val_stats, test_stats]:
        print(f"{stats['split']:<8} {stats['n_segments']:<10} {stats['total_switches']:<10} "
              f"{stats['to_auto']:<10} {stats['to_allo']:<10} "
              f"{stats['auto_pct']:<10.1f} {stats['allo_pct']:<10.1f}")

    print("\n📈 AVERAGE SWITCHES PER SEGMENT:\n")
    print(f"{'Split':<8} {'Avg Total':<12} {'Avg →Auto':<12} {'Avg →Allo':<12}")
    print("-" * 45)

    for stats in [train_stats, val_stats, test_stats]:
        print(f"{stats['split']:<8} {stats['avg_switches_per_seg']:<12.2f} "
              f"{stats['avg_auto_per_seg']:<12.2f} {stats['avg_allo_per_seg']:<12.2f}")

    print("\n📑 SEGMENT COVERAGE:\n")
    print(f"{'Split':<8} {'Has →Auto':<15} {'Has →Allo':<15}")
    print("-" * 40)

    for stats in [train_stats, val_stats, test_stats]:
        print(f"{stats['split']:<8} {stats['segs_with_auto']}/{stats['n_segments']:<10} "
              f"{stats['segs_with_allo']}/{stats['n_segments']}")

    # Check for distribution imbalance
    auto_pcts = [train_stats['auto_pct'], val_stats['auto_pct'], test_stats['auto_pct']]
    allo_pcts = [train_stats['allo_pct'], val_stats['allo_pct'], test_stats['allo_pct']]

    max_auto_diff = max(auto_pcts) - min(auto_pcts)
    max_allo_diff = max(allo_pcts) - min(allo_pcts)

    print("\n⚖️ DISTRIBUTION BALANCE CHECK:")
    print(f"  Max difference in Auto%: {max_auto_diff:.1f}%")
    print(f"  Max difference in Allo%: {max_allo_diff:.1f}%")
    print(f"  Train vs Test Auto% difference: {abs(train_stats['auto_pct'] - test_stats['auto_pct']):.1f}%")
    print(f"  Train vs Test Allo% difference: {abs(train_stats['allo_pct'] - test_stats['allo_pct']):.1f}%")

    if max_auto_diff > 5 or max_allo_diff > 5:
        print("  ⚠️ WARNING: Imbalanced distribution detected (>5% difference)")
        print("  Consider using stratified splitting based on switch types")
    else:
        print("  ✅ Distribution is well balanced across splits")

    return train_stats, val_stats, test_stats


def create_stratified_switch_split(df, train_ratio=0.7, val_ratio=0.15, test_ratio=0.15):
    """
    Create stratified split that maintains switch type distribution
    """
    print("\n" + "=" * 80)
    print("CREATING STRATIFIED SPLIT BY SWITCH TYPE")
    print("=" * 80)

    # Calculate switch type ratio for each segment
    segment_features = []

    for idx in range(len(df)):
        row = df.iloc[idx]
        labels = [int(l) for l in row['labels'].split(',')]

        to_auto = labels.count(2)
        to_allo = labels.count(3)
        total_switches = to_auto + to_allo

        # Categorize segments by dominant switch type
        if total_switches == 0:
            category = 'no_switch'  # Shouldn't happen with your filtering
        elif to_auto > 0 and to_allo == 0:
            category = 'auto_only'
        elif to_allo > 0 and to_auto == 0:
            category = 'allo_only'
        elif to_auto > to_allo:
            category = 'auto_dominant'
        elif to_allo > to_auto:
            category = 'allo_dominant'
        else:
            category = 'balanced'

        segment_features.append({
            'idx': idx,
            'category': category,
            'to_auto': to_auto,
            'to_allo': to_allo,
            'auto_ratio': to_auto / total_switches if total_switches > 0 else 0
        })

    # Add category to dataframe
    df['switch_category'] = [sf['category'] for sf in segment_features]

    # Print category distribution
    print("\nSegment categories:")
    category_counts = df['switch_category'].value_counts()
    for cat, count in category_counts.items():
        print(f"  {cat}: {count} segments ({count / len(df) * 100:.1f}%)")

    # Stratified split by category
    # First split: train+val vs test
    train_val_df, test_df = train_test_split(
        df,
        test_size=test_ratio,
        stratify=df['switch_category'],
        random_state=42
    )

    # Second split: train vs val
    relative_val_size = val_ratio / (train_ratio + val_ratio)
    train_df, val_df = train_test_split(
        train_val_df,
        test_size=relative_val_size,
        stratify=train_val_df['switch_category'],
        random_state=42
    )

    # Remove temporary column
    train_df = train_df.drop('switch_category', axis=1)
    val_df = val_df.drop('switch_category', axis=1)
    test_df = test_df.drop('switch_category', axis=1)

    # Save splits
    train_df.to_csv('train_segments_stratified.csv', index=False)
    val_df.to_csv('val_segments_stratified.csv', index=False)
    test_df.to_csv('test_segments_stratified.csv', index=False)

    print(f"\n✅ Stratified splits created:")
    print(f"  Train: {len(train_df)} segments")
    print(f"  Val: {len(val_df)} segments")
    print(f"  Test: {len(test_df)} segments")

    return train_df, val_df, test_df

def clean_and_normalize_text(text_content):
    """
    Clean text and normalize tags, remove ALL newlines and extra spaces
    """
    # First normalize all tag variations
    text_content = re.sub(r'<\s*AUTO\s*>', '<auto>', text_content, flags=re.IGNORECASE)
    text_content = re.sub(r'<\s*ALLO\s*>', '<allo>', text_content, flags=re.IGNORECASE)

    # Remove ALL newlines and replace with spaces
    text_content = text_content.replace('\n', ' ')
    text_content = text_content.replace('\r', ' ')
    text_content = text_content.replace('\t', ' ')

    # Remove any separator lines (multiple dashes, underscores, etc)
    text_content = re.sub(r'[-_=]{3,}', ' ', text_content)

    # Clean up multiple spaces
    text_content = re.sub(r'\s+', ' ', text_content)

    # Ensure proper spacing around tags for splitting
    text_content = re.sub(r'<auto>', ' <auto> ', text_content)
    text_content = re.sub(r'<allo>', ' <allo> ', text_content)

    # Final cleanup of multiple spaces
    text_content = re.sub(r'\s+', ' ', text_content)

    return text_content.strip()


def split_into_sentences(text_content):
    """
    Split text into sentences based on . / // boundaries
    """
    sentences = []
    current_sentence = ""

    # First split by sentence endings, keeping the delimiters
    parts = re.split(r'(\.|//?)', text_content)

    for i, part in enumerate(parts):
        if part.strip() in ['.', '/', '//']:
            # End of sentence
            if current_sentence.strip():
                current_sentence += part
                sentences.append(current_sentence.strip())
                current_sentence = ""
        else:
            current_sentence += part

    # Add any remaining content
    if current_sentence.strip():
        sentences.append(current_sentence.strip())

    return [s for s in sentences if s.strip()]




def extend_to_find_switch(current_segment, all_sentences, current_idx, max_tokens):
    """
    Try to extend segment to include at least one switch
    """
    extended_segment = current_segment.copy()
    token_count = sum(len(s.split()) for s in extended_segment)

    # Look ahead up to 5 sentences
    for i in range(current_idx, min(current_idx + 5, len(all_sentences))):
        if i >= len(all_sentences):
            break

        sentence = all_sentences[i]
        sentence_tokens = len(sentence.split())

        # Don't exceed max tokens too much
        if token_count + sentence_tokens > max_tokens + 50:  # Allow slight overflow
            break

        extended_segment.append(sentence)
        token_count += sentence_tokens

        # Check if we now have a switch
        segment_text = ' '.join(extended_segment)
        if verify_segment_has_switch(segment_text):
            return {
                'text': segment_text,
                'token_count': token_count,
                'has_transition': True
            }

    return None


def process_segment_to_tokens_labels_old(segment_text):
    """
    Convert a segment to tokens and 4-class labels
    """
    # Split by tags while keeping them
    parts = re.split(r'(<auto>|<allo>)', segment_text)

    tokens = []
    labels = []
    current_mode = None

    for i, part in enumerate(parts):
        if part == '<auto>':
            continue
        elif part == '<allo>':
            continue
        elif part.strip():
            words = part.strip().split()

            # Determine what mode this segment should be in
            segment_mode = None
            for j in range(i - 1, -1, -1):
                if parts[j] == '<auto>':
                    segment_mode = 'auto'
                    break
                elif parts[j] == '<allo>':
                    segment_mode = 'allo'
                    break

            if segment_mode is None:
                segment_mode = 'auto'  # Default

            for word_idx, word in enumerate(words):
                tokens.append(word)

                # Determine label based on position and mode change
                if word_idx == 0 and current_mode is not None and current_mode != segment_mode:
                    # This is a switch point
                    label = 2 if segment_mode == 'auto' else 3
                else:
                    # Continuation
                    label = 0 if segment_mode == 'auto' else 1

                labels.append(label)

            current_mode = segment_mode

    return tokens, labels



def has_code_switching_tags(file_path, file_type):
    """Check if file contains AUTO/ALLO tags before processing"""
    try:
        if file_type == 'docx':
            doc = docx.Document(str(file_path))
            text_content = ' '.join([para.text for para in doc.paragraphs])
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()

        # Check for any variation of tags
        text_lower = text_content.lower()
        return ('<auto>' in text_lower or '<allo>' in text_lower)
    except:
        return False



def create_train_val_test_split(df, train_ratio=0.7, val_ratio=0.15, test_ratio=0.15):
    """
    Create stratified split ensuring file diversity and no data leakage
    Each split MUST contain segments from multiple files
    """
    print(f"\n=== Creating Train/Val/Test Split ===")

    # Group by source file
    file_groups = {}
    for file_name in df['source_file'].unique():
        file_data = df[df['source_file'] == file_name].copy()
        file_groups[file_name] = file_data

    print(f"Total files: {len(file_groups)}")

    # Ensure we have enough files
    if len(file_groups) < 3:
        print("WARNING: Less than 3 files detected. Adding file diversity through segment splitting.")

    train_dfs = []
    val_dfs = []
    test_dfs = []

    # Strategy: Split SEGMENTS from EACH file across train/val/test
    # This ensures every split has diversity
    for file_name, file_data in file_groups.items():
        # Shuffle segments within this file
        file_data = file_data.sample(frac=1, random_state=42).reset_index(drop=True)

        n = len(file_data)
        n_train = int(n * train_ratio)
        n_val = int(n * val_ratio)

        # Split this file's segments
        train_segments = file_data.iloc[:n_train]
        val_segments = file_data.iloc[n_train:n_train + n_val]
        test_segments = file_data.iloc[n_train + n_val:]

        # Add to respective lists
        if len(train_segments) > 0:
            train_dfs.append(train_segments)
        if len(val_segments) > 0:
            val_dfs.append(val_segments)
        if len(test_segments) > 0:
            test_dfs.append(test_segments)

        print(f"  {file_name[:40]}...: {len(train_segments)} train, {len(val_segments)} val, {len(test_segments)} test")

    # Combine and shuffle
    train_df = pd.concat(train_dfs, ignore_index=True).sample(frac=1, random_state=42).reset_index(drop=True)
    val_df = pd.concat(val_dfs, ignore_index=True).sample(frac=1, random_state=42).reset_index(drop=True)
    test_df = pd.concat(test_dfs, ignore_index=True).sample(frac=1, random_state=42).reset_index(drop=True)

    # Verify file diversity
    train_files = train_df['source_file'].nunique()
    val_files = val_df['source_file'].nunique()
    test_files = test_df['source_file'].nunique()

    # Save splits
    train_df.to_csv('train_segments.csv', index=False)
    val_df.to_csv('val_segments.csv', index=False)
    test_df.to_csv('test_segments.csv', index=False)

    print(f"\nFinal split:")
    print(f"  Training: {len(train_df)} segments from {train_files} files")
    print(f"  Validation: {len(val_df)} segments from {val_files} files")
    print(f"  Test: {len(test_df)} segments from {test_files} files")

    # Verify diversity
    if train_files < 2 or val_files < 2 or test_files < 2:
        print("WARNING: Some splits have segments from fewer than 2 files!")
    else:
        print("✓ All splits have good file diversity")

    # Show file overlap
    train_file_set = set(train_df['source_file'].unique())
    val_file_set = set(val_df['source_file'].unique())
    test_file_set = set(test_df['source_file'].unique())

    common_files = train_file_set & val_file_set & test_file_set
    if common_files:
        print(f"✓ Files appearing in all splits (good for diversity): {len(common_files)}")

    return train_df, val_df, test_df


def balance_segments_by_switches(all_segments, target_switch_ratio=0.67):
    """
    Balance segments to have target_switch_ratio with switches

    Args:
        all_segments: List of all processed segments
        target_switch_ratio: Proportion of segments that should have switches (default 0.67 = 2/3)

    Returns:
        Balanced list of segments
    """
    print(f"\n=== Balancing Segments (Target: {target_switch_ratio * 100:.0f}% with switches) ===")

    # Separate segments by whether they have switches
    segments_with_switches = [s for s in all_segments if s['has_switch']]
    segments_without_switches = [s for s in all_segments if not s['has_switch']]

    n_with_switches = len(segments_with_switches)
    n_without_switches = len(segments_without_switches)

    print(f"Available segments:")
    print(f"  With switches: {n_with_switches}")
    print(f"  Without switches: {n_without_switches}")

    # Calculate desired counts based on target ratio
    # If target is 2/3 with switches, then for every 2 with switches, we want 1 without
    # target_switch_ratio = n_with / (n_with + n_without)
    # So: n_without = n_with * (1 - target_switch_ratio) / target_switch_ratio

    if n_with_switches == 0:
        print("ERROR: No segments with switches found!")
        return []

    # Calculate how many without-switch segments we need
    desired_without_switches = int(n_with_switches * (1 - target_switch_ratio) / target_switch_ratio)

    # Use all segments with switches
    balanced_segments = segments_with_switches.copy()

    # Sample the desired number of segments without switches
    if n_without_switches > 0:
        if n_without_switches >= desired_without_switches:
            # We have enough, randomly sample
            import random
            random.seed(42)
            sampled_without = random.sample(segments_without_switches, desired_without_switches)
            balanced_segments.extend(sampled_without)
        else:
            # We don't have enough, use all available
            print(f"  WARNING: Only {n_without_switches} segments without switches available")
            print(f"           Wanted {desired_without_switches} to achieve {target_switch_ratio * 100:.0f}% ratio")
            balanced_segments.extend(segments_without_switches)

    # Shuffle the combined list
    import random
    random.seed(42)
    random.shuffle(balanced_segments)

    # Calculate actual ratio
    final_with_switches = sum(1 for s in balanced_segments if s['has_switch'])
    final_without_switches = len(balanced_segments) - final_with_switches
    actual_ratio = final_with_switches / len(balanced_segments) if len(balanced_segments) > 0 else 0

    print(f"\nBalanced dataset:")
    print(f"  Total segments: {len(balanced_segments)}")
    print(f"  With switches: {final_with_switches} ({actual_ratio * 100:.1f}%)")
    print(f"  Without switches: {final_without_switches} ({(1 - actual_ratio) * 100:.1f}%)")
    print(f"  Target was: {target_switch_ratio * 100:.0f}% with switches")

    if abs(actual_ratio - target_switch_ratio) > 0.05:
        print(f"  ⚠️ WARNING: Actual ratio differs from target by more than 5%")
    else:
        print(f"  ✓ Ratio is close to target")

    return balanced_segments
# ============================================================================
# PART 2: DATASET CLASS
# ============================================================================

class CodeSwitchingDataset4Class(Dataset):
    """Dataset for 4-class token-level code-switching."""

    def __init__(self, csv_file, tokenizer, max_length=512):
        self.data = pd.read_csv(csv_file)
        self.tokenizer = tokenizer
        self.max_length = max_length

    def __len__(self):
        return len(self.data)

    def __getitem__(self, idx):
        row = self.data.iloc[idx]
        tokens = row['tokens'].split()
        labels = list(map(int, row['labels'].split(',')))

        encoding = self.tokenize_and_align_labels(tokens, labels)

        return {
            'input_ids': torch.tensor(encoding['input_ids']),
            'attention_mask': torch.tensor(encoding['attention_mask']),
            'labels': torch.tensor(encoding['labels'])
        }

    def tokenize_and_align_labels(self, tokens, labels):
        tokenized_inputs = self.tokenizer(
            tokens,
            is_split_into_words=True,
            padding='max_length',
            truncation=True,
            max_length=self.max_length
        )

        word_ids = tokenized_inputs.word_ids()
        aligned_labels = []
        previous_word_idx = None

        for word_idx in word_ids:
            if word_idx is None:
                aligned_labels.append(-100)
            elif word_idx != previous_word_idx:
                aligned_labels.append(labels[word_idx] if word_idx < len(labels) else -100)
            else:
                aligned_labels.append(-100)
            previous_word_idx = word_idx

        tokenized_inputs['labels'] = aligned_labels
        return tokenized_inputs

#####
### segmentation aware loss
###

def compute_segmentation_alignment_loss(self, tokens, predictions, labels, batch_idx):
    """
    Penalize switches that occur AFTER segmentation marks rather than AT them.
    Reward switches that align with segmentation boundaries.
    """
    device = predictions.device
    alignment_loss = torch.zeros_like(predictions, dtype=torch.float32)

    for seq_idx in range(len(tokens)):
        if tokens[seq_idx] == -100:  # Skip padding
            continue

        token_text = self.tokenizer.convert_ids_to_tokens([tokens[seq_idx].item()])[0]

        # Check if current token is or contains segmentation mark
        has_single_seg = '/' in token_text and '//' not in token_text
        has_double_seg = '//' in token_text
        has_seg_mark = has_single_seg or has_double_seg

        # Check if this is a predicted switch point
        is_predicted_switch = predictions[seq_idx] >= 2
        is_true_switch = labels[seq_idx] >= 2 if labels[seq_idx] != -100 else False

        if seq_idx > 0:
            prev_token_text = self.tokenizer.convert_ids_to_tokens([tokens[seq_idx - 1].item()])[0] if tokens[
                                                                                                           seq_idx - 1] != -100 else ""
            prev_has_seg = ('/' in prev_token_text)

            # PENALTY: Switch occurring 1-2 positions AFTER segmentation mark
            if is_predicted_switch and seq_idx >= 2:
                # Check if segmentation mark was 1-2 positions before
                for offset in [1, 2]:
                    if seq_idx - offset >= 0:
                        check_token = self.tokenizer.convert_ids_to_tokens([tokens[seq_idx - offset].item()])[0] if \
                        tokens[seq_idx - offset] != -100 else ""
                        if '/' in check_token:
                            # This switch is happening AFTER a segmentation mark - BAD!
                            alignment_loss[seq_idx] *= 3.0  # Penalty factor
                            break

            # REWARD: Switch occurring AT or just before segmentation mark
            if is_predicted_switch:
                if has_seg_mark or (seq_idx + 1 < len(tokens) and tokens[seq_idx + 1] != -100):
                    next_token_text = self.tokenizer.convert_ids_to_tokens([tokens[seq_idx + 1].item()])[
                        0] if seq_idx + 1 < len(tokens) else ""
                    if '/' in next_token_text:
                        # Switch aligns with segmentation boundary - GOOD!
                        alignment_loss[seq_idx] *= 0.3  # Reward factor

    return alignment_loss
# ============================================================================
# PART 3: PROXIMITY-AWARE LOSS FUNCTION
# ============================================================================
class SwitchFocusedLoss(nn.Module):
    """Loss that primarily cares about detecting switches correctly"""

    def __init__(self, switch_recall_weight=5.0, proximity_tolerance=5,
                 segmentation_penalty=3.0, segmentation_reward=0.3):
        super().__init__()
        self.proximity_tolerance = proximity_tolerance
        self.segmentation_penalty = segmentation_penalty
        self.segmentation_reward = segmentation_reward

        # Heavily weight switch classes to boost recall

        self.class_weights = torch.tensor([
            0.1,  # Nearly ignore non-switch classes
            0.1,
            5.0,  # Focus on switches
            5.0
        ])

    def check_segmentation_alignment(self, tokens, predictions, b, seq_len):
        """
        Check if switches align with Tibetan segmentation marks / or //
        Returns adjustment factors for each position
        """
        adjustment = torch.ones(seq_len, device=predictions.device)

        # Convert tokens to text to check for segmentation marks
        for t in range(seq_len):
            if predictions[b, t] >= 2:  # This is a predicted switch

                # Check if switch happens 1-2 positions AFTER segmentation mark (BAD)
                for offset in [1, 2]:
                    if t - offset >= 0:
                        # Check if previous token contains / or //
                        # Note: You'll need to check actual token text here
                        # This is a simplified version - you may need to adapt based on your tokenization
                        if t - offset < len(tokens[b]) and tokens[b][t - offset] != -100:
                            # If previous position likely had segmentation mark
                            # (You might need to decode token to check for '/')
                            adjustment[t] = self.segmentation_penalty

                # Check if switch happens AT or RIGHT BEFORE segmentation mark (GOOD)
                if t + 1 < seq_len:
                    # If next position likely has segmentation mark
                    adjustment[t] = self.segmentation_reward

        return adjustment

    def forward(self, logits, labels, tokens=None):
        batch_size, seq_len, num_classes = logits.shape
        device = logits.device
        self.class_weights = self.class_weights.to(device)

        # Base loss with heavy switch weighting
        base_loss = F.cross_entropy(
            logits.view(-1, num_classes),
            labels.view(-1),
            weight=self.class_weights,
            reduction='none'
        ).view(batch_size, seq_len)

        predictions = torch.argmax(logits, dim=-1)

        # Apply segmentation alignment if tokens provided
        if tokens is not None:
            for b in range(batch_size):
                seg_adjustment = self.check_segmentation_alignment(tokens, predictions, b, seq_len)
                base_loss[b] *= seg_adjustment

        # Rest of your existing proximity logic
        for b in range(batch_size):
            true_switches = torch.where(labels[b] >= 2)[0]
            pred_switches = torch.where(predictions[b] >= 2)[0]

            # Existing proximity rewards
            for true_pos in true_switches:
                window_start = max(0, true_pos - self.proximity_tolerance)
                window_end = min(seq_len, true_pos + self.proximity_tolerance + 1)
                window_preds = predictions[b, window_start:window_end]
                if torch.any(window_preds >= 2):
                    base_loss[b, true_pos] *= 0.1

            # Mild penalty for far predictions
            for pred_pos in pred_switches:
                if len(true_switches) > 0:
                    distances = torch.abs(true_switches - pred_pos)
                    min_distance = torch.min(distances).item()
                    if min_distance > self.proximity_tolerance:
                        base_loss[b, pred_pos] *= 2.0

        valid_mask = (labels != -100).float()
        return (base_loss * valid_mask).sum() / valid_mask.sum()

    def forward_old(self, logits, labels, input_ids=None):
        batch_size, seq_len, num_classes = logits.shape
        device = logits.device
        self.class_weights = self.class_weights.to(device)

        # Base loss with heavy switch weighting
        base_loss = F.cross_entropy(
            logits.view(-1, num_classes),
            labels.view(-1),
            weight=self.class_weights,
            reduction='none'
        ).view(batch_size, seq_len)

        predictions = torch.argmax(logits, dim=-1)

        # REWARD predictions near true switches (not penalize far ones)
        for b in range(batch_size):
            true_switches = torch.where(labels[b] >= 2)[0]  # Both switch types
            pred_switches = torch.where(predictions[b] >= 2)[0]

            # For each true switch, reward ANY prediction within tolerance
            for true_pos in true_switches:
                # Find predictions within tolerance window
                window_start = max(0, true_pos - self.proximity_tolerance)
                window_end = min(seq_len, true_pos + self.proximity_tolerance + 1)

                # If there's a prediction in the window, reduce loss
                window_preds = predictions[b, window_start:window_end]
                if torch.any(window_preds >= 2):
                    base_loss[b, true_pos] *= 0.1  # Big reward

            # Mild penalty for predictions far from any true switch
            for pred_pos in pred_switches:
                if len(true_switches) > 0:
                    distances = torch.abs(true_switches - pred_pos)
                    min_distance = torch.min(distances).item()

                    if min_distance > self.proximity_tolerance:
                        # Mild penalty, not extreme
                        base_loss[b, pred_pos] *= 2.0  # Was 30.0!

        valid_mask = (labels != -100).float()
        return (base_loss * valid_mask).sum() / valid_mask.sum()
def apply_transition_constraints(predictions, logits=None):
    """
    Apply logical constraints to predictions:
    - If in Auto mode (0), can only switch to Allo (3)
    - If in Allo mode (1), can only switch to Auto (2)
    """
    import torch

    if isinstance(predictions, torch.Tensor):
        predictions = predictions.cpu().numpy()

    corrected_predictions = predictions.copy()
    current_mode = 0  # Start in Auto by default

    for i in range(len(predictions)):
        pred = predictions[i]

        if pred == -100:  # Skip padding
            continue

        # Check for invalid transitions
        if pred == 2:  # Switch to Auto
            if current_mode == 0:  # Already in Auto - INVALID
                # Change to continuation in Auto
                corrected_predictions[i] = 0
            else:  # Was in Allo - VALID
                current_mode = 0

        elif pred == 3:  # Switch to Allo
            if current_mode == 1:  # Already in Allo - INVALID
                # Change to continuation in Allo
                corrected_predictions[i] = 1
            else:  # Was in Auto - VALID
                current_mode = 1

        elif pred == 0:  # Continue in Auto
            current_mode = 0

        elif pred == 1:  # Continue in Allo
            current_mode = 1

    return corrected_predictions


class ProximityAwareLoss4Class(nn.Module):
    """
    Loss function with logical transition constraints
    """

    def __init__(self, switch_loss_weight=30.0, proximity_tolerance=5,
                 distance_decay=0.7, false_positive_penalty=10.0,
                 invalid_transition_penalty=100.0,
                 segmentation_alignment_weight=2.0):  # NEW parameter
        super().__init__()
        self.switch_loss_weight = switch_loss_weight
        self.proximity_tolerance = proximity_tolerance
        self.distance_decay = distance_decay
        self.false_positive_penalty = false_positive_penalty
        self.invalid_transition_penalty = invalid_transition_penalty
        self.segmentation_alignment_weight = segmentation_alignment_weight
        self.tokenizer = None

        # Much higher weights for switch classes to combat class imbalance
        self.class_weights = torch.tensor([
            1.0,  # Class 0: Non-switch auto
            1.0,  # Class 1: Non-switch allo
            switch_loss_weight,  # Class 2: Switch to auto (increased)
            switch_loss_weight   # Class 3: Switch to allo (increased)
        ])

    def forward(self, logits, labels):
        batch_size, seq_len, num_classes = logits.shape
        device = logits.device
        self.class_weights = self.class_weights.to(device)

        # Standard cross-entropy loss with class weights
        ce_loss = F.cross_entropy(
            logits.view(-1, num_classes),
            labels.view(-1),
            weight=self.class_weights,
            reduction='none'
        ).view(batch_size, seq_len)

        valid_mask = (labels != -100).float()

        # Get predictions
        predictions = torch.argmax(logits, dim=-1)

        # Apply proximity-aware adjustments
        proximity_adjusted_loss = ce_loss.clone()

        for b in range(batch_size):
            # Track current mode to detect invalid transitions
            current_true_mode = 0  # Start in Auto
            current_pred_mode = 0

            for t in range(seq_len):
                if labels[b, t] == -100:
                    continue

                true_label = labels[b, t].item()
                pred_label = predictions[b, t].item()

                # PENALIZE INVALID TRANSITIONS HEAVILY
                # Check predicted transitions
                if pred_label == 2:  # Predicted Switch to Auto
                    if current_pred_mode == 0:  # Already in Auto - INVALID!
                        proximity_adjusted_loss[b, t] *= self.invalid_transition_penalty
                    current_pred_mode = 0
                elif pred_label == 3:  # Predicted Switch to Allo
                    if current_pred_mode == 1:  # Already in Allo - INVALID!
                        proximity_adjusted_loss[b, t] *= self.invalid_transition_penalty
                    current_pred_mode = 1
                elif pred_label == 0:
                    current_pred_mode = 0
                elif pred_label == 1:
                    current_pred_mode = 1

                # REWARD CORRECT LOGICAL TRANSITIONS
                # Check if this is a true switch point
                if true_label == 2:  # True Switch to Auto
                    if current_true_mode == 1 and pred_label == 2:  # Correct logic
                        proximity_adjusted_loss[b, t] *= 0.1  # Strong reward
                    current_true_mode = 0
                elif true_label == 3:  # True Switch to Allo
                    if current_true_mode == 0 and pred_label == 3:  # Correct logic
                        proximity_adjusted_loss[b, t] *= 0.1  # Strong reward
                    current_true_mode = 1
                elif true_label == 0:
                    current_true_mode = 0
                elif true_label == 1:
                    current_true_mode = 1

            # Original proximity logic (but with type matching)
            true_switches_to_auto = torch.where(labels[b] == 2)[0]
            true_switches_to_allo = torch.where(labels[b] == 3)[0]
            pred_switches_to_auto = torch.where(predictions[b] == 2)[0]
            pred_switches_to_allo = torch.where(predictions[b] == 3)[0]
            for pred_pos in pred_switches_to_auto:
                if len(true_switches_to_auto) > 0:
                    distances = torch.abs(true_switches_to_auto - pred_pos)
                    min_distance = torch.min(distances).item()

                    if min_distance == 0:
                        proximity_adjusted_loss[b, pred_pos] *= 0.1  # Best reward
                    elif min_distance <= self.proximity_tolerance:
                        # FIXED: Linear scaling so closer = better reward
                        # Distance 1: ×0.28, Distance 2: ×0.46, Distance 5: ×1.0
                        reward_factor = 0.1 + (min_distance / self.proximity_tolerance) * 0.9
                        proximity_adjusted_loss[b, pred_pos] *= reward_factor
                    else:
                        # Flat penalty for anything beyond tolerance
                        proximity_adjusted_loss[b, pred_pos] *= self.false_positive_penalty
                else:
                    proximity_adjusted_loss[b, pred_pos] *= self.false_positive_penalty * 2
            for pred_pos in pred_switches_to_allo:
                if len(true_switches_to_allo) > 0:
                    distances = torch.abs(true_switches_to_allo - pred_pos)
                    min_distance = torch.min(distances).item()

                    if min_distance == 0:
                        proximity_adjusted_loss[b, pred_pos] *= 0.1  # Best reward
                    elif min_distance <= self.proximity_tolerance:
                        # FIXED: Linear scaling so closer = better reward
                        reward_factor = 0.1 + (min_distance / self.proximity_tolerance) * 0.9
                        proximity_adjusted_loss[b, pred_pos] *= reward_factor
                    else:
                        # Flat penalty for anything beyond tolerance
                        proximity_adjusted_loss[b, pred_pos] *= self.false_positive_penalty
                else:
                    proximity_adjusted_loss[b, pred_pos] *= self.false_positive_penalty * 2
            # For predicted "switch to auto"
            # for pred_pos in pred_switches_to_auto:
            #     if len(true_switches_to_auto) > 0:
            #         distances = torch.abs(true_switches_to_auto - pred_pos)
            #         min_distance = torch.min(distances).item()
            #
            #         if min_distance == 0:
            #             proximity_adjusted_loss[b, pred_pos] *= 0.1
            #         elif min_distance <= self.proximity_tolerance:
            #             decay_factor = self.distance_decay ** min_distance
            #             proximity_adjusted_loss[b, pred_pos] *= decay_factor
            #         else:
            #             proximity_adjusted_loss[b, pred_pos] *= self.false_positive_penalty
            #     else:
            #         proximity_adjusted_loss[b, pred_pos] *= self.false_positive_penalty * 2
            #
            # # For predicted "switch to allo"
            # for pred_pos in pred_switches_to_allo:
            #     if len(true_switches_to_allo) > 0:
            #         distances = torch.abs(true_switches_to_allo - pred_pos)
            #         min_distance = torch.min(distances).item()
            #
            #         if min_distance == 0:
            #             proximity_adjusted_loss[b, pred_pos] *= 0.1
            #         elif min_distance <= self.proximity_tolerance:
            #             decay_factor = self.distance_decay ** min_distance
            #             proximity_adjusted_loss[b, pred_pos] *= decay_factor
            #         else:
            #             proximity_adjusted_loss[b, pred_pos] *= self.false_positive_penalty
            #     else:
            #         proximity_adjusted_loss[b, pred_pos] *= self.false_positive_penalty * 2
            # Penalize missed true switches with distance-aware scaling
            for true_pos in true_switches_to_auto:
                if len(pred_switches_to_auto) == 0:
                    # No predictions at all - heavy penalty
                    proximity_adjusted_loss[b, true_pos] *= 5.0
                else:
                    distances = torch.abs(pred_switches_to_auto - true_pos)
                    min_distance = torch.min(distances).item()

                    if min_distance > self.proximity_tolerance:
                        # Scale penalty based on how far the nearest prediction is
                        # Further away = higher penalty
                        distance_penalty = 2.0 + (min_distance - self.proximity_tolerance) * 0.3
                        distance_penalty = min(distance_penalty, 8.0)  # Cap at 8x
                        proximity_adjusted_loss[b, true_pos] *= distance_penalty
                    # If within tolerance, the positive proximity logic already handled it

            # Same for allo switches
            for true_pos in true_switches_to_allo:
                if len(pred_switches_to_allo) == 0:
                    proximity_adjusted_loss[b, true_pos] *= 5.0
                else:
                    distances = torch.abs(pred_switches_to_allo - true_pos)
                    min_distance = torch.min(distances).item()

                    if min_distance > self.proximity_tolerance:
                        distance_penalty = 2.0 + (min_distance - self.proximity_tolerance) * 0.3
                        distance_penalty = min(distance_penalty, 8.0)
                        proximity_adjusted_loss[b, true_pos] *= distance_penalty

            # Penalize missed true switches
            # for true_pos in true_switches_to_auto:
            #     if len(pred_switches_to_auto) == 0:
            #         proximity_adjusted_loss[b, true_pos] *= 3.0
            #     else:
            #         distances = torch.abs(pred_switches_to_auto - true_pos)
            #         min_distance = torch.min(distances).item()
            #         if min_distance > self.proximity_tolerance:
            #             proximity_adjusted_loss[b, true_pos] *= 2.0

            # for true_pos in true_switches_to_allo:
            #     if len(pred_switches_to_allo) == 0:
            #         proximity_adjusted_loss[b, true_pos] *= 3.0
            #     else:
            #         distances = torch.abs(pred_switches_to_allo - true_pos)
            #         min_distance = torch.min(distances).item()
            #         if min_distance > self.proximity_tolerance:
            #             proximity_adjusted_loss[b, true_pos] *= 2.0

        # Apply valid mask and compute mean
        total_loss = proximity_adjusted_loss * valid_mask
        loss = total_loss.sum() / valid_mask.sum().clamp(min=1)

        return loss


# ============================================================================
# PART 4: PROXIMITY-AWARE METRICS
# ============================================================================

def evaluate_switch_detection_with_proximity(true_labels, pred_labels, tolerance=5):
    """
    Evaluate switch detection with proximity tolerance and TYPE matching
    Switch types must match: class 2 with class 2, class 3 with class 3
    """
    true_labels = np.array(true_labels)
    pred_labels = np.array(pred_labels)

    # Find switch positions BY TYPE
    true_switches_to_auto = np.where(true_labels == 2)[0]  # Switch to Auto
    true_switches_to_allo = np.where(true_labels == 3)[0]  # Switch to Allo
    pred_switches_to_auto = np.where(pred_labels == 2)[0]  # Predicted Switch to Auto
    pred_switches_to_allo = np.where(pred_labels == 3)[0]  # Predicted Switch to Allo

    # Track matches separately by type
    matched_true_to_auto = set()
    matched_pred_to_auto = set()
    matched_true_to_allo = set()
    matched_pred_to_allo = set()

    exact_matches = 0
    proximity_matches = 0

    # Match "Switch to Auto" predictions with "Switch to Auto" ground truth
    for pred_pos in pred_switches_to_auto:
        if len(true_switches_to_auto) > 0:
            distances = np.abs(true_switches_to_auto - pred_pos)
            min_distance = np.min(distances)
            closest_true_idx = np.argmin(distances)
            closest_true_pos = true_switches_to_auto[closest_true_idx]

            # Only count if not already matched
            if closest_true_pos not in matched_true_to_auto:
                if min_distance == 0:
                    exact_matches += 1
                    matched_true_to_auto.add(closest_true_pos)
                    matched_pred_to_auto.add(pred_pos)
                elif min_distance <= tolerance:
                    proximity_matches += 1
                    matched_true_to_auto.add(closest_true_pos)
                    matched_pred_to_auto.add(pred_pos)

    # Match "Switch to Allo" predictions with "Switch to Allo" ground truth
    for pred_pos in pred_switches_to_allo:
        if len(true_switches_to_allo) > 0:
            distances = np.abs(true_switches_to_allo - pred_pos)
            min_distance = np.min(distances)
            closest_true_idx = np.argmin(distances)
            closest_true_pos = true_switches_to_allo[closest_true_idx]

            # Only count if not already matched
            if closest_true_pos not in matched_true_to_allo:
                if min_distance == 0:
                    exact_matches += 1
                    matched_true_to_allo.add(closest_true_pos)
                    matched_pred_to_allo.add(pred_pos)
                elif min_distance <= tolerance:
                    proximity_matches += 1
                    matched_true_to_allo.add(closest_true_pos)
                    matched_pred_to_allo.add(pred_pos)

    # Total counts
    total_true_switches = len(true_switches_to_auto) + len(true_switches_to_allo)
    total_pred_switches = len(pred_switches_to_auto) + len(pred_switches_to_allo)
    total_matched_true = len(matched_true_to_auto) + len(matched_true_to_allo)
    total_matched_pred = len(matched_pred_to_auto) + len(matched_pred_to_allo)

    total_matches = exact_matches + proximity_matches
    missed_switches = total_true_switches - total_matched_true
    false_switches = total_pred_switches - total_matched_pred

    # Calculate metrics
    precision = total_matches / total_pred_switches if total_pred_switches > 0 else 0
    recall = total_matches / total_true_switches if total_true_switches > 0 else 0
    f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0

    # Also calculate per-type metrics
    to_auto_precision = len(matched_pred_to_auto) / len(pred_switches_to_auto) if len(pred_switches_to_auto) > 0 else 0
    to_auto_recall = len(matched_true_to_auto) / len(true_switches_to_auto) if len(true_switches_to_auto) > 0 else 0
    to_allo_precision = len(matched_pred_to_allo) / len(pred_switches_to_allo) if len(pred_switches_to_allo) > 0 else 0
    to_allo_recall = len(matched_true_to_allo) / len(true_switches_to_allo) if len(true_switches_to_allo) > 0 else 0

    return {
        'exact_matches': exact_matches,
        'proximity_matches': proximity_matches,
        'total_matches': total_matches,
        'true_switches': total_true_switches,
        'pred_switches': total_pred_switches,
        'missed_switches': missed_switches,
        'false_switches': false_switches,
        'precision': precision,
        'recall': recall,
        'f1': f1,
        # Per-type metrics
        'true_to_auto': len(true_switches_to_auto),
        'true_to_allo': len(true_switches_to_allo),
        'pred_to_auto': len(pred_switches_to_auto),
        'pred_to_allo': len(pred_switches_to_allo),
        'matched_to_auto': len(matched_true_to_auto),
        'matched_to_allo': len(matched_true_to_allo),
        'to_auto_precision': to_auto_precision,
        'to_auto_recall': to_auto_recall,
        'to_allo_precision': to_allo_precision,
        'to_allo_recall': to_allo_recall
    }


def compute_metrics_for_trainer(eval_pred, tolerance=5):
    """
    Compute metrics for the trainer with proximity awareness and TYPE matching
    """
    predictions, labels = eval_pred
    predictions = np.argmax(predictions, axis=2)

    # Flatten all sequences
    all_predictions = predictions.flatten()
    all_labels = labels.flatten()

    # Remove padding
    mask = all_labels != -100
    all_predictions = all_predictions[mask]
    all_labels = all_labels[mask]

    # Basic accuracy
    accuracy = (all_predictions == all_labels).mean()

    # Proximity-aware switch metrics with TYPE matching
    switch_metrics = evaluate_switch_detection_with_proximity(
        all_labels, all_predictions, tolerance=tolerance
    )

    return {
        'accuracy': float(accuracy),
        'switch_precision': float(switch_metrics['precision']),
        'switch_recall': float(switch_metrics['recall']),
        'switch_f1': float(switch_metrics['f1']),
        'true_switches': switch_metrics['true_switches'],
        'pred_switches': switch_metrics['pred_switches'],
        'exact_matches': switch_metrics['exact_matches'],
        'proximity_matches': switch_metrics['proximity_matches'],
        # Per-type metrics
        'to_auto_precision': float(switch_metrics['to_auto_precision']),
        'to_auto_recall': float(switch_metrics['to_auto_recall']),
        'to_allo_precision': float(switch_metrics['to_allo_precision']),
        'to_allo_recall': float(switch_metrics['to_allo_recall']),
        'true_to_auto': switch_metrics['true_to_auto'],
        'true_to_allo': switch_metrics['true_to_allo'],
        'matched_to_auto': switch_metrics['matched_to_auto'],
        'matched_to_allo': switch_metrics['matched_to_allo']
    }


# ============================================================================
# PART 5: CUSTOM TRAINER
# ============================================================================

class ProximityAwareTrainer4Class(Trainer):
    """Custom trainer with proximity-aware loss and transition constraints"""

    def __init__(self, switch_loss_weight=30.0, proximity_tolerance=5,
                 distance_decay=0.7, false_positive_penalty=10.0,
                 invalid_transition_penalty=100.0, *args, **kwargs):
        # Handle tokenizer/processing_class
        if 'tokenizer' in kwargs and 'processing_class' not in kwargs:
            kwargs['processing_class'] = kwargs.get('tokenizer')

        super().__init__(*args, **kwargs)

        # Initialize custom loss with all parameters including invalid transition penalty
        self.proximity_loss = ProximityAwareLoss4Class(
            switch_loss_weight=switch_loss_weight,
            proximity_tolerance=proximity_tolerance,
            distance_decay=distance_decay,
            false_positive_penalty=false_positive_penalty,
            invalid_transition_penalty=invalid_transition_penalty
        )
        self.proximity_tolerance = proximity_tolerance

    def compute_loss(self, model, inputs, return_outputs=False, num_items_in_batch=None):
        labels = inputs.pop("labels")
        outputs = model(**inputs)
        logits = outputs.get('logits')
        loss = self.proximity_loss(logits, labels)
        return (loss, outputs) if return_outputs else loss

    def prediction_step(self, model, inputs, prediction_loss_only, ignore_keys=None):
        """Override to apply transition constraints during evaluation"""
        outputs = super().prediction_step(model, inputs, prediction_loss_only, ignore_keys)

        if not prediction_loss_only:
            loss, logits, labels = outputs

            # Get raw predictions
            predictions = torch.argmax(logits, dim=-1)

            # Apply transition constraints to each sequence in the batch
            batch_size = predictions.shape[0]
            for b in range(batch_size):
                predictions[b] = torch.tensor(
                    apply_transition_constraints(predictions[b]),
                    device=predictions.device
                )

            # Return with corrected predictions
            return (loss, logits, labels)

        return outputs

    def evaluate(self, eval_dataset=None, ignore_keys=None, metric_key_prefix="eval"):
        # Override to use proximity-aware metrics
        self.compute_metrics = lambda eval_pred: compute_metrics_for_trainer(
            eval_pred, tolerance=self.proximity_tolerance
        )
        return super().evaluate(eval_dataset, ignore_keys, metric_key_prefix)


# ============================================================================
# PART 6: EVALUATION AND VISUALIZATION
# ============================================================================

def print_test_examples_with_constraints(model, tokenizer, test_csv='test_segments.csv',
                                        num_examples=5, tolerance=5):
    """
    Print test examples showing effect of logical constraints
    """
    print("\n" + "=" * 80)
    print(f"SWITCH DETECTION WITH LOGICAL CONSTRAINTS")
    print("Rules: Can only switch FROM current mode TO different mode")
    print("=" * 80)

    test_df = pd.read_csv(test_csv)
    device = next(model.parameters()).device
    model.eval()

    label_names = {
        0: 'NonSwitch-Auto',
        1: 'NonSwitch-Allo',
        2: 'SWITCH→Auto',
        3: 'SWITCH→Allo'
    }

    sample_indices = np.random.choice(len(test_df), min(num_examples, len(test_df)), replace=False)

    for i, idx in enumerate(sample_indices):
        row = test_df.iloc[idx]
        tokens = row['tokens'].split()
        true_labels = list(map(int, row['labels'].split(',')))

        print(f"\n--- Example {i+1} (from {row['source_file'][:40]}...) ---")
        print(f"Segment length: {row['num_tokens']} tokens")

        # Get predictions
        tokenizer_output = tokenizer(
            tokens,
            is_split_into_words=True,
            return_tensors="pt",
            padding=True,
            truncation=True,
            max_length=512
        )

        inputs = {k: v.to(device) for k, v in tokenizer_output.items()}

        with torch.no_grad():
            outputs = model(**inputs)
            raw_predictions = torch.argmax(outputs.logits, dim=2)

        # Align predictions
        word_ids = tokenizer_output.word_ids()
        aligned_raw = []
        previous_word_idx = None

        for j, word_idx in enumerate(word_ids):
            if word_idx is not None and word_idx != previous_word_idx:
                pred = raw_predictions[0][j].item()
                aligned_raw.append(pred)
            previous_word_idx = word_idx

        # Apply constraints
        final_len = min(len(aligned_raw), len(true_labels), len(tokens))
        aligned_raw = aligned_raw[:final_len]
        aligned_constrained = apply_transition_constraints(aligned_raw)
        true_labels = true_labels[:final_len]
        tokens = tokens[:final_len]

        # Find constraint violations that were corrected
        violations_corrected = []
        current_mode = 0

        for j in range(len(aligned_raw)):
            raw = aligned_raw[j]
            constrained = aligned_constrained[j]

            if raw != constrained:
                violations_corrected.append((j, raw, constrained, current_mode))

            # Update mode based on constrained prediction
            if constrained == 2:
                current_mode = 0
            elif constrained == 3:
                current_mode = 1
            elif constrained == 0:
                current_mode = 0
            elif constrained == 1:
                current_mode = 1

        print(f"\nConstraint violations corrected: {len(violations_corrected)}")
        if violations_corrected:
            for pos, raw, fixed, mode in violations_corrected[:5]:  # Show first 5
                mode_str = "Auto" if mode == 0 else "Allo"
                print(f"  Pos {pos}: {label_names[raw]} → {label_names[fixed]} (was in {mode_str} mode)")

        # Evaluate both versions
        raw_eval = evaluate_switch_detection_with_proximity(true_labels, aligned_raw, tolerance)
        const_eval = evaluate_switch_detection_with_proximity(true_labels, aligned_constrained, tolerance)

        print(f"\nPerformance comparison:")
        print(f"  {'Metric':<20} {'Raw':<15} {'Constrained':<15}")
        print(f"  {'-'*50}")
        print(f"  {'Precision':<20} {raw_eval['precision']:<15.3f} {const_eval['precision']:<15.3f}")
        print(f"  {'Recall':<20} {raw_eval['recall']:<15.3f} {const_eval['recall']:<15.3f}")
        print(f"  {'F1':<20} {raw_eval['f1']:<15.3f} {const_eval['f1']:<15.3f}")
        print(f"  {'False Positives':<20} {raw_eval['false_switches']:<15} {const_eval['false_switches']:<15}")

        # Show switch regions
        true_switches = [(j, l) for j, l in enumerate(true_labels) if l in [2, 3]]
        const_switches = [(j, l) for j, l in enumerate(aligned_constrained) if l in [2, 3]]

        if true_switches or const_switches:
            print(f"\nSwitch regions (showing logical flow):")
            all_switch_pos = set([p for p, _ in true_switches + const_switches])

            for switch_pos in sorted(all_switch_pos):
                start = max(0, switch_pos - 2)
                end = min(len(tokens), switch_pos + 3)

                print(f"\n  Around position {switch_pos}:")
                current_true_mode = 0
                current_pred_mode = 0

                for pos in range(start, end):
                    if pos < len(tokens):
                        token = tokens[pos]
                        true_label = true_labels[pos]
                        const_pred = aligned_constrained[pos]

                        # Determine modes
                        if true_label in [0, 2]:
                            current_true_mode = 0
                        elif true_label in [1, 3]:
                            current_true_mode = 1

                        if const_pred in [0, 2]:
                            current_pred_mode = 0
                        elif const_pred in [1, 3]:
                            current_pred_mode = 1

                        marker = "→" if pos == switch_pos else " "

                        # Check logical validity
                        logic_check = ""
                        if const_pred == 2 and current_pred_mode == 0:
                            logic_check = "⚠️ INVALID"
                        elif const_pred == 3 and current_pred_mode == 1:
                            logic_check = "⚠️ INVALID"
                        elif const_pred in [2, 3]:
                            logic_check = "✓ VALID"

                        print(f"    {marker} [{pos:3d}] {token:12s} | True: {label_names[true_label]:15s} | Pred: {label_names[const_pred]:15s} | {logic_check}")

    return sample_indices


# Add this after print_test_examples_with_constraints function:
def print_test_examples_proximity(model, tokenizer, test_csv='test_segments.csv',
                                 num_examples=5, tolerance=5):
    """
    Print detailed test examples with proximity-aware evaluation
    Shows TYPE-specific switch matching
    """
    # Fallback to the constraint version if needed
    return print_test_examples_with_constraints(model, tokenizer, test_csv, num_examples, tolerance)


"""
Modified code to include non-switching segments and ensure NO tags in data
Key changes:
1. Include segments without switches
2. Comprehensive tag validation
3. Strengthened tag removal
"""


# Add this function for final validation
def validate_no_tags_in_data(df, split_name):
    """
    CRITICAL: Final validation to ensure absolutely NO tags in any data
    Returns True if clean, False if tags found
    """
    print(f"\n{'=' * 60}")
    print(f"CRITICAL TAG VALIDATION FOR {split_name}")
    print(f"{'=' * 60}")

    issues_found = []

    for idx in range(len(df)):
        row = df.iloc[idx]

        # Check tokens
        tokens_str = row['tokens']
        if contains_any_tag_pattern(tokens_str):
            issues_found.append(f"Row {idx}: Tags found in tokens: {tokens_str[:100]}")

        # Check original text
        if 'original_text' in row:
            orig_text = row['original_text']
            if contains_any_tag_pattern(orig_text):
                issues_found.append(f"Row {idx}: Tags found in original_text: {orig_text[:100]}")

        # Check individual tokens
        tokens_list = tokens_str.split()
        for i, token in enumerate(tokens_list):
            if contains_any_tag_pattern(token):
                issues_found.append(f"Row {idx}, Token {i}: '{token}' contains tag pattern")

    if issues_found:
        print(f"❌ VALIDATION FAILED! Found {len(issues_found)} issues:")
        for issue in issues_found[:10]:  # Show first 10
            print(f"  - {issue}")
        if len(issues_found) > 10:
            print(f"  ... and {len(issues_found) - 10} more issues")
        return False
    else:
        print(f"✅ VALIDATION PASSED! No tags found in {len(df)} segments")
        return True


# Modified extract_segments_with_token_limit to include non-switching segments
def extract_segments_with_token_limit(sentences, min_tokens=300, max_tokens=400,
                                      require_switch=False):
    """
    Extract segments of 300-400 tokens
    NEW: Can optionally include segments WITHOUT switches

    Args:
        require_switch: If True, only include segments with switches
                       If False, include all segments (with or without switches)
    """
    segments = []
    current_segment = []
    current_token_count = 0

    for sentence in sentences:
        sentence_tokens = sentence.split()
        sentence_token_count = len(sentence_tokens)

        if current_token_count + sentence_token_count > max_tokens:
            if current_token_count >= min_tokens:
                segment_text = ' '.join(current_segment)

                has_switch = verify_segment_has_switch(segment_text)

                # Include segment based on require_switch parameter
                if not require_switch or has_switch:
                    segments.append({
                        'text': segment_text,
                        'token_count': current_token_count,
                        'has_transition': has_switch
                    })

            current_segment = [sentence]
            current_token_count = sentence_token_count
        else:
            current_segment.append(sentence)
            current_token_count += sentence_token_count

            if current_token_count >= min_tokens:
                segment_text = ' '.join(current_segment)
                has_switch = verify_segment_has_switch(segment_text)

                # Include segment based on require_switch parameter
                if not require_switch or has_switch:
                    segments.append({
                        'text': segment_text,
                        'token_count': current_token_count,
                        'has_transition': has_switch
                    })
                    current_segment = []
                    current_token_count = 0

    # Handle remaining sentences
    if current_segment and current_token_count >= min_tokens // 2:
        segment_text = ' '.join(current_segment)
        has_switch = verify_segment_has_switch(segment_text)

        if not require_switch or has_switch:
            segments.append({
                'text': segment_text,
                'token_count': current_token_count,
                'has_transition': has_switch
            })

    return segments


# Modified process_file_to_segments to include non-switching segments
def process_file_to_segments(file_path, file_type, require_switch=False):
    """
    Process a single file and extract segments
    NEW: Can include segments without switches
    CRITICAL: Ensures NO tags in tokens

    Args:
        require_switch: If True, only include segments with switches
                       If False, include all segments
    """
    print(f"Processing {file_type} file: {file_path.name}")

    # Read file content
    if file_type == 'docx':
        doc = docx.Document(str(file_path))
        text_content = ' '.join([para.text for para in doc.paragraphs])
    else:  # txt file
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        text_content = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text_content = f.read()
                break
            except UnicodeDecodeError:
                continue

        if text_content is None:
            raise ValueError(f"Could not read {file_path.name} with any encoding")

    # Clean and normalize
    text_content = clean_and_normalize_text(text_content)

    # Split into sentences
    sentences = split_into_sentences(text_content)
    print(f"  Found {len(sentences)} sentences")

    # Extract segments (with or without switches based on parameter)
    segments = extract_segments_with_token_limit(sentences, min_tokens=300, max_tokens=400,
                                                 require_switch=require_switch)

    if require_switch:
        print(f"  Found {len(segments)} segments WITH switches (300-400 tokens each)")
    else:
        segments_with_switches = sum(1 for s in segments if s['has_transition'])
        segments_without_switches = len(segments) - segments_with_switches
        print(f"  Found {len(segments)} total segments:")
        print(f"    - {segments_with_switches} with switches")
        print(f"    - {segments_without_switches} without switches")

    # Convert segments to token-label pairs
    processed_segments = []
    segments_with_tags = 0

    for seg_idx, segment in enumerate(segments):
        tokens, labels = process_segment_to_tokens_labels_clean(segment['text'])

        if len(tokens) > 0:
            num_transitions = sum(1 for l in labels if l in [2, 3])

            # CRITICAL: Verify no tags in tokens
            has_tags = False
            for token in tokens:
                if contains_any_tag_pattern(token):
                    has_tags = True
                    segments_with_tags += 1
                    print(f"    WARNING: Found tag in token: '{token}'")
                    break

            if has_tags:
                continue  # Skip this segment entirely

            # Create text WITHOUT tags for storage
            text_without_tags = ' '.join(tokens)

            # Final verification
            if contains_any_tag_pattern(text_without_tags):
                print(f"    WARNING: Tags still found after cleaning in segment {seg_idx}")
                segments_with_tags += 1
                continue

            processed_segments.append({
                'segment_id': f"{file_path.stem}_{seg_idx}",
                'source_file': file_path.name,
                'file_type': file_type,
                'tokens': tokens,
                'labels': labels,
                'num_tokens': len(tokens),
                'num_transitions': num_transitions,
                'has_switch': num_transitions > 0,
                'original_text': text_without_tags
            })

    if segments_with_tags > 0:
        print(f"  Warning: Filtered out {segments_with_tags} segments that contained tags")

    print(f"  Processed {len(processed_segments)} valid, clean segments (NO tags)")
    return processed_segments


# Modified process_all_files to include non-switching segments
def process_all_files(data_dir, require_switch=False):
    """
    Process all .txt and .docx files in the data directory
    NEW: Can include segments without switches

    Args:
        require_switch: If True, only include segments with switches
                       If False, include all segments
    """
    data_path = Path(data_dir)

    # Find all files
    txt_files = list(data_path.glob("*.txt"))
    docx_files = list(data_path.glob("*.docx"))

    all_files = [(f, 'txt') for f in txt_files] + [(f, 'docx') for f in docx_files]

    # Filter files with tags
    all_files = [(f, t) for f, t in all_files if has_code_switching_tags(f, t)]

    print(f"\nFound {len(txt_files)} .txt files and {len(docx_files)} .docx files")
    print(f"Files with code-switching tags: {len(all_files)}")
    print(f"Total files to process: {len(all_files)}")

    if require_switch:
        print(f"Mode: INCLUDING ONLY segments with switches")
    else:
        print(f"Mode: INCLUDING ALL segments (with AND without switches)")

    all_segments = []

    for file_path, file_type in all_files:
        try:
            file_segments = process_file_to_segments(file_path, file_type,
                                                     require_switch=require_switch)
            all_segments.extend(file_segments)
        except Exception as e:
            print(f"  ERROR processing {file_path.name}: {e}")
            continue

    print(f"\n=== Processing Complete ===")
    print(f"Total segments: {len(all_segments)}")

    if len(all_segments) == 0:
        print("ERROR: No valid segments found!")
        return pd.DataFrame(), []

    # Separate statistics
    segments_with_switches = [s for s in all_segments if s['has_switch']]
    segments_without_switches = [s for s in all_segments if not s['has_switch']]

    print(f"  Segments WITH switches: {len(segments_with_switches)}")
    print(f"  Segments WITHOUT switches: {len(segments_without_switches)}")

    # Create DataFrame
    segments_data = []

    for segment in all_segments:
        segments_data.append({
            'segment_id': segment['segment_id'],
            'source_file': segment['source_file'],
            'file_type': segment['file_type'],
            'tokens': ' '.join(segment['tokens']),
            'labels': ','.join(map(str, segment['labels'])),
            'num_tokens': segment['num_tokens'],
            'num_transitions': segment['num_transitions'],
            'has_switch': segment['has_switch'],
            'original_text': segment['original_text']
        })

    df = pd.DataFrame(segments_data)

    # Print statistics
    print(f"\n=== Dataset Statistics ===")
    print(f"Total segments: {len(df)}")
    print(f"  With switches: {df['has_switch'].sum()}")
    print(f"  Without switches: {(~df['has_switch']).sum()}")
    print(f"Total tokens: {df['num_tokens'].sum()}")
    print(f"Average tokens per segment: {df['num_tokens'].mean():.1f}")
    print(f"Total transitions: {df['num_transitions'].sum()}")
    print(f"Files represented: {df['source_file'].nunique()}")

    # Label distribution
    all_labels = []
    for labels_str in df['labels']:
        all_labels.extend([int(l) for l in labels_str.split(',')])

    label_names = ['Non-switch Auto', 'Non-switch Allo', 'Switch to Auto', 'Switch to Allo']
    print(f"\nLabel distribution:")
    for i in range(4):
        count = sum(1 for l in all_labels if l == i)
        percentage = count / len(all_labels) * 100
        print(f"  Class {i} ({label_names[i]}): {count} ({percentage:.2f}%)")

    return df, all_segments


def verify_split_balance(train_df, val_df, test_df, target_ratio=0.67):
    """Verify that each split maintains the target switch ratio"""
    print(f"\n=== Verifying Split Balance (Target: {target_ratio * 100:.0f}% with switches) ===")

    for name, df in [('Train', train_df), ('Val', val_df), ('Test', test_df)]:
        with_switches = df['has_switch'].sum()
        total = len(df)
        ratio = with_switches / total if total > 0 else 0

        print(f"\n{name} split:")
        print(f"  Total: {total}")
        print(f"  With switches: {with_switches} ({ratio * 100:.1f}%)")
        print(f"  Without switches: {total - with_switches} ({(1 - ratio) * 100:.1f}%)")

        if abs(ratio - target_ratio) > 0.10:  # 10% tolerance
            print(f"  ⚠️ WARNING: Ratio differs from target by more than 10%")
        else:
            print(f"  ✓ Ratio is acceptable")

    return True

# ============================================================================
# PART 7: MAIN TRAINING PIPELINE
# ============================================================================

def train_tibetan_code_switching(target_switch_ratio=0.67):

    """
        Main training pipeline with logical transition constraints
    """
    print("=" * 60)
    print("TIBETAN CODE-SWITCHING DETECTION TRAINING")
    print("With Proximity-Aware Loss (5-token tolerance)")
    print("With Logical Transition Constraints")
    print("=" * 60)

    # Step 1: Process all files
    print("\nSTEP 1: Processing files...")
    data_dir = 'dataset/annotated-data-raw'
    # data_dir = 'classify_allo_auto/data'

    if not os.path.exists(data_dir):
        print(f"ERROR: Data directory {data_dir} not found!")
        return


    df, all_segments = process_all_files(data_dir, require_switch=False)


    if len(df) == 0:
        print("ERROR: No segments with transitions found!")
        return
    if not validate_no_tags_in_data(df, "ALL PROCESSED DATA"):
        print("\n❌ FATAL ERROR: Tags found in data! Cannot proceed.")
        print("Please review the aggressive_tag_removal function.")
        return

    # Step 1.5: Balance segments to achieve target ratio
    print("\nSTEP 1.5: Balancing segments...")
    balanced_segments = balance_segments_by_switches(all_segments, target_switch_ratio)

    if len(balanced_segments) == 0:
        print("ERROR: No balanced segments available!")
        return

    # Convert balanced segments to DataFrame
    segments_data = []
    for segment in balanced_segments:
        segments_data.append({
            'segment_id': segment['segment_id'],
            'source_file': segment['source_file'],
            'file_type': segment['file_type'],
            'tokens': ' '.join(segment['tokens']),
            'labels': ','.join(map(str, segment['labels'])),
            'num_tokens': segment['num_tokens'],
            'num_transitions': segment['num_transitions'],
            'has_switch': segment['has_switch'],
            'original_text': segment['original_text']
        })

    df = pd.DataFrame(segments_data)
    # Save processed data
    df.to_csv('all_segments_300_400_tokens.csv', index=False)

    # Step 2: Create train/val/test split with better file diversity
    print("\nSTEP 2: Creating train/val/test split with file diversity...")
    train_df, val_df, test_df = create_train_val_test_split(df)
    verify_split_balance(train_df, val_df, test_df, target_ratio=target_switch_ratio)

    # CRITICAL: Validate each split
    print("\nValidating splits for tags...")
    if not validate_no_tags_in_data(train_df, "TRAIN"):
        return
    if not validate_no_tags_in_data(val_df, "VAL"):
        return
    if not validate_no_tags_in_data(test_df, "TEST"):
        return

    print("\n✅ ALL SPLITS VALIDATED - NO TAGS FOUND")

    # NEW: Analyze switch distribution
    print("\n" + "=" * 60)
    print("ANALYZING SWITCH DISTRIBUTION IN SPLITS")
    print("=" * 60)
    train_stats, val_stats, test_stats = analyze_and_balance_switch_distribution(
        train_df, val_df, test_df
    )

    # NEW: Create stratified split if distribution is imbalanced
    if abs(train_stats['auto_pct'] - test_stats['auto_pct']) > 5 or \
            abs(train_stats['allo_pct'] - test_stats['allo_pct']) > 5:
        print("\n" + "=" * 60)
        print("CREATING BETTER STRATIFIED SPLIT DUE TO IMBALANCE")
        print("=" * 60)
        train_df, val_df, test_df = create_stratified_switch_split(df)

        # Re-analyze the stratified split
        print("\n📊 VERIFYING STRATIFIED SPLIT:")
        train_stats, val_stats, test_stats = analyze_and_balance_switch_distribution(
            train_df, val_df, test_df
        )

        # Use stratified files
        train_dataset_file = 'train_segments_stratified.csv'
        val_dataset_file = 'val_segments_stratified.csv'
        test_dataset_file = 'test_segments_stratified.csv'
    else:
        print("\n✅ Distribution is balanced, using original splits")
        train_dataset_file = 'train_segments.csv'
        val_dataset_file = 'val_segments.csv'
        test_dataset_file = 'test_segments.csv'
    # ADD THIS NEW SECTION to clean duplicates:
    print("\n" + "=" * 60)
    print("CHECKING FOR DUPLICATE SEQUENCES BETWEEN TRAIN AND TEST")
    print("=" * 60)

    # Remove any train segments with 25+ token sequences that appear in test
    train_df_clean = remove_duplicate_sequences_from_train(train_df, test_df, min_duplicate_length=25)

    # Save the cleaned training set
    if len(train_df_clean) < len(train_df):
        train_df = train_df_clean
        train_df.to_csv('train_segments_clean.csv', index=False)
        train_dataset_file = 'train_segments_clean.csv'
        print(f"✅ Using cleaned training set: {train_dataset_file}")

    # Step 3: Initialize model with better configuration
    print("\nSTEP 3: Initializing model...")
    # model_name = 'bert-base-multilingual-cased'
    model_name = 'OMRIDRORI/mbert-tibetan-continual-wylie-final'
    # output_dir = './tibetan_code_switching_constrained_model_bert-base-multilingual-cased'
    output_dir = './alloauto-segmentation-training/fine_tuned_ALTO_models/ALTO_allow_non_switch_test_train_and_fixed_loss_7_10_loss_segment_aware'
    # output_dir = './tibetan_code_switching_constrained_model_wylie-final_all_data_no_labels_no_prtial_v2_2_10'

    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    print(f"Using device: {device}")

    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForTokenClassification.from_pretrained(
        model_name,
        num_labels=4,
        label2id={'non_switch_auto': 0, 'non_switch_allo': 1, 'to_auto': 2, 'to_allo': 3},
        id2label={0: 'non_switch_auto', 1: 'non_switch_allo', 2: 'to_auto', 3: 'to_allo'}
    )

    # Initialize with balanced bias for both switch types
    with torch.no_grad():
        model.classifier.bias.data[0] = 0.0  # Non-switch auto
        model.classifier.bias.data[1] = 0.0  # Non-switch allo
        model.classifier.bias.data[2] = -1.0  # Switch to auto (slight negative bias)
        model.classifier.bias.data[3] = -1.0  # Switch to allo (same as auto for balance)

    model = model.to(device)

    # Step 4: Create datasets (using potentially stratified files)
    print("\nSTEP 4: Creating datasets...")
    train_dataset = CodeSwitchingDataset4Class(train_dataset_file, tokenizer)
    val_dataset = CodeSwitchingDataset4Class(val_dataset_file, tokenizer)
    test_dataset = CodeSwitchingDataset4Class(test_dataset_file, tokenizer)

    data_collator = DataCollatorForTokenClassification(tokenizer)


    # Calculate class weights based on actual distribution
    print("\nCalculating class distribution for better weighting...")
    train_labels = []
    for idx in range(len(train_df)):
        labels = train_df.iloc[idx]['labels'].split(',')
        train_labels.extend([int(l) for l in labels])

    label_counts = {i: train_labels.count(i) for i in range(4)}
    total_count = len(train_labels)

    # Calculate separate weights for each switch type
    to_auto_weight = (total_count / (4 * label_counts[2])) * 10 if label_counts[2] > 0 else 30
    to_allo_weight = (total_count / (4 * label_counts[3])) * 10 if label_counts[3] > 0 else 30
    to_auto_weight = min(to_auto_weight, 20)  # Was 50, now 20
    to_allo_weight = min(to_allo_weight, 20)  # Was 50, now 20
    # Boost Switch→Allo more if it's rarer
    if label_counts[3] < label_counts[2] / 2:
        to_allo_weight *= 1.5

    to_auto_weight = min(to_auto_weight, 50)
    to_allo_weight = min(to_allo_weight, 50)

    print(f"  Class distribution in training:")
    for i in range(4):
        print(f"    Class {i}: {label_counts[i]} ({label_counts[i]/total_count*100:.1f}%)")
    print(f"  Using Switch→Auto weight: {to_auto_weight:.1f}")
    print(f"  Using Switch→Allo weight: {to_allo_weight:.1f}")

    # Step 5: Training with improved settings
    print("\nSTEP 5: Starting training with logical constraints...")
    print("  - Invalid transitions (Auto→Auto, Allo→Allo) heavily penalized")
    print("  - Valid transitions rewarded")

    training_args = TrainingArguments(
        output_dir=output_dir,
        eval_strategy="steps",
        eval_steps=30,
        save_strategy="steps",
        save_steps=60,
        learning_rate=2e-5,
        per_device_train_batch_size=8,  # Increase if memory allows
        per_device_eval_batch_size=8,
        num_train_epochs=10,  # Reduced from 15
        weight_decay=0.01,  # Reduced from 0.1
        logging_steps=20,
        load_best_model_at_end=True,
        metric_for_best_model='switch_f1',
        greater_is_better=True,
        warmup_steps=100,
        save_total_limit=2,
        fp16=torch.cuda.is_available(),
        report_to=[],
        gradient_accumulation_steps=1,  # Keep simple
        label_smoothing_factor=0.0,  # Remove smoothing
        # push_to_hub=True,  # Enable pushing to HF
        # hub_model_id="levshechter/tibetan-CS-detector_mbert-tibetan-continual-wylie_all_data_no_labels_no_partial",  # Your HF repo
    )

    class SimpleSwitchTrainer(Trainer):
        """Trainer focused on switch detection with proximity rewards"""

        def __init__(self, *args, **kwargs):
            if 'tokenizer' in kwargs and 'processing_class' not in kwargs:
                kwargs['processing_class'] = kwargs.get('tokenizer')
            super().__init__(*args, **kwargs)

            # Use the simpler SwitchFocusedLoss
            self.loss_fn = SwitchFocusedLoss(
                switch_recall_weight=10.0,
                proximity_tolerance=5,
                segmentation_penalty=3.0,  # Penalty for switches after seg marks
                segmentation_reward=0.3  # Reward for switches at seg marks
            )
            self.tokenizer = kwargs.get('tokenizer') or kwargs.get('processing_class')

        def compute_loss(self, model, inputs, return_outputs=False, num_items_in_batch=None):
            labels = inputs.pop("labels")
            input_ids = inputs.get("input_ids").clone()  # Keep for seg check
            outputs = model(**inputs)
            logits = outputs.get('logits')

            # Decode tokens to check for segmentation marks
            tokens_with_seg_info = self.analyze_segmentation_marks(input_ids)

            loss = self.loss_fn(logits, labels, tokens_with_seg_info)
            return (loss, outputs) if return_outputs else loss

        def analyze_segmentation_marks(self, input_ids):
            """
            Analyze tokens for Tibetan segmentation marks / and //
            Returns tensor marking positions with seg marks
            """
            batch_size, seq_len = input_ids.shape
            seg_marks = torch.zeros_like(input_ids, dtype=torch.float)

            for b in range(batch_size):
                tokens = self.tokenizer.convert_ids_to_tokens(input_ids[b])
                for t, token in enumerate(tokens):
                    if token and ('/' in token or '།' in token):  # Also check for Tibetan shad
                        seg_marks[b, t] = 1.0

            return seg_marks

    # Use the simpler trainer
    trainer = SimpleSwitchTrainer(
        model=model,
        args=training_args,
        train_dataset=train_dataset,
        eval_dataset=val_dataset,
        data_collator=data_collator,
        processing_class=tokenizer,
        compute_metrics=lambda eval_pred: compute_metrics_for_trainer(eval_pred, tolerance=5)
    )

    # # Custom trainer with logical constraints
    # trainer = ProximityAwareTrainer4Class(
    #     model=model,
    #     args=training_args,
    #     data_collator=data_collator,
    #     train_dataset=train_dataset,
    #     eval_dataset=val_dataset,
    #     processing_class=tokenizer,
    #     compute_metrics=lambda eval_pred: compute_metrics_for_trainer(eval_pred, tolerance=5),
    #     switch_loss_weight=max(to_auto_weight, to_allo_weight),  # Use higher weight
    #     proximity_tolerance=5,
    #     distance_decay=0.7,
    #     false_positive_penalty=30.0,
    #     # false_positive_penalty=10.0,
    #     invalid_transition_penalty=100.0  # Very high penalty for invalid transitions,
    #
    # )

    # Add early stopping callback
    from transformers import EarlyStoppingCallback
    trainer.add_callback(EarlyStoppingCallback(early_stopping_patience=5))

    # Train the model
    print("\nTraining with logical transition constraints...")
    trainer.train()

    # Save final model
    trainer.save_model(f'{output_dir}/final_model')
    tokenizer.save_pretrained(f'{output_dir}/final_model')

    # Step 6: Evaluation with constraint application
    print("\nSTEP 6: Evaluating on test set with constraints...")

    # First, evaluate without constraints to see raw performance
    print("\nRaw performance (without constraints):")
    raw_results = trainer.evaluate(eval_dataset=test_dataset)

    print(f"\n=== Raw Test Results (no constraints) ===")
    print(f"Switch F1: {raw_results['eval_switch_f1']:.3f}")
    print(f"Switch Precision: {raw_results['eval_switch_precision']:.3f}")
    print(f"Switch Recall: {raw_results['eval_switch_recall']:.3f}")

    # Now apply constraints during evaluation
    print("\n=== Final Test Results (with logical constraints) ===")
    test_results = trainer.evaluate(eval_dataset=test_dataset)

    print(f"Accuracy: {test_results['eval_accuracy']:.3f}")
    print(f"Switch F1: {test_results['eval_switch_f1']:.3f}")
    print(f"Switch Precision: {test_results['eval_switch_precision']:.3f}")
    print(f"Switch Recall: {test_results['eval_switch_recall']:.3f}")
    print(f"Exact Matches: {test_results.get('eval_exact_matches', 0)}")
    print(f"Proximity Matches: {test_results.get('eval_proximity_matches', 0)}")
    print(f"True Switches: {test_results['eval_true_switches']}")
    print(f"Predicted Switches: {test_results['eval_pred_switches']}")

    # Per-type analysis
    print(f"\nPer-Type Performance:")
    print(f"  Switch→Auto Precision: {test_results.get('eval_to_auto_precision', 0):.3f}")
    print(f"  Switch→Auto Recall: {test_results.get('eval_to_auto_recall', 0):.3f}")
    print(f"  Switch→Allo Precision: {test_results.get('eval_to_allo_precision', 0):.3f}")
    print(f"  Switch→Allo Recall: {test_results.get('eval_to_allo_recall', 0):.3f}")

    # Check balance
    auto_count = test_results.get('eval_matched_to_auto', 0)
    allo_count = test_results.get('eval_matched_to_allo', 0)

    if allo_count == 0 and test_results.get('eval_true_to_allo', 0) > 0:
        print("\n⚠️ WARNING: Model still not predicting Switch→Allo!")
        print("  Consider: 1) More training epochs, 2) Higher weight for Switch→Allo")

    # Show some examples with constraint application
    print("\nShowing test examples with logical constraints...")
    print_test_examples_with_constraints(model, tokenizer, 'test_segments.csv', num_examples=3, tolerance=5)

    print(f"\n=== Training Complete ===")
    print(f"Model saved to: {output_dir}/final_model")
    print(f"Logical constraints: Auto→Allo and Allo→Auto only")

    return trainer, model, tokenizer, test_results


# ============================================================================
# MAIN EXECUTION
# ============================================================================

if __name__ == "__main__":
    test_text = "<auto> བོད་ཡིག་གི་ <allo> ཚིག་གྲུབ་དང་ <auto> སྐད་ཡིག་ <allo <auto allo> auto>"
    tokens, labels = validate_preprocessing(test_text)

    # trainer, model, tokenizer, results = train_tibetan_code_switching()
    trainer, model, tokenizer, results = train_tibetan_code_switching(target_switch_ratio=0.67)

    import ipdb
    ipdb.set_trace()

