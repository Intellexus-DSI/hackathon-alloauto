"""
4-Class Code-Switching Detection System for Tibetan Text
Classes:
  0: Non-switching Auto (continuing in Auto mode)
  1: Non-switching Allo (continuing in Allo mode)
  2: Switch TO Auto
  3: Switch TO Allo
"""

import torch
import torch.nn as nn
import torch.nn.functional as F
import numpy as np
import pandas as pd
import json
import re
import docx
import os
import ssl
from typing import List, Tuple, Dict
from collections import defaultdict
from sklearn.metrics import classification_report
from transformers import (
    AutoTokenizer,
    AutoModelForTokenClassification,
    TrainingArguments,
    DataCollatorForTokenClassification,
    Trainer
)
from torch.utils.data import Dataset

# Environment setup
os.environ["CUDA_VISIBLE_DEVICES"] = "3"
ssl._create_default_https_context = ssl._create_unverified_context

print(f"Number of GPUs available: {torch.cuda.device_count()}")
for i in range(torch.cuda.device_count()):
    print(f"GPU {i}: {torch.cuda.get_device_name(i)}")


# ============================================================================
# PART 1: DATA PROCESSING
# ============================================================================
def robust_preprocess_text(text_content):
    """
    Robust preprocessing that handles various tag formats and text issues.
    """
    # Handle all possible tag variations
    tag_patterns = [
        (r'<\s*ALLO\s*>', '<allo>'),
        (r'<\s*allo\s*>', '<allo>'),
        (r'<\s*Allo\s*>', '<allo>'),
        (r'<\s*AUTO\s*>', '<auto>'),
        (r'<\s*auto\s*>', '<auto>'),
        (r'<\s*Auto\s*>', '<auto>'),
    ]

    for pattern, replacement in tag_patterns:
        text_content = re.sub(pattern, replacement, text_content, flags=re.IGNORECASE)

    # Clean up whitespace while preserving tag boundaries
    # Replace newlines with spaces
    text_content = text_content.replace('\n', ' ')
    text_content = text_content.replace('\r', ' ')
    text_content = text_content.replace('\t', ' ')

    # Replace multiple spaces with single space
    text_content = re.sub(r'\s+', ' ', text_content)

    # Ensure spaces around tags for cleaner splitting
    text_content = re.sub(r'<allo>', ' <allo> ', text_content)
    text_content = re.sub(r'<auto>', ' <auto> ', text_content)

    # Clean up multiple spaces again
    text_content = re.sub(r'\s+', ' ', text_content)

    return text_content.strip()



def process_tibetan_4class(docx_path, tokens_csv_path, sequences_csv_path, sequence_length=512):
    """
    Process Tibetan text with 4-class labeling system.
    """
    doc = docx.Document(docx_path)
    full_text = ' '.join([para.text for para in doc.paragraphs])

    tokens = []
    labels = []
    current_mode = None  # Start with unknown mode

    parts = re.split(r'(<auto>|<allo>)', full_text)

    for i, part in enumerate(parts):
        if part == '<auto>':
            next_mode = 'auto'
            # Find next non-empty part to label correctly
            for j in range(i + 1, len(parts)):
                if parts[j].strip() and parts[j] not in ['<auto>', '<allo>']:
                    break
        elif part == '<allo>':
            next_mode = 'allo'
        elif part.strip():
            words = part.strip().split()

            for word_idx, word in enumerate(words):
                tokens.append(word)

                # Determine label
                if word_idx == 0 and i > 0:  # First word after a tag
                    prev_part = parts[i - 1]
                    if prev_part == '<auto>':
                        if current_mode == 'allo':
                            labels.append(2)  # Switch TO auto
                        else:
                            labels.append(0)  # Continue in auto
                        current_mode = 'auto'
                    elif prev_part == '<allo>':
                        if current_mode == 'auto':
                            labels.append(3)  # Switch TO allo
                        else:
                            labels.append(1)  # Continue in allo
                        current_mode = 'allo'
                    else:
                        # Continue current mode
                        if current_mode == 'auto':
                            labels.append(0)
                        elif current_mode == 'allo':
                            labels.append(1)
                        else:
                            labels.append(0)  # Default to auto if unknown
                else:
                    # Not a switch point
                    if current_mode == 'auto':
                        labels.append(0)
                    elif current_mode == 'allo':
                        labels.append(1)
                    else:
                        labels.append(0)  # Default

    # Create DataFrames
    token_df = pd.DataFrame({
        'token': tokens,
        'label': labels,
        'mode': ['auto' if l in [0, 2] else 'allo' for l in labels]
    })

    sequences = []
    for i in range(0, len(tokens), sequence_length):
        seq_tokens = tokens[i:i + sequence_length]
        seq_labels = labels[i:i + sequence_length]

        if len(seq_tokens) > 0:
            num_switches = sum(1 for l in seq_labels if l in [2, 3])
            contains_switch = 1 if num_switches > 0 else 0

            sequences.append({
                'sequence_id': len(sequences),
                'tokens': ' '.join(seq_tokens),
                'labels': ','.join(map(str, seq_labels)),
                'length': len(seq_tokens),
                'num_switches': num_switches,
                'contains_switch': contains_switch,
                'num_to_auto': sum(1 for l in seq_labels if l == 2),
                'num_to_allo': sum(1 for l in seq_labels if l == 3)
            })

    sequences_df = pd.DataFrame(sequences)

    # Save
    token_df.to_csv(tokens_csv_path, index=False)
    sequences_df.to_csv(sequences_csv_path, index=False)

    # Print statistics
    print(f"\n=== 4-Class Processing Complete ===")
    print(f"Total tokens: {len(tokens)}")
    print(f"Total sequences: {len(sequences)}")
    print(f"\nLabel distribution:")
    label_names = ['Non-switch Auto', 'Non-switch Allo', 'Switch to Auto', 'Switch to Allo']
    for i in range(4):
        count = sum(1 for l in labels if l == i)
        print(f"  Class {i} ({label_names[i]}): {count} ({count / len(labels) * 100:.2f}%)")

    return token_df, sequences_df


# ============================================================================
# PART 2: DATASET CLASS
# ============================================================================

class CodeSwitchingDataset4Class(Dataset):
    """Dataset for 4-class token-level code-switching."""

    def __init__(self, csv_file, tokenizer, max_length=512):
        self.data = pd.read_csv(csv_file)
        self.tokenizer = tokenizer
        self.max_length = max_length

    def __len__(self):
        return len(self.data)

    def __getitem__(self, idx):
        row = self.data.iloc[idx]
        tokens = row['tokens'].split()
        labels = list(map(int, row['labels'].split(',')))

        encoding = self.tokenize_and_align_labels(tokens, labels)

        return {
            'input_ids': torch.tensor(encoding['input_ids']),
            'attention_mask': torch.tensor(encoding['attention_mask']),
            'labels': torch.tensor(encoding['labels'])
        }

    def tokenize_and_align_labels(self, tokens, labels):
        tokenized_inputs = self.tokenizer(
            tokens,
            is_split_into_words=True,
            padding='max_length',
            truncation=True,
            max_length=self.max_length
        )

        word_ids = tokenized_inputs.word_ids()
        aligned_labels = []
        previous_word_idx = None

        for word_idx in word_ids:
            if word_idx is None:
                aligned_labels.append(-100)
            elif word_idx != previous_word_idx:
                aligned_labels.append(labels[word_idx] if word_idx < len(labels) else -100)
            else:
                aligned_labels.append(-100)
            previous_word_idx = word_idx

        tokenized_inputs['labels'] = aligned_labels
        return tokenized_inputs


# ============================================================================
# PART 3: LOSS FUNCTION
# ============================================================================

class ProximityAwareLoss4Class(nn.Module):
    """Loss function for 4-class code-switching with proximity awareness."""

    def __init__(self, max_distance=10, distance_weights=None,
                 switch_loss_weight=10.0, proximity_bonus=True,
                 false_positive_penalty=2, focal_alpha=0.25, focal_gamma=2.0):
        super().__init__()
        self.max_distance = max_distance

        if distance_weights is None:
            self.distance_weights = torch.tensor([1.0, 0.99, 0.95, 0.9, 0.85, 0.80, 0.75, 0.70, 0.65, 0.60, 0.55])
        else:
            self.distance_weights = torch.tensor(distance_weights)

        # Class weights for 4 classes
        self.class_weights = torch.tensor([
            1.0,  # Class 0: Non-switch auto
            1.0,  # Class 1: Non-switch allo
            switch_loss_weight,  # Class 2: Switch to auto
            switch_loss_weight  # Class 3: Switch to allo
        ])

        self.proximity_bonus = proximity_bonus
        self.false_positive_penalty = false_positive_penalty
        self.focal_alpha = focal_alpha
        self.focal_gamma = focal_gamma

    def forward(self, logits, labels):
        batch_size, seq_len, num_classes = logits.shape
        device = logits.device

        self.distance_weights = self.distance_weights.to(device)
        self.class_weights = self.class_weights.to(device)

        # Cross-entropy loss
        ce_loss = F.cross_entropy(
            logits.view(-1, num_classes),
            labels.view(-1),
            weight=self.class_weights,
            reduction='none'
        ).view(batch_size, seq_len)

        valid_mask = (labels != -100).float()

        # Switch classes are 2 and 3
        switch_mask = ((labels == 2) | (labels == 3)).float()
        predictions = torch.argmax(logits, dim=-1)
        pred_switch_mask = ((predictions == 2) | (predictions == 3)).float()

        proximity_bonus = torch.zeros_like(ce_loss)

        if self.proximity_bonus:
            for b in range(batch_size):
                true_switch_pos = torch.where(switch_mask[b] == 1)[0]
                pred_switch_pos = torch.where(pred_switch_mask[b] == 1)[0]

                # Bonus for predictions near true switches
                for pred_pos in pred_switch_pos:
                    if valid_mask[b, pred_pos] == 0:
                        continue

                    if len(true_switch_pos) > 0:
                        distances = torch.abs(true_switch_pos - pred_pos)
                        min_distance = torch.min(distances).item()

                        if min_distance <= self.max_distance:
                            weight = self.distance_weights[min(min_distance, len(self.distance_weights) - 1)]
                            proximity_bonus[b, pred_pos] = ce_loss[b, pred_pos] * weight * 0.5

                # Penalty for missed switches
                for true_pos in true_switch_pos:
                    if valid_mask[b, true_pos] == 0:
                        continue

                    if len(pred_switch_pos) > 0:
                        distances = torch.abs(pred_switch_pos - true_pos)
                        min_distance = torch.min(distances).item()

                        if min_distance > self.max_distance:
                            ce_loss[b, true_pos] *= 2.0

        # Apply bonuses and penalties
        ce_loss = ce_loss - proximity_bonus
        masked_loss = ce_loss * valid_mask

        # False positive penalty
        false_positive_mask = (pred_switch_mask * (1 - switch_mask)) * valid_mask
        false_positive_penalty = false_positive_mask * self.false_positive_penalty

        total_loss = masked_loss + false_positive_penalty
        loss = total_loss.sum() / valid_mask.sum().clamp(min=1)

        return loss


# ============================================================================
# PART 4: METRICS
# ============================================================================

class ProximityAwareMetrics4Class:
    """Metrics for 4-class system with proximity awareness."""

    def __init__(self, max_distance=10):
        self.max_distance = max_distance

    def compute_proximity_metrics(self, predictions: np.ndarray, labels: np.ndarray) -> Dict[str, float]:
        mask = labels != -100
        predictions = predictions[mask]
        labels = labels[mask]

        # Basic accuracy
        correct = (predictions == labels).sum()
        total = len(labels)
        accuracy = correct / total if total > 0 else 0

        results = {
            'accuracy': accuracy,
            'total_true_switches': 0,
            'total_pred_switches': 0,
            'exact_switch_matches': 0,
            'proximity_matches': {i: 0 for i in range(1, self.max_distance + 1)},
            'missed_switches': 0,
            'false_switches': 0
        }

        # Find switches (classes 2 and 3)
        true_switches = []
        pred_switches = []

        for i, (true_label, pred_label) in enumerate(zip(labels, predictions)):
            if true_label in [2, 3]:  # Switch classes
                true_switches.append((i, true_label))
            if pred_label in [2, 3]:
                pred_switches.append((i, pred_label))

        results['total_true_switches'] = len(true_switches)
        results['total_pred_switches'] = len(pred_switches)

        # Match predictions to true switches
        matched_true = set()
        matched_pred = set()

        # Exact matches
        for i, (true_pos, true_class) in enumerate(true_switches):
            for j, (pred_pos, pred_class) in enumerate(pred_switches):
                if true_pos == pred_pos and true_class == pred_class:
                    results['exact_switch_matches'] += 1
                    matched_true.add(i)
                    matched_pred.add(j)

        # Proximity matches
        for i, (true_pos, true_class) in enumerate(true_switches):
            if i in matched_true:
                continue

            best_distance = float('inf')
            best_pred_idx = None

            for j, (pred_pos, pred_class) in enumerate(pred_switches):
                if j in matched_pred:
                    continue

                distance = abs(true_pos - pred_pos)
                if distance <= self.max_distance and distance < best_distance:
                    if true_class == pred_class or distance <= 2:
                        best_distance = distance
                        best_pred_idx = j

            if best_pred_idx is not None and best_distance > 0:
                results['proximity_matches'][int(best_distance)] += 1
                matched_true.add(i)
                matched_pred.add(best_pred_idx)

        results['missed_switches'] = len(true_switches) - len(matched_true)
        results['false_switches'] = len(pred_switches) - len(matched_pred)

        # Calculate F1 scores
        if results['total_true_switches'] > 0:
            proximity_recall_scores = []
            for dist in range(self.max_distance + 1):
                if dist == 0:
                    matches = results['exact_switch_matches']
                else:
                    matches = results['proximity_matches'].get(dist, 0)
                weight = 1.0 - (dist * 0.08)
                proximity_recall_scores.append(matches * weight)
            results['proximity_recall'] = sum(proximity_recall_scores) / results['total_true_switches']
        else:
            results['proximity_recall'] = 0.0

        if results['total_pred_switches'] > 0:
            total_valid_predictions = (results['exact_switch_matches'] +
                                       sum(results['proximity_matches'].values()))
            results['proximity_precision'] = total_valid_predictions / results['total_pred_switches']
        else:
            results['proximity_precision'] = 0.0

        if results['proximity_precision'] + results['proximity_recall'] > 0:
            results['proximity_f1'] = (2 * results['proximity_precision'] * results['proximity_recall'] /
                                       (results['proximity_precision'] + results['proximity_recall']))
        else:
            results['proximity_f1'] = 0.0

        return results


def compute_metrics_with_proximity_4class(eval_pred, max_distance=10):
    """Compute metrics for 4-class system."""
    predictions, labels = eval_pred
    predictions = np.argmax(predictions, axis=2)

    metrics_calc = ProximityAwareMetrics4Class(max_distance=max_distance)

    all_metrics = []
    for pred_seq, label_seq in zip(predictions, labels):
        seq_metrics = metrics_calc.compute_proximity_metrics(pred_seq, label_seq)
        all_metrics.append(seq_metrics)

    # Aggregate
    aggregated = {}
    for key in all_metrics[0].keys():
        if isinstance(all_metrics[0][key], dict):
            aggregated[key] = {}
            for subkey in all_metrics[0][key].keys():
                values = [m[key][subkey] for m in all_metrics if key in m and subkey in m[key]]
                aggregated[key][subkey] = sum(values) / len(values) if values else 0
        else:
            values = [m[key] for m in all_metrics if key in m]
            aggregated[key] = sum(values) / len(values) if values else 0

    flat_metrics = {
        'accuracy': aggregated['accuracy'],
        'proximity_f1': aggregated['proximity_f1'],
        'proximity_recall': aggregated['proximity_recall'],
        'proximity_precision': aggregated['proximity_precision'],
        'exact_matches': aggregated.get('exact_switch_matches', 0),
        'missed_switches': aggregated.get('missed_switches', 0),
        'false_switches': aggregated.get('false_switches', 0)
    }

    for dist in range(1, max_distance + 1):
        if 'proximity_matches' in aggregated and dist in aggregated['proximity_matches']:
            flat_metrics[f'matches_at_{dist}_words'] = aggregated['proximity_matches'][dist]

    return flat_metrics


# ============================================================================
# PART 5: TRAINER
# ============================================================================
class ProximityAwareTrainer4Class(Trainer):
    """Custom trainer for 4-class system."""

    def __init__(self, max_distance=10, distance_weights=None,
                 switch_loss_weight=10.0, false_positive_penalty=2.0,  # Add this parameter
                 *args, **kwargs):
        if 'tokenizer' in kwargs and 'processing_class' not in kwargs:
            kwargs['processing_class'] = kwargs.get('tokenizer')

        super().__init__(*args, **kwargs)
        self.proximity_loss = ProximityAwareLoss4Class(
            max_distance=max_distance,
            distance_weights=distance_weights,
            switch_loss_weight=switch_loss_weight,
            false_positive_penalty=false_positive_penalty  # Pass it to the loss function
        )

    def compute_loss(self, model, inputs, return_outputs=False, num_items_in_batch=None):
        labels = inputs.pop("labels")
        outputs = model(**inputs)
        logits = outputs.get('logits')
        loss = self.proximity_loss(logits, labels)
        return (loss, outputs) if return_outputs else loss
class ProximityAwareTrainer4Class_old(Trainer):
    """Custom trainer for 4-class system."""

    def __init__(self, max_distance=10, distance_weights=None,
                 switch_loss_weight=10.0, *args, **kwargs):
        if 'tokenizer' in kwargs and 'processing_class' not in kwargs:
            kwargs['processing_class'] = kwargs.get('tokenizer')

        super().__init__(*args, **kwargs)
        self.proximity_loss = ProximityAwareLoss4Class(
            max_distance=max_distance,
            distance_weights=distance_weights,
            switch_loss_weight=switch_loss_weight
        )

    def compute_loss(self, model, inputs, return_outputs=False, num_items_in_batch=None):
        labels = inputs.pop("labels")
        outputs = model(**inputs)
        logits = outputs.get('logits')
        loss = self.proximity_loss(logits, labels)
        return (loss, outputs) if return_outputs else loss


# ============================================================================
# PART 6: DATA AUGMENTATION
# ============================================================================

def augment_training_data_4class(train_csv='train_sequences.csv', augmentation_factor=2):
    """Augment training data for 4-class system."""
    train_df = pd.read_csv(train_csv)

    switch_sequences = train_df[train_df['contains_switch'] == 1]
    no_switch_sequences = train_df[train_df['contains_switch'] == 0]

    print(f"\nOriginal training data:")
    print(f"  With switches: {len(switch_sequences)}")
    print(f"  Without switches: {len(no_switch_sequences)}")

    augmented_switches = pd.concat([switch_sequences] * augmentation_factor)
    n_no_switch = min(len(no_switch_sequences), len(augmented_switches) * 2)
    balanced_no_switch = no_switch_sequences.sample(n=n_no_switch, random_state=42)

    augmented_train = pd.concat([augmented_switches, balanced_no_switch])
    augmented_train = augmented_train.sample(frac=1, random_state=42).reset_index(drop=True)

    augmented_train.to_csv('train_sequences_augmented.csv', index=False)

    print(f"\nAugmented training data:")
    print(f"  Total sequences: {len(augmented_train)}")
    print(
        f"  With switches: {augmented_train['contains_switch'].sum()} ({augmented_train['contains_switch'].mean() * 100:.1f}%)")

    return augmented_train


# ============================================================================
# PART 7: DATA SPLITTING
# ============================================================================

def stratified_split_for_sequences_4class(sequences_csv, train_ratio=0.7, val_ratio=0.15, test_ratio=0.15):
    """Create stratified train/val/test split for 4-class system."""
    assert abs(train_ratio + val_ratio + test_ratio - 1.0) < 0.001

    df = pd.read_csv(sequences_csv)

    switch_sequences = df[df['contains_switch'] == 1]
    no_switch_sequences = df[df['contains_switch'] == 0]

    print(f"Total sequences: {len(df)}")
    print(f"  With switches: {len(switch_sequences)}")
    print(f"  Without switches: {len(no_switch_sequences)}")

    def split_group(group_df, train_r, val_r, test_r):
        n = len(group_df)
        n_train = int(n * train_r)
        n_val = int(n * val_r)

        train = group_df.iloc[:n_train]
        val = group_df.iloc[n_train:n_train + n_val]
        test = group_df.iloc[n_train + n_val:]

        return train, val, test

    switch_train, switch_val, switch_test = split_group(switch_sequences, train_ratio, val_ratio, test_ratio)
    no_switch_train, no_switch_val, no_switch_test = split_group(no_switch_sequences, train_ratio, val_ratio,
                                                                 test_ratio)

    train_data = pd.concat([switch_train, no_switch_train]).sample(frac=1, random_state=42).reset_index(drop=True)
    val_data = pd.concat([switch_val, no_switch_val]).sample(frac=1, random_state=42).reset_index(drop=True)
    test_data = pd.concat([switch_test, no_switch_test]).sample(frac=1, random_state=42).reset_index(drop=True)

    train_data.to_csv('train_sequences.csv', index=False)
    val_data.to_csv('val_sequences.csv', index=False)
    test_data.to_csv('test_sequences.csv', index=False)

    print(f"\n=== Stratified Train/Val/Test Split ===")
    print(f"Training: {len(train_data)} sequences")
    print(f"  With switches: {train_data['contains_switch'].sum()} ({train_data['contains_switch'].mean() * 100:.1f}%)")
    print(f"Validation: {len(val_data)} sequences")
    print(f"  With switches: {val_data['contains_switch'].sum()} ({val_data['contains_switch'].mean() * 100:.1f}%)")
    print(f"Test: {len(test_data)} sequences")
    print(f"  With switches: {test_data['contains_switch'].sum()} ({test_data['contains_switch'].mean() * 100:.1f}%)")

    return train_data, val_data, test_data


# ============================================================================
# PART 8: EVALUATION
# ============================================================================

def evaluate_on_test_set_4class(model, tokenizer, test_csv='test_sequences.csv', max_distance=10):
    """Evaluation function for 4-class system."""
    print("\n" + "=" * 60)
    print("EVALUATING ON TEST SET (4-CLASS)")
    print("=" * 60)

    test_df = pd.read_csv(test_csv)
    device = next(model.parameters()).device
    model.eval()

    all_predictions = []
    all_labels = []
    detailed_results = []

    for idx, row in test_df.iterrows():
        # Get the raw tokens
        raw_tokens = row['tokens'].split()

        # Filter out IOB labels that got mixed in
        tokens = []
        for token in raw_tokens:
            # Skip IOB labels
            if token not in ['I-ALLO', 'B-ALLO', 'I-AUTO', 'B-AUTO', 'O']:
                tokens.append(token)

        # Get true labels
        true_labels = list(map(int, row['labels'].split(',')))

        # Make sure labels match tokens length
        true_labels = true_labels[:len(tokens)]

        if len(tokens) == 0:
            continue

        tokenizer_output = tokenizer(
            tokens,
            is_split_into_words=True,
            return_tensors="pt",
            padding=True,
            truncation=True,
            max_length=512
        )

        # Rest of the function continues as before...
        word_ids = tokenizer_output.word_ids()
        inputs = {k: v.to(device) for k, v in tokenizer_output.items()}

        with torch.no_grad():
            outputs = model(**inputs)
            predictions = torch.argmax(outputs.logits, dim=2)
            probs = torch.softmax(outputs.logits, dim=2)

        aligned_predictions = []
        aligned_probs = []
        previous_word_idx = None

        for i, word_idx in enumerate(word_ids):
            if word_idx is not None and word_idx != previous_word_idx:
                pred = predictions[0][i].item()
                aligned_predictions.append(pred)
                aligned_probs.append(probs[0][i].cpu().numpy())
            previous_word_idx = word_idx

        final_len = min(len(aligned_predictions), len(true_labels), len(tokens))
        aligned_predictions = aligned_predictions[:final_len]
        aligned_labels = true_labels[:final_len]

        if final_len > 0:
            all_predictions.extend(aligned_predictions)
            all_labels.extend(aligned_labels)

        # Store results
        true_switches = [(i, l) for i, l in enumerate(aligned_labels) if l in [2, 3]]
        pred_switches = [(i, p) for i, p in enumerate(aligned_predictions) if p in [2, 3]]

        detailed_results.append({
            'sequence_idx': idx,
            'tokens': tokens[:final_len],  # Now contains only actual tokens
            'true_labels': aligned_labels,
            'predictions': aligned_predictions,
            'has_true_switches': len(true_switches) > 0,
            'has_pred_switches': len(pred_switches) > 0,
            'true_switches': true_switches,
            'pred_switches': pred_switches
        })

    all_predictions_array = np.array(all_predictions)
    all_labels_array = np.array(all_labels)

    print(f"Total tokens processed: {len(all_predictions_array)}")

    # Compute metrics
    metrics_calculator = ProximityAwareMetrics4Class(max_distance=max_distance)
    proximity_metrics = metrics_calculator.compute_proximity_metrics(
        all_predictions_array,
        all_labels_array
    )

    # Print results
    print("\n=== Overall Test Metrics ===")
    print(f"Total test tokens: {len(all_labels_array)}")
    print(f"Accuracy: {proximity_metrics['accuracy']:.3f}")
    print(f"\nProximity-aware metrics:")
    print(f"  F1 Score: {proximity_metrics['proximity_f1']:.3f}")
    print(f"  Precision: {proximity_metrics['proximity_precision']:.3f}")
    print(f"  Recall: {proximity_metrics['proximity_recall']:.3f}")

    print(f"\n=== Switch Detection Performance ===")
    print(f"Total true switches: {proximity_metrics['total_true_switches']}")
    print(f"Total predicted switches: {proximity_metrics['total_pred_switches']}")
    print(f"Exact matches: {proximity_metrics['exact_switch_matches']}")

    print("\nDistance breakdown:")
    for dist in range(1, max_distance + 1):
        matches = proximity_metrics['proximity_matches'].get(dist, 0)
        print(f"  {dist} word{'s' if dist > 1 else ''} off: {matches}")

    print(f"\nMissed switches: {proximity_metrics['missed_switches']}")
    print(f"False switches: {proximity_metrics['false_switches']}")

    # Classification report for 4 classes
    if proximity_metrics['total_true_switches'] > 10:
        print("\n=== Per-Class Performance ===")
        print(classification_report(
            all_labels_array,
            all_predictions_array,
            labels=[0, 1, 2, 3],
            target_names=['Non-switch Auto', 'Non-switch Allo', 'Switch to Auto', 'Switch to Allo'],
            digits=3,
            zero_division=0
        ))

    # Save results
    with open('test_results_4class.json', 'w') as f:
        results_summary = {
            'overall_metrics': {k: float(v) if not isinstance(v, dict) else v
                                for k, v in proximity_metrics.items()},
            'total_sequences': len(test_df),
            'sequences_with_true_switches': sum(1 for r in detailed_results if r['has_true_switches']),
            'sequences_with_pred_switches': sum(1 for r in detailed_results if r['has_pred_switches'])
        }
        json.dump(results_summary, f, indent=2)

    print(f"\nDetailed results saved to test_results_4class.json")

    return proximity_metrics, detailed_results


# ============================================================================
# PART 9: MAIN PIPELINE
# ============================================================================


import os
from pathlib import Path
import pandas as pd
from typing import List, Tuple


def process_multiple_docx_files(data_dir: str, output_dir: str = None, sequence_length: int = 512):
    """
    Process all .docx files in a directory and combine them into single dataset.

    Args:
        data_dir: Directory containing .docx files
        output_dir: Output directory for combined CSV files
        sequence_length: Length of sequences for splitting

    Returns:
        Combined token_df and sequences_df
    """
    if output_dir is None:
        output_dir = data_dir

    # Find all .docx files
    docx_files = list(Path(data_dir).glob("*.docx"))
    print("docx_files: ")
    print(docx_files)


    print(f"Found {len(docx_files)} .docx files in {data_dir}")

    all_tokens_dfs = []
    all_sequences_dfs = []

    # Process each file
    for idx, docx_path in enumerate(docx_files):
        print(f"\n[{idx + 1}/{len(docx_files)}] Processing: {docx_path.name}")

        try:
            # Create temporary output paths for this file
            temp_tokens_path = Path(output_dir) / f"temp_tokens_{idx}.csv"
            temp_sequences_path = Path(output_dir) / f"temp_sequences_{idx}.csv"

            # Process the file
            token_df, seq_df = process_tibetan_4class(
                str(docx_path),
                str(temp_tokens_path),
                str(temp_sequences_path),
                sequence_length=sequence_length
            )

            # Add file source information
            token_df['source_file'] = docx_path.name
            token_df['file_id'] = idx

            seq_df['source_file'] = docx_path.name
            seq_df['file_id'] = idx

            all_tokens_dfs.append(token_df)
            all_sequences_dfs.append(seq_df)

            # Clean up temp files
            temp_tokens_path.unlink(missing_ok=True)
            temp_sequences_path.unlink(missing_ok=True)

        except Exception as e:
            print(f"  ERROR processing {docx_path.name}: {e}")
            continue

    if not all_tokens_dfs:
        raise ValueError("No files were successfully processed!")

    # Combine all dataframes
    print("\n=== Combining all processed files ===")
    combined_tokens_df = pd.concat(all_tokens_dfs, ignore_index=True)
    combined_sequences_df = pd.concat(all_sequences_dfs, ignore_index=True)

    # Reassign sequence IDs to be unique across all files
    combined_sequences_df['sequence_id'] = range(len(combined_sequences_df))

    # Save combined data
    combined_tokens_path = Path(output_dir) / 'all_tokens_combined.csv'
    combined_sequences_path = Path(output_dir) / 'all_sequences_combined.csv'

    combined_tokens_df.to_csv(combined_tokens_path, index=False)
    combined_sequences_df.to_csv(combined_sequences_path, index=False)

    # Print combined statistics
    print(f"\n=== Combined Dataset Statistics ===")
    print(f"Total files processed: {len(docx_files)}")
    print(f"Total tokens: {len(combined_tokens_df)}")
    print(f"Total sequences: {len(combined_sequences_df)}")

    # Label distribution across all files
    print(f"\nCombined label distribution:")
    label_names = ['Non-switch Auto', 'Non-switch Allo', 'Switch to Auto', 'Switch to Allo']
    for i in range(4):
        count = (combined_tokens_df['label'] == i).sum()
        percentage = count / len(combined_tokens_df) * 100
        print(f"  Class {i} ({label_names[i]}): {count} ({percentage:.2f}%)")

    # Switch statistics
    switch_sequences = combined_sequences_df[combined_sequences_df['contains_switch'] == 1]
    print(
        f"\nSequences with switches: {len(switch_sequences)} ({len(switch_sequences) / len(combined_sequences_df) * 100:.1f}%)")
    print(f"Total switches to auto: {combined_sequences_df['num_to_auto'].sum()}")
    print(f"Total switches to allo: {combined_sequences_df['num_to_allo'].sum()}")

    # Per-file statistics
    print(f"\n=== Per-file Statistics ===")
    for file_id, file_name in combined_sequences_df[['file_id', 'source_file']].drop_duplicates().values:
        file_seqs = combined_sequences_df[combined_sequences_df['file_id'] == file_id]
        file_switches = file_seqs['contains_switch'].sum()
        print(f"  {file_name}: {len(file_seqs)} sequences, {file_switches} with switches")

    return combined_tokens_df, combined_sequences_df


def stratified_split_combined_data(combined_sequences_path: str,
                                   train_ratio: float = 0.7,
                                   val_ratio: float = 0.15,
                                   test_ratio: float = 0.15,
                                   ensure_file_diversity: bool = True):
    """
    Create stratified train/val/test split from combined data.

    Args:
        combined_sequences_path: Path to combined sequences CSV
        train_ratio: Proportion for training
        val_ratio: Proportion for validation
        test_ratio: Proportion for testing
        ensure_file_diversity: If True, ensure each split has data from multiple files
    """
    df = pd.read_csv(combined_sequences_path)

    if ensure_file_diversity and 'file_id' in df.columns:
        # Group by file and switch status for more balanced splitting
        groups = []
        for file_id in df['file_id'].unique():
            file_df = df[df['file_id'] == file_id]

            # Split this file's data
            switch_seqs = file_df[file_df['contains_switch'] == 1]
            no_switch_seqs = file_df[file_df['contains_switch'] == 0]

            for seq_group in [switch_seqs, no_switch_seqs]:
                if len(seq_group) > 0:
                    n = len(seq_group)
                    n_train = int(n * train_ratio)
                    n_val = int(n * val_ratio)

                    train = seq_group.iloc[:n_train]
                    val = seq_group.iloc[n_train:n_train + n_val]
                    test = seq_group.iloc[n_train + n_val:]

                    groups.append({
                        'train': train,
                        'val': val,
                        'test': test
                    })

        # Combine all groups
        train_data = pd.concat([g['train'] for g in groups if len(g['train']) > 0])
        val_data = pd.concat([g['val'] for g in groups if len(g['val']) > 0])
        test_data = pd.concat([g['test'] for g in groups if len(g['test']) > 0])

    else:
        # Use original stratified split logic
        switch_sequences = df[df['contains_switch'] == 1]
        no_switch_sequences = df[df['contains_switch'] == 0]

        def split_group(group_df, train_r, val_r, test_r):
            n = len(group_df)
            n_train = int(n * train_r)
            n_val = int(n * val_r)

            train = group_df.iloc[:n_train]
            val = group_df.iloc[n_train:n_train + n_val]
            test = group_df.iloc[n_train + n_val:]

            return train, val, test

        switch_train, switch_val, switch_test = split_group(switch_sequences, train_ratio, val_ratio, test_ratio)
        no_switch_train, no_switch_val, no_switch_test = split_group(no_switch_sequences, train_ratio, val_ratio,
                                                                     test_ratio)

        train_data = pd.concat([switch_train, no_switch_train])
        val_data = pd.concat([switch_val, no_switch_val])
        test_data = pd.concat([switch_test, no_switch_test])

    # Shuffle
    train_data = train_data.sample(frac=1, random_state=42).reset_index(drop=True)
    val_data = val_data.sample(frac=1, random_state=42).reset_index(drop=True)
    test_data = test_data.sample(frac=1, random_state=42).reset_index(drop=True)

    # Save splits
    train_data.to_csv('train_sequences_combined.csv', index=False)
    val_data.to_csv('val_sequences_combined.csv', index=False)
    test_data.to_csv('test_sequences_combined.csv', index=False)

    print(f"\n=== Combined Data Train/Val/Test Split ===")
    print(f"Training: {len(train_data)} sequences")
    print(f"  With switches: {train_data['contains_switch'].sum()} ({train_data['contains_switch'].mean() * 100:.1f}%)")
    if 'file_id' in train_data.columns:
        print(f"  From {train_data['file_id'].nunique()} different files")

    print(f"Validation: {len(val_data)} sequences")
    print(f"  With switches: {val_data['contains_switch'].sum()} ({val_data['contains_switch'].mean() * 100:.1f}%)")
    if 'file_id' in val_data.columns:
        print(f"  From {val_data['file_id'].nunique()} different files")

    print(f"Test: {len(test_data)} sequences")
    print(f"  With switches: {test_data['contains_switch'].sum()} ({test_data['contains_switch'].mean() * 100:.1f}%)")
    if 'file_id' in test_data.columns:
        print(f"  From {test_data['file_id'].nunique()} different files")

    return train_data, val_data, test_data

def process_tibetan_4class_from_text(text_content, tokens_csv_path, sequences_csv_path, sequence_length=512):
    """
    Process Tibetan text content with 4-class labeling system.
    This is a refactored version that works with raw text content.
    """
    text_content = robust_preprocess_text(text_content)

    # Now continue with the rest of your processing...
    parts = re.split(r'(<auto>|<allo>)', text_content)
    tokens = []
    labels = []
    current_mode = None  # Start with unknown mode


    for i, part in enumerate(parts):
        if part == '<auto>':
            next_mode = 'auto'
        elif part == '<allo>':
            next_mode = 'allo'
        elif part.strip():
            words = part.strip().split()

            for word_idx, word in enumerate(words):
                tokens.append(word)

                # Determine label
                if word_idx == 0 and i > 0:  # First word after a tag
                    prev_part = parts[i - 1]
                    if prev_part == '<auto>':
                        if current_mode == 'allo':
                            labels.append(2)  # Switch TO auto
                        else:
                            labels.append(0)  # Continue in auto
                        current_mode = 'auto'
                    elif prev_part == '<allo>':
                        if current_mode == 'auto':
                            labels.append(3)  # Switch TO allo
                        else:
                            labels.append(1)  # Continue in allo
                        current_mode = 'allo'
                    else:
                        # Continue current mode
                        if current_mode == 'auto':
                            labels.append(0)
                        elif current_mode == 'allo':
                            labels.append(1)
                        else:
                            labels.append(0)  # Default to auto if unknown
                else:
                    # Not a switch point
                    if current_mode == 'auto':
                        labels.append(0)
                    elif current_mode == 'allo':
                        labels.append(1)
                    else:
                        labels.append(0)  # Default

    # Create DataFrames
    token_df = pd.DataFrame({
        'token': tokens,
        'label': labels,
        'mode': ['auto' if l in [0, 2] else 'allo' for l in labels]
    })

    sequences = []
    for i in range(0, len(tokens), sequence_length):
        seq_tokens = tokens[i:i + sequence_length]
        seq_labels = labels[i:i + sequence_length]

        if len(seq_tokens) > 0:
            num_switches = sum(1 for l in seq_labels if l in [2, 3])
            contains_switch = 1 if num_switches > 0 else 0

            sequences.append({
                'sequence_id': len(sequences),
                'tokens': ' '.join(seq_tokens),
                'labels': ','.join(map(str, seq_labels)),
                'length': len(seq_tokens),
                'num_switches': num_switches,
                'contains_switch': contains_switch,
                'num_to_auto': sum(1 for l in seq_labels if l == 2),
                'num_to_allo': sum(1 for l in seq_labels if l == 3)
            })

    sequences_df = pd.DataFrame(sequences)

    # Save
    token_df.to_csv(tokens_csv_path, index=False)
    sequences_df.to_csv(sequences_csv_path, index=False)

    # Print statistics
    print(f"\n=== 4-Class Processing Complete ===")
    print(f"Total tokens: {len(tokens)}")
    print(f"Total sequences: {len(sequences)}")
    print(f"\nLabel distribution:")
    label_names = ['Non-switch Auto', 'Non-switch Allo', 'Switch to Auto', 'Switch to Allo']
    for i in range(4):
        count = sum(1 for l in labels if l == i)
        print(f"  Class {i} ({label_names[i]}): {count} ({count / len(labels) * 100:.2f}%)")

    return token_df, sequences_df


def process_mixed_files(data_dir: str, output_dir: str = None, sequence_length: int = 512):
    """
    Process all .docx and .txt files in a directory and combine them into single dataset.

    Args:
        data_dir: Directory containing .docx and .txt files
        output_dir: Output directory for combined CSV files
        sequence_length: Length of sequences for splitting

    Returns:
        Combined token_df and sequences_df
    """
    if output_dir is None:
        output_dir = data_dir

    # Find all .docx and .txt files
    docx_files = list(Path(data_dir).glob("*.docx"))
    txt_files = list(Path(data_dir).glob("*.txt"))

    all_files = [(f, 'docx') for f in docx_files] + [(f, 'txt') for f in txt_files]

    print(f"Found {len(docx_files)} .docx files and {len(txt_files)} .txt files in {data_dir}")
    print(f"Total files to process: {len(all_files)}")

    all_tokens_dfs = []
    all_sequences_dfs = []

    # Process each file
    for idx, (file_path, file_type) in enumerate(all_files):
        print(f"\n[{idx + 1}/{len(all_files)}] Processing ({file_type}): {file_path.name}")

        try:
            # Create temporary output paths for this file
            temp_tokens_path = Path(output_dir) / f"temp_tokens_{idx}.csv"
            temp_sequences_path = Path(output_dir) / f"temp_sequences_{idx}.csv"

            # Process based on file type
            if file_type == 'docx':
                # Read docx file
                doc = docx.Document(str(file_path))
                text_content = ' '.join([para.text for para in doc.paragraphs])
            else:  # txt file
                # Read txt file with various encodings
                encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
                text_content = None

                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            text_content = f.read()
                        print(f"  Successfully read with {encoding} encoding")
                        break
                    except UnicodeDecodeError:
                        continue

                if text_content is None:
                    print(f"  ERROR: Could not read {file_path.name} with any encoding")
                    continue

            # Process the text content
            token_df, seq_df = process_tibetan_4class_from_text(
                text_content,
                str(temp_tokens_path),
                str(temp_sequences_path),
                sequence_length=sequence_length
            )

            # Add file source information
            token_df['source_file'] = file_path.name
            token_df['file_id'] = idx
            token_df['file_type'] = file_type

            seq_df['source_file'] = file_path.name
            seq_df['file_id'] = idx
            seq_df['file_type'] = file_type

            all_tokens_dfs.append(token_df)
            all_sequences_dfs.append(seq_df)

            # Clean up temp files
            temp_tokens_path.unlink(missing_ok=True)
            temp_sequences_path.unlink(missing_ok=True)

        except Exception as e:
            print(f"  ERROR processing {file_path.name}: {e}")
            continue

    if not all_tokens_dfs:
        raise ValueError("No files were successfully processed!")

    # Combine all dataframes
    print("\n=== Combining all processed files ===")
    combined_tokens_df = pd.concat(all_tokens_dfs, ignore_index=True)
    combined_sequences_df = pd.concat(all_sequences_dfs, ignore_index=True)

    # Reassign sequence IDs to be unique across all files
    combined_sequences_df['sequence_id'] = range(len(combined_sequences_df))

    # Save combined data
    combined_tokens_path = Path(output_dir) / 'all_tokens_combined.csv'
    combined_sequences_path = Path(output_dir) / 'all_sequences_combined.csv'

    combined_tokens_df.to_csv(combined_tokens_path, index=False)
    combined_sequences_df.to_csv(combined_sequences_path, index=False)

    # Print combined statistics
    print(f"\n=== Combined Dataset Statistics ===")
    print(f"Total files processed: {len(all_files)}")
    print(
        f"  .docx files: {len([f for f in all_sequences_dfs if 'file_type' in f.columns and f['file_type'].iloc[0] == 'docx'])}")
    print(
        f"  .txt files: {len([f for f in all_sequences_dfs if 'file_type' in f.columns and f['file_type'].iloc[0] == 'txt'])}")
    print(f"Total tokens: {len(combined_tokens_df)}")
    print(f"Total sequences: {len(combined_sequences_df)}")

    # Label distribution across all files
    print(f"\nCombined label distribution:")
    label_names = ['Non-switch Auto', 'Non-switch Allo', 'Switch to Auto', 'Switch to Allo']
    for i in range(4):
        count = (combined_tokens_df['label'] == i).sum()
        percentage = count / len(combined_tokens_df) * 100
        print(f"  Class {i} ({label_names[i]}): {count} ({percentage:.2f}%)")

    # Switch statistics
    switch_sequences = combined_sequences_df[combined_sequences_df['contains_switch'] == 1]
    print(
        f"\nSequences with switches: {len(switch_sequences)} ({len(switch_sequences) / len(combined_sequences_df) * 100:.1f}%)")
    print(f"Total switches to auto: {combined_sequences_df['num_to_auto'].sum()}")
    print(f"Total switches to allo: {combined_sequences_df['num_to_allo'].sum()}")

    # Per-file statistics
    print(f"\n=== Per-file Statistics ===")
    for file_id, file_name, file_type in combined_sequences_df[
        ['file_id', 'source_file', 'file_type']].drop_duplicates().values:
        file_seqs = combined_sequences_df[combined_sequences_df['file_id'] == file_id]
        file_switches = file_seqs['contains_switch'].sum()
    return combined_tokens_df, combined_sequences_df



def train_on_combined_data_pipeline():
    """
    Complete pipeline for training on multiple documents.
    """

    data_dir = 'classify_allo_auto/data'  # Directory with your files
    output_dir = 'classify_allo_auto/combined_data'  # Output directory

    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)

    # Process all files (both .docx and .txt)
    combined_tokens, combined_sequences = process_mixed_files(
        data_dir=data_dir,
        output_dir=output_dir,
        sequence_length=512
    )

    # Rest of the pipeline remains the same...
    # Step 2: Create train/val/test split
    print("\n" + "=" * 60)
    print("STEP 2: Creating train/val/test split")
    print("=" * 60)

    train_df, val_df, test_df = stratified_split_combined_data(
        combined_sequences_path=f'{output_dir}/all_sequences_combined.csv',
        train_ratio=0.7,
        val_ratio=0.15,
        test_ratio=0.15,
        ensure_file_diversity=True
    )

    # Step 3: Augment training data
    print("\n" + "=" * 60)
    print("STEP 3: Augmenting training data")
    print("=" * 60)

    train_csv = 'train_sequences_combined.csv'
    train_df = pd.read_csv(train_csv)

    switch_sequences = train_df[train_df['contains_switch'] == 1]
    no_switch_sequences = train_df[train_df['contains_switch'] == 0]

    augmentation_factor = 5
    augmented_switches = pd.concat([switch_sequences] * augmentation_factor)

    n_no_switch = min(len(no_switch_sequences), len(augmented_switches) * 2)
    balanced_no_switch = no_switch_sequences.sample(n=n_no_switch, random_state=42, replace=True)

    augmented_train = pd.concat([augmented_switches, balanced_no_switch])
    augmented_train = augmented_train.sample(frac=1, random_state=42).reset_index(drop=True)
    augmented_train.to_csv('train_sequences_combined_augmented.csv', index=False)

    print(f"Augmented training data:")
    print(f"  Total sequences: {len(augmented_train)}")

    print(f"Augmented training data:")
    print(f"  Total sequences: {len(augmented_train)}")
    print(
        f"  With switches: {augmented_train['contains_switch'].sum()} ({augmented_train['contains_switch'].mean() * 100:.1f}%)")

    # Step 4: Initialize model and train
    print("\n" + "=" * 60)
    print("STEP 4: Training model on combined data")
    print("=" * 60)

    model_name = 'OMRIDRORI/mbert-tibetan-continual-unicode-240k'
    output_model_dir = './classify_allo_auto/combined_model_4class'

    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    print(f"Using device: {device}")

    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForTokenClassification.from_pretrained(
        model_name,
        num_labels=4,
        label2id={
            'non_switch_auto': 0,
            'non_switch_allo': 1,
            'to_auto': 2,
            'to_allo': 3
        },
        id2label={
            0: 'non_switch_auto',
            1: 'non_switch_allo',
            2: 'to_auto',
            3: 'to_allo'
        }
    )

    # Initialize with better bias for rare classes
    with torch.no_grad():
        model.classifier.bias.data[2] = 5.0  # Bias for switch to auto
        model.classifier.bias.data[3] = 5.0  # Bias for switch to allo

    model = model.to(device)

    # Create datasets
    train_dataset = CodeSwitchingDataset4Class('train_sequences_combined_augmented.csv', tokenizer)
    val_dataset = CodeSwitchingDataset4Class('val_sequences_combined.csv', tokenizer)

    data_collator = DataCollatorForTokenClassification(tokenizer)

    # Training arguments with adjustments for combined data
    # Training arguments with adjustments for combined data
    training_args = TrainingArguments(
        output_dir=output_model_dir,
        eval_strategy="steps",
        eval_steps=200,
        save_strategy="steps",
        save_steps=400,  # Changed from 500 to 400 (multiple of 200)
        learning_rate=1e-5,
        per_device_train_batch_size=16,
        per_device_eval_batch_size=16,
        num_train_epochs=10,
        weight_decay=0.01,
        logging_dir=f'{output_model_dir}/logs',
        logging_steps=50,
        load_best_model_at_end=True,
        metric_for_best_model='proximity_f1',
        greater_is_better=True,
        warmup_steps=1000,
        save_total_limit=2,
        gradient_accumulation_steps=2,
        fp16=torch.cuda.is_available(),
    )

    def compute_metrics(eval_pred):
        return compute_metrics_with_proximity_4class(eval_pred, max_distance=10)

    # Use higher switch loss weight for combined data
    trainer = ProximityAwareTrainer4Class(
        model=model,
        args=training_args,
        data_collator=data_collator,
        train_dataset=train_dataset,
        eval_dataset=val_dataset,
        processing_class=tokenizer,
        compute_metrics=compute_metrics,
        max_distance=10,
        distance_weights=[1.0, 0.99, 0.95, 0.9, 0.85, 0.80, 0.75, 0.70, 0.65, 0.60, 0.55],
        switch_loss_weight=30.0,  # Much lower!
        false_positive_penalty=1.0  # Increase this!
    )

    # Train
    trainer.train()

    # Save model
    trainer.save_model(f'{output_model_dir}/final_model')
    tokenizer.save_pretrained(f'{output_model_dir}/final_model')

    model.push_to_hub("levshechter/tibetan-code-switching-detector")
    tokenizer.push_to_hub("levshechter/tibetan-code-switching-detector")

    # Or push the trainer (includes training args)
    trainer.push_to_hub("levshechter/tibetan-code-switching-detector")

    # Step 5: Evaluate on test set
    print("\n" + "=" * 60)
    print("STEP 5: Evaluating on combined test set")
    print("=" * 60)
    # Step 5: Evaluate on test set (existing code)
    print("\n" + "=" * 60)
    print("STEP 5: Evaluating on combined test set")
    print("=" * 60)

    test_metrics, test_results = evaluate_on_test_set_4class(
        model,
        tokenizer,
        'test_sequences_combined.csv',
        max_distance=10
    )

    # Step 5b: NEW - Switch detection evaluation
    print("\n" + "=" * 60)
    print("STEP 5b: Switch Detection Evaluation")
    print("=" * 60)
    from switch_detection_evaluation import evaluate_model_on_test_switch_detection

    switch_metrics, switch_results = evaluate_model_on_test_switch_detection(
        model,
        tokenizer,
        'test_sequences_combined.csv',
        tolerance=5
    )

    print("\n" + "=" * 60)
    print("COMBINED DATA PIPELINE COMPLETE!")
    print("=" * 60)
    print(f"Model saved to: {output_model_dir}/final_model")
    print(f"Token-level F1 Score: {test_metrics['proximity_f1']:.3f}")
    print(f"Switch Detection F1 Score (±5 tokens): {switch_metrics['f1']:.3f}")
    import ipdb
    ipdb.set_trace()
    return trainer, model, tokenizer, test_metrics, switch_metrics

    # test_metrics, test_results = evaluate_on_test_set_4class(
    #     model,
    #     tokenizer,
    #     'test_sequences_combined.csv',
    #     max_distance=10
    # )
    #
    # print("\n" + "=" * 60)
    # print("COMBINED DATA PIPELINE COMPLETE!")
    # print("=" * 60)
    # print(f"Model saved to: {output_model_dir}/final_model")
    # print(f"Test F1 Score: {test_metrics['proximity_f1']:.3f}")
    # import ipdb
    # ipdb.set_trace()

# Run the combined pipeline
if __name__ == "__main__":
    # First, process all documents
    data_dir = 'classify_allo_auto/data'  # Change this to your data directory
    output_dir = 'classify_allo_auto/combined_data'

    os.makedirs(output_dir, exist_ok=True)

    # Process and combine all documents
    combined_tokens, combined_sequences = process_multiple_docx_files(
        data_dir=data_dir,
        output_dir=output_dir,
        sequence_length=512
    )

    # Then run the full training pipeline
    trainer, model, tokenizer, test_metrics = train_on_combined_data_pipeline()
    import ipdb; ipdb.set_trace()
#
# if __name__ == "__main__":
#     train_and_evaluate_4class_pipeline()