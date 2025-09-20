import re
from docx import Document


def convert_auto_allo_to_iob(docx_path, output_path):
    """Convert docx with <auto> and <allo> tags to IOB format"""

    # Read docx file
    doc = Document(docx_path)
    full_text = ' '.join([paragraph.text for paragraph in doc.paragraphs])

    # Process the text
    iob_data = []

    # Split by tags while keeping them
    parts = re.split(r'(<auto>|<allo>)', full_text)

    current_label = None
    first_token_in_entity = True

    for part in parts:
        if part == '<auto>':
            current_label = 'AUTO'
            first_token_in_entity = True
            continue
        elif part == '<allo>':
            current_label = 'ALLO'
            first_token_in_entity = True
            continue
        elif part.strip():  # If part has content
            # Tokenize this part
            tokens = part.strip().split()

            for token in tokens:
                if current_label:
                    if first_token_in_entity:
                        label = f'B-{current_label}'
                        first_token_in_entity = False
                    else:
                        label = f'I-{current_label}'
                else:
                    # If no current label (text before first tag)
                    label = 'O'

                iob_data.append((token, label))

    # Write to output file
    with open(output_path, 'w', encoding='utf-8') as f:
        for token, label in iob_data:
            f.write(f"{token}\t{label}\n")

    return iob_data


def convert_auto_allo_to_json(docx_path, output_path):
    """Convert to JSON format for modern NLP libraries"""
    import json

    # Read docx file
    doc = Document(docx_path)
    full_text = ' '.join([paragraph.text for paragraph in doc.paragraphs])

    # Process into sentences (you might want to adjust sentence splitting)
    sentences = []

    # Split by double slashes as sentence boundaries (adjust as needed)
    text_parts = full_text.split('//')

    for text_part in text_parts:
        if not text_part.strip():
            continue

        tokens = []
        labels = []

        # Process this sentence part
        parts = re.split(r'(<auto>|<allo>)', text_part)
        current_label = None
        first_token = True

        for part in parts:
            if part == '<auto>':
                current_label = 'AUTO'
                first_token = True
            elif part == '<allo>':
                current_label = 'ALLO'
                first_token = True
            elif part.strip():
                part_tokens = part.strip().split()

                for token in part_tokens:
                    tokens.append(token)

                    if current_label:
                        if first_token:
                            labels.append(f'B-{current_label}')
                            first_token = False
                        else:
                            labels.append(f'I-{current_label}')
                    else:
                        labels.append('O')

        if tokens:  # Only add if we have tokens
            sentences.append({
                'tokens': tokens,
                'labels': labels
            })

    # Write JSON output
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(sentences, f, ensure_ascii=False, indent=2)

    return sentences


# More robust version that handles edge cases
def robust_convert_to_iob(docx_path, output_path):
    """Robust conversion handling various edge cases"""

    # Read docx
    doc = Document(docx_path)
    full_text = ' '.join([paragraph.text for paragraph in doc.paragraphs])

    # Clean up any extra spaces around tags
    full_text = re.sub(r'\s*<auto>\s*', '<auto>', full_text)
    full_text = re.sub(r'\s*<allo>\s*', '<allo>', full_text)

    iob_data = []

    # Use regex to find all tags and text between them
    pattern = r'(<auto>|<allo>)([^<]*)'
    matches = re.findall(pattern, full_text)

    # Handle any text before the first tag
    first_tag_pos = full_text.find('<')
    if first_tag_pos > 0:
        pre_text = full_text[:first_tag_pos].strip()
        if pre_text:
            tokens = pre_text.split()
            for token in tokens:
                iob_data.append((token, 'O'))

    # Process matched segments
    for tag, text in matches:
        if not text.strip():
            continue

        entity_type = tag.replace('<', '').replace('>', '').upper()
        tokens = text.strip().split()

        for i, token in enumerate(tokens):
            if i == 0:
                label = f'B-{entity_type}'
            else:
                label = f'I-{entity_type}'
            iob_data.append((token, label))

    # Write output in different formats

    # 1. Simple TSV format
    with open(output_path, 'w', encoding='utf-8') as f:
        for token, label in iob_data:
            f.write(f"{token}\t{label}\n")

    # 2. CoNLL format with sentence boundaries (based on //)
    conll_path = output_path.replace('.txt', '_conll.txt')
    with open(conll_path, 'w', encoding='utf-8') as f:
        for token, label in iob_data:
            f.write(f"{token}\t{label}\n")
            if token.endswith('//'):
                f.write("\n")  # Empty line for sentence boundary

    # 3. Training-ready format for spaCy
    spacy_path = output_path.replace('.txt', '_spacy.json')
    import json

    # Convert to spaCy format
    training_data = []
    current_sentence = []
    current_labels = []

    for token, label in iob_data:
        current_sentence.append(token)
        current_labels.append(label)

        # If sentence ends (you might adjust this logic)
        if token.endswith('//') or token.endswith('/'):
            if current_sentence:
                training_data.append({
                    'tokens': current_sentence,
                    'labels': current_labels
                })
                current_sentence = []
                current_labels = []

    # Don't forget the last sentence
    if current_sentence:
        training_data.append({
            'tokens': current_sentence,
            'labels': current_labels
        })

    with open(spacy_path, 'w', encoding='utf-8') as f:
        json.dump(training_data, f, ensure_ascii=False, indent=2)

    return iob_data


# Usage
robust_convert_to_iob('classify_allo_auto/data/Nicola_Bajetta_rNam_gsum_bshad_pa_Auto_vs_Allo_signals_alo_and_auto.docx', 'classify_allo_auto/data/output_iob_nicola_rNam_gsum_alo_auto.txt')