import re
import csv
from docx import Document


def convert_auto_allo_to_csv(docx_path, output_csv_path):
    """Convert docx with <auto> and <allo> tags to CSV format"""

    # Read docx file
    doc = Document(docx_path)
    full_text = ' '.join([paragraph.text for paragraph in doc.paragraphs])

    # Process the text
    iob_data = []

    # Clean up any extra spaces around tags
    full_text = re.sub(r'\s*<auto>\s*', '<auto>', full_text)
    full_text = re.sub(r'\s*<allo>\s*', '<allo>', full_text)

    # Split by tags while keeping them
    parts = re.split(r'(<auto>|<allo>)', full_text)

    current_label = None
    first_token_in_entity = True

    for part in parts:
        if part == '<auto>':
            current_label = 'AUTO'
            first_token_in_entity = True
            continue
        elif part == '<allo>':
            current_label = 'ALLO'
            first_token_in_entity = True
            continue
        elif part.strip():  # If part has content
            # Tokenize this part
            tokens = part.strip().split()

            for token in tokens:
                if current_label:
                    if first_token_in_entity:
                        label = f'B-{current_label}'
                        first_token_in_entity = False
                    else:
                        label = f'I-{current_label}'
                else:
                    # If no current label (text before first tag)
                    label = 'O'

                iob_data.append((token, label))

    # Write to CSV file
    with open(output_csv_path, 'w', encoding='utf-8', newline='') as f:
        writer = csv.writer(f)

        # Write header
        writer.writerow(['token', 'label'])

        # Write data
        for token, label in iob_data:
            writer.writerow([token, label])

    print(f"Saved CSV to: {output_csv_path}")
    print(f"Total tokens processed: {len(iob_data)}")

    # Print statistics
    stats = {'AUTO': 0, 'ALLO': 0, 'O': 0}
    for token, label in iob_data:
        if label.startswith('B-'):
            entity_type = label[2:]
            stats[entity_type] = stats.get(entity_type, 0) + 1
        elif label == 'O':
            stats['O'] += 1

    print(f"\nStatistics:")
    print(f"  AUTO entities: {stats.get('AUTO', 0)}")
    print(f"  ALLO entities: {stats.get('ALLO', 0)}")
    print(f"  Outside tokens: {stats['O']}")

    return iob_data


def convert_docx_to_enhanced_csv(docx_path, output_csv_path):
    """Convert DOCX to CSV with additional features"""

    # Read docx file
    doc = Document(docx_path)
    full_text = ' '.join([paragraph.text for paragraph in doc.paragraphs])

    # Process the text
    iob_data = []

    # Clean up spaces
    full_text = re.sub(r'\s*<auto>\s*', '<auto>', full_text)
    full_text = re.sub(r'\s*<allo>\s*', '<allo>', full_text)

    # Use regex to find all tags and text between them
    pattern = r'(<auto>|<allo>)([^<]*)'
    matches = re.findall(pattern, full_text)

    # Handle any text before the first tag
    first_tag_pos = full_text.find('<')
    if first_tag_pos > 0:
        pre_text = full_text[:first_tag_pos].strip()
        if pre_text:
            tokens = pre_text.split()
            for token in tokens:
                iob_data.append((token, 'O'))

    # Process matched segments
    for tag, text in matches:
        if not text.strip():
            continue

        entity_type = tag.replace('<', '').replace('>', '').upper()
        tokens = text.strip().split()

        for i, token in enumerate(tokens):
            if i == 0:
                label = f'B-{entity_type}'
            else:
                label = f'I-{entity_type}'
            iob_data.append((token, label))

    # Write enhanced CSV with sentence grouping
    with open(output_csv_path, 'w', encoding='utf-8', newline='') as f:
        writer = csv.writer(f)

        # Write header with more columns
        writer.writerow(['sentence_id', 'token_position', 'token', 'label', 'entity_type'])

        sentence_id = 0
        token_position = 0

        for token, label in iob_data:
            # Extract entity type (AUTO, ALLO, or O)
            if label == 'O':
                entity_type = 'O'
            else:
                entity_type = label[2:]  # Remove B- or I- prefix

            writer.writerow([sentence_id, token_position, token, label, entity_type])

            token_position += 1

            # Check for sentence boundary
            if '//' in token or token.endswith('.'):
                sentence_id += 1
                token_position = 0

    print(f"Saved enhanced CSV to: {output_csv_path}")


# Main execution
if __name__ == "__main__":
    # Your file paths
    input_docx = 'classify_allo_auto/data/Nicola_Bajetta_rNam_gsum_bshad_pa_Auto_vs_Allo_signals_alo_and_auto_cleaned.docx'

    # Simple CSV output
    output_csv = 'classify_allo_auto/data/output_iob_nicola_rNam_gsum_alo_auto.csv'
    convert_auto_allo_to_csv(input_docx, output_csv)

    # Enhanced CSV output with more features
    enhanced_csv = 'classify_allo_auto/data/output_iob_nicola_rNam_gsum_alo_auto_enhanced.csv'
    convert_docx_to_enhanced_csv(input_docx, enhanced_csv)

    # You can also create multiple output formats
    print("\nCreated CSV files:")
    print(f"1. Simple format: {output_csv}")
    print(f"2. Enhanced format: {enhanced_csv}")