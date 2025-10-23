"""
Exact Tibetan Code-Switching Data Preprocessing
Extracted from original training code - EXACT SAME logic
"""

import re
import os
import json
import docx
import hashlib
import pandas as pd
import random
from pathlib import Path
from typing import List, Tuple, Dict
from collections import defaultdict


# ============================================================================
# TAG CLEANING AND VALIDATION
# ============================================================================


def verify_split_balance(train_df, val_df, test_df, target_ratio=0.67):
    """Verify that each split maintains the target switch ratio"""
    print(f"\n=== Verifying Split Balance (Target: {target_ratio * 100:.0f}% with switches) ===")

    for name, df in [('Train', train_df), ('Val', val_df), ('Test', test_df)]:
        with_switches = df['has_switch'].sum()
        total = len(df)
        ratio = with_switches / total if total > 0 else 0

        print(f"\n{name} split:")
        print(f"  Total: {total}")
        print(f"  With switches: {with_switches} ({ratio * 100:.1f}%)")
        print(f"  Without switches: {total - with_switches} ({(1 - ratio) * 100:.1f}%)")

        if abs(ratio - target_ratio) > 0.10:  # 10% tolerance
            print(f"  ⚠️ WARNING: Ratio differs from target by more than 10%")
        else:
            print(f"  ✓ Ratio is acceptable")

    return True
def validate_no_tags_in_data(df, split_name):
    """
    CRITICAL: Final validation to ensure absolutely NO tags in any data
    Returns True if clean, False if tags found
    """
    print(f"\n{'=' * 60}")
    print(f"CRITICAL TAG VALIDATION FOR {split_name}")
    print(f"{'=' * 60}")

    issues_found = []

    for idx in range(len(df)):
        row = df.iloc[idx]

        # Check tokens
        tokens_str = row['tokens']
        if contains_any_tag_pattern(tokens_str):
            issues_found.append(f"Row {idx}: Tags found in tokens: {tokens_str[:100]}")

        # Check original text
        if 'original_text' in row:
            orig_text = row['original_text']
            if contains_any_tag_pattern(orig_text):
                issues_found.append(f"Row {idx}: Tags found in original_text: {orig_text[:100]}")

        # Check individual tokens
        tokens_list = tokens_str.split()
        for i, token in enumerate(tokens_list):
            if contains_any_tag_pattern(token):
                issues_found.append(f"Row {idx}, Token {i}: '{token}' contains tag pattern")

    if issues_found:
        print(f"❌ VALIDATION FAILED! Found {len(issues_found)} issues:")
        for issue in issues_found[:10]:  # Show first 10
            print(f"  - {issue}")
        if len(issues_found) > 10:
            print(f"  ... and {len(issues_found) - 10} more issues")
        return False
    else:
        print(f"✅ VALIDATION PASSED! No tags found in {len(df)} segments")
        return True

def aggressive_tag_removal(text):
    """
    Aggressively remove ALL forms of tags and partial tags from text
    """
    tag_patterns = [
        # Full tags with any case and spacing
        r'<\s*[Aa][Uu][Tt][Oo]\s*>',
        r'<\s*[Aa][Ll][Ll][Oo]\s*>',
        # Partial tags (opening)
        r'<\s*[Aa][Uu][Tt][Oo]\b',
        r'<\s*[Aa][Ll][Ll][Oo]\b',
        r'<\s*[Aa][Uu][Tt]\b',
        r'<\s*[Aa][Ll][Ll]\b',
        r'<\s*[Aa][Uu]\b',
        r'<\s*[Aa][Ll]\b',
        # Partial tags (closing)
        r'\b[Aa][Uu][Tt][Oo]\s*>',
        r'\b[Aa][Ll][Ll][Oo]\s*>',
        r'\b[Uu][Tt][Oo]\s*>',
        r'\b[Ll][Ll][Oo]\s*>',
        # Any word containing tag patterns
        r'\S*<[Aa][Uu][Tt][Oo]\S*',
        r'\S*<[Aa][Ll][Ll][Oo]\S*',
        r'\S*[Aa][Uu][Tt][Oo]>\S*',
        r'\S*[Aa][Ll][Ll][Oo]>\S*',
    ]

    cleaned_text = text
    for pattern in tag_patterns:
        cleaned_text = re.sub(pattern, ' ', cleaned_text)

    cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
    return cleaned_text.strip()


def contains_any_tag_pattern(text):
    """
    Check if text contains ANY form of tag or partial tag
    Returns True if ANY tag pattern is found
    """
    # Comprehensive list of patterns to check
    forbidden_patterns = [
        # Full tags
        r'<auto>', r'<AUTO>', r'<allo>', r'<ALLO>',
        r'<Auto>', r'<Allo>',
        # With spaces
        r'<\s+auto\s+>', r'<\s+AUTO\s+>', r'<\s+allo\s+>', r'<\s+ALLO\s+>',
        # Partial tags (missing closing)
        r'<auto\b', r'<AUTO\b', r'<allo\b', r'<ALLO\b',
        r'<aut\b', r'<AUT\b', r'<all\b', r'<ALL\b',
        r'<au\b', r'<AU\b', r'<al\b', r'<AL\b',
        # Partial tags (missing opening)
        r'\bauto>', r'\bAUTO>', r'\ballo>', r'\bALLO>',
        r'\buto>', r'\bUTO>', r'\bllo>', r'\bLLO>',
        r'\bto>', r'\bTO>', r'\blo>', r'\bLO>',
        # Check for fragments within words
        r'<auto', r'<AUTO', r'<allo', r'<ALLO',
        r'auto>', r'AUTO>', r'allo>', r'ALLO>',
    ]

    text_lower = text.lower()

    for pattern in forbidden_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return True

    # Also check for simple substring presence
    forbidden_substrings = [
        '<auto', '<AUTO', '<allo', '<ALLO',
        'auto>', 'AUTO>', 'allo>', 'ALLO>',
        '<Auto', '<Allo', 'Auto>', 'Allo>'
    ]

    for substring in forbidden_substrings:
        if substring in text or substring.lower() in text_lower:
            return True

    return False


# ============================================================================
# SEGMENT PROCESSING
# ============================================================================

def process_segment_to_tokens_labels_clean(segment_text):
    """
    Convert segment to tokens and labels, ensuring NO tags in tokens
    """
    # Normalize tags for consistent splitting
    normalized_text = re.sub(r'<\s*AUTO\s*>', '<auto>', segment_text, flags=re.IGNORECASE)
    normalized_text = re.sub(r'<\s*ALLO\s*>', '<allo>', normalized_text, flags=re.IGNORECASE)

    # Split by tags
    parts = re.split(r'(<auto>|<allo>)', normalized_text.lower())

    tokens = []
    labels = []
    current_mode = None

    for i, part in enumerate(parts):
        if part == '<auto>':
            continue  # Skip tag
        elif part == '<allo>':
            continue  # Skip tag
        elif part.strip():
            # Clean any remaining tag fragments
            clean_part = aggressive_tag_removal(part)
            words = clean_part.strip().split()

            # Determine mode from previous tags
            segment_mode = None
            for j in range(i - 1, -1, -1):
                if parts[j] == '<auto>':
                    segment_mode = 'auto'
                    break
                elif parts[j] == '<allo>':
                    segment_mode = 'allo'
                    break

            if segment_mode is None:
                segment_mode = 'auto'  # Default to auto if no tag

            for word_idx, word in enumerate(words):
                # Double-check word doesn't contain tags
                if contains_any_tag_pattern(word):
                    continue  # Skip tag fragments

                tokens.append(word)

                # Assign label
                if word_idx == 0 and current_mode is not None and current_mode != segment_mode:
                    label = 2 if segment_mode == 'auto' else 3
                else:
                    label = 0 if segment_mode == 'auto' else 1

                labels.append(label)

            current_mode = segment_mode

    return tokens, labels


def verify_segment_has_switch(segment_text):
    """
    Verify that a segment contains at least one actual switch point
    Returns True if segment has at least one transition between auto and allo
    """
    # Must have both tags
    if '<auto>' not in segment_text.lower() or '<allo>' not in segment_text.lower():
        return False

    # Check for actual switches by looking at tag sequences
    parts = re.split(r'(<auto>|<allo>)', segment_text.lower())

    last_tag = None
    switches_found = 0

    for part in parts:
        if part == '<auto>':
            if last_tag == '<allo>':
                switches_found += 1
            last_tag = '<auto>'
        elif part == '<allo>':
            if last_tag == '<auto>':
                switches_found += 1
            last_tag = '<allo>'

    return switches_found > 0


# ============================================================================
# TEXT PROCESSING
# ============================================================================

def clean_and_normalize_text(text):
    """Clean and normalize text while preserving tags and Tibetan content"""
    # Remove excessive whitespace but preserve tags
    text = re.sub(r'\s+', ' ', text)

    # Ensure proper spacing around tags
    text = re.sub(r'(<auto>)', r' \1 ', text, flags=re.IGNORECASE)
    text = re.sub(r'(<allo>)', r' \1 ', text, flags=re.IGNORECASE)

    # Clean up multiple spaces
    text = re.sub(r'\s+', ' ', text)

    return text.strip()


def split_into_sentences(text_content):
    """
    Split text into sentences based on Tibetan punctuation and markers
    """
    sentences = []
    current_sentence = ""

    # First split by sentence endings, keeping the delimiters
    parts = re.split(r'(\.|//?)', text_content)

    for i, part in enumerate(parts):
        if part.strip() in ['.', '/', '//']:
            # End of sentence
            if current_sentence.strip():
                current_sentence += part
                sentences.append(current_sentence.strip())
                current_sentence = ""
        else:
            current_sentence += part

    # Add any remaining content
    if current_sentence.strip():
        sentences.append(current_sentence.strip())

    return [s for s in sentences if s.strip()]


# ============================================================================
# SEGMENT EXTRACTION
# ============================================================================

def extract_segments_with_token_limit(sentences, min_tokens=300, max_tokens=400, require_switch=False):
    """
    Extract segments of 300-400 tokens
    Can optionally include segments WITHOUT switches

    Args:
        require_switch: If True, only include segments with switches
                       If False, include all segments (with or without switches)
    """
    segments = []
    current_segment = []
    current_token_count = 0

    for sentence in sentences:
        sentence_tokens = sentence.split()
        sentence_token_count = len(sentence_tokens)

        if current_token_count + sentence_token_count > max_tokens:
            if current_token_count >= min_tokens:
                segment_text = ' '.join(current_segment)

                has_switch = verify_segment_has_switch(segment_text)

                # Include segment based on require_switch parameter
                if not require_switch or has_switch:
                    segments.append({
                        'text': segment_text,
                        'token_count': current_token_count,
                        'has_transition': has_switch
                    })

            current_segment = [sentence]
            current_token_count = sentence_token_count
        else:
            current_segment.append(sentence)
            current_token_count += sentence_token_count

            if current_token_count >= min_tokens:
                segment_text = ' '.join(current_segment)
                has_switch = verify_segment_has_switch(segment_text)

                # Include segment based on require_switch parameter
                if not require_switch or has_switch:
                    segments.append({
                        'text': segment_text,
                        'token_count': current_token_count,
                        'has_transition': has_switch
                    })
                    current_segment = []
                    current_token_count = 0

    # Handle remaining sentences
    if current_segment and current_token_count >= min_tokens // 2:
        segment_text = ' '.join(current_segment)
        has_switch = verify_segment_has_switch(segment_text)

        if not require_switch or has_switch:
            segments.append({
                'text': segment_text,
                'token_count': current_token_count,
                'has_transition': has_switch
            })

    return segments


# ============================================================================
# FILE PROCESSING
# ============================================================================

def has_code_switching_tags(file_path, file_type):
    """Check if file contains AUTO/ALLO tags before processing"""
    try:
        if file_type == 'docx':
            doc = docx.Document(str(file_path))
            text_content = ' '.join([para.text for para in doc.paragraphs])
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()

        # Check for any variation of tags
        text_lower = text_content.lower()
        return ('<auto>' in text_lower or '<allo>' in text_lower)
    except:
        return False


def process_file_to_segments(file_path, file_type, require_switch=False):
    """
    Process a single file and extract segments
    Can include segments without switches
    CRITICAL: Ensures NO tags in tokens

    Args:
        require_switch: If True, only include segments with switches
                       If False, include all segments
    """
    print(f"Processing {file_type} file: {file_path.name}")

    # Read file content
    if file_type == 'docx':
        doc = docx.Document(str(file_path))
        text_content = ' '.join([para.text for para in doc.paragraphs])
    else:  # txt file
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        text_content = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text_content = f.read()
                break
            except UnicodeDecodeError:
                continue

        if text_content is None:
            raise ValueError(f"Could not read {file_path.name} with any encoding")

    # Clean and normalize
    text_content = clean_and_normalize_text(text_content)

    # Split into sentences
    sentences = split_into_sentences(text_content)
    print(f"  Found {len(sentences)} sentences")

    # Extract segments (with or without switches based on parameter)
    segments = extract_segments_with_token_limit(sentences, min_tokens=300, max_tokens=400,
                                                 require_switch=require_switch)

    if require_switch:
        print(f"  Found {len(segments)} segments WITH switches (300-400 tokens each)")
    else:
        segments_with_switches = sum(1 for s in segments if s['has_transition'])
        segments_without_switches = len(segments) - segments_with_switches
        print(f"  Found {len(segments)} total segments:")
        print(f"    - {segments_with_switches} with switches")
        print(f"    - {segments_without_switches} without switches")

    # Convert segments to token-label pairs
    processed_segments = []
    segments_with_tags = 0

    for seg_idx, segment in enumerate(segments):
        tokens, labels = process_segment_to_tokens_labels_clean(segment['text'])

        if len(tokens) > 0:
            num_transitions = sum(1 for l in labels if l in [2, 3])

            # CRITICAL: Verify no tags in tokens
            has_tags = False
            for token in tokens:
                if contains_any_tag_pattern(token):
                    has_tags = True
                    segments_with_tags += 1
                    print(f"    WARNING: Found tag in token: '{token}'")
                    break

            if has_tags:
                continue  # Skip this segment entirely

            # Create text WITHOUT tags for storage
            text_without_tags = ' '.join(tokens)

            # Final verification
            if contains_any_tag_pattern(text_without_tags):
                print(f"    WARNING: Tags still found after cleaning in segment {seg_idx}")
                segments_with_tags += 1
                continue

            processed_segments.append({
                'segment_id': f"{file_path.stem}_{seg_idx}",
                'source_file': file_path.name,
                'file_type': file_type,
                'tokens': tokens,
                'labels': labels,
                'num_tokens': len(tokens),
                'num_transitions': num_transitions,
                'has_switch': num_transitions > 0,
                'original_text': text_without_tags
            })

    if segments_with_tags > 0:
        print(f"  Warning: Filtered out {segments_with_tags} segments that contained tags")

    print(f"  Processed {len(processed_segments)} valid, clean segments (NO tags)")
    return processed_segments


def process_all_files(data_dir, require_switch=False):
    """
    Process all .txt and .docx files in the data directory
    Can include segments without switches

    Args:
        require_switch: If True, only include segments with switches
                       If False, include all segments
    """
    data_path = Path(data_dir)

    # Find all files
    txt_files = list(data_path.glob("*.txt"))
    docx_files = list(data_path.glob("*.docx"))

    all_files = [(f, 'txt') for f in txt_files] + [(f, 'docx') for f in docx_files]

    # Filter files with tags
    all_files = [(f, t) for f, t in all_files if has_code_switching_tags(f, t)]

    print(f"\nFound {len(txt_files)} .txt files and {len(docx_files)} .docx files")
    print(f"Files with code-switching tags: {len(all_files)}")
    print(f"Total files to process: {len(all_files)}")

    if require_switch:
        print(f"Mode: INCLUDING ONLY segments with switches")
    else:
        print(f"Mode: INCLUDING ALL segments (with AND without switches)")

    all_segments = []

    for file_path, file_type in all_files:
        try:
            file_segments = process_file_to_segments(file_path, file_type,
                                                     require_switch=require_switch)
            all_segments.extend(file_segments)
        except Exception as e:
            print(f"  ERROR processing {file_path.name}: {e}")
            continue

    print(f"\n=== Processing Complete ===")
    print(f"Total segments: {len(all_segments)}")

    if len(all_segments) == 0:
        print("ERROR: No valid segments found!")
        return pd.DataFrame(), []

    # Separate statistics
    segments_with_switches = [s for s in all_segments if s['has_switch']]
    segments_without_switches = [s for s in all_segments if not s['has_switch']]

    print(f"  Segments WITH switches: {len(segments_with_switches)}")
    print(f"  Segments WITHOUT switches: {len(segments_without_switches)}")

    # Create DataFrame
    segments_data = []

    for segment in all_segments:
        segments_data.append({
            'segment_id': segment['segment_id'],
            'source_file': segment['source_file'],
            'file_type': segment['file_type'],
            'tokens': ' '.join(segment['tokens']),
            'labels': ','.join(map(str, segment['labels'])),
            'num_tokens': segment['num_tokens'],
            'num_transitions': segment['num_transitions'],
            'has_switch': segment['has_switch'],
            'original_text': segment['original_text']
        })

    df = pd.DataFrame(segments_data)

    # Print statistics
    print(f"\n=== Dataset Statistics ===")
    print(f"Total segments: {len(df)}")
    print(f"  With switches: {df['has_switch'].sum()}")
    print(f"  Without switches: {(~df['has_switch']).sum()}")
    print(f"Total tokens: {df['num_tokens'].sum()}")
    print(f"Average tokens per segment: {df['num_tokens'].mean():.1f}")
    print(f"Total transitions: {df['num_transitions'].sum()}")
    print(f"Files represented: {df['source_file'].nunique()}")

    # Label distribution
    all_labels = []
    for labels_str in df['labels']:
        all_labels.extend([int(l) for l in labels_str.split(',')])

    label_names = ['Non-switch Auto', 'Non-switch Allo', 'Switch to Auto', 'Switch to Allo']
    print(f"\nLabel distribution:")
    for i in range(4):
        count = sum(1 for l in all_labels if l == i)
        percentage = count / len(all_labels) * 100
        print(f"  Class {i} ({label_names[i]}): {count} ({percentage:.2f}%)")

    return df, all_segments


# ============================================================================
# BALANCING AND SPLITTING
# ============================================================================

def balance_segments_by_switches(all_segments, target_switch_ratio=0.67):
    """
    Balance segments to have target_switch_ratio with switches

    Args:
        all_segments: List of all processed segments
        target_switch_ratio: Proportion of segments that should have switches (default 0.67 = 2/3)

    Returns:
        Balanced list of segments
    """
    print(f"\n=== Balancing Segments (Target: {target_switch_ratio * 100:.0f}% with switches) ===")

    # Separate segments by whether they have switches
    segments_with_switches = [s for s in all_segments if s['has_switch']]
    segments_without_switches = [s for s in all_segments if not s['has_switch']]

    n_with_switches = len(segments_with_switches)
    n_without_switches = len(segments_without_switches)

    print(f"Available segments:")
    print(f"  With switches: {n_with_switches}")
    print(f"  Without switches: {n_without_switches}")

    if n_with_switches == 0:
        print("ERROR: No segments with switches found!")
        return []

    # Calculate how many without-switch segments we need
    desired_without_switches = int(n_with_switches * (1 - target_switch_ratio) / target_switch_ratio)

    # Use all segments with switches
    balanced_segments = segments_with_switches.copy()

    # Sample the desired number of segments without switches
    if n_without_switches > 0:
        if n_without_switches >= desired_without_switches:
            # We have enough, randomly sample
            random.seed(42)
            sampled_without = random.sample(segments_without_switches, desired_without_switches)
            balanced_segments.extend(sampled_without)
        else:
            # We don't have enough, use all available
            print(f"  WARNING: Only {n_without_switches} segments without switches available")
            print(f"           Wanted {desired_without_switches} to achieve {target_switch_ratio * 100:.0f}% ratio")
            balanced_segments.extend(segments_without_switches)

    # Shuffle the combined list
    random.seed(42)
    random.shuffle(balanced_segments)

    # Calculate actual ratio
    final_with_switches = sum(1 for s in balanced_segments if s['has_switch'])
    final_without_switches = len(balanced_segments) - final_with_switches
    actual_ratio = final_with_switches / len(balanced_segments) if len(balanced_segments) > 0 else 0

    print(f"\nBalanced dataset:")
    print(f"  Total segments: {len(balanced_segments)}")
    print(f"  With switches: {final_with_switches} ({actual_ratio * 100:.1f}%)")
    print(f"  Without switches: {final_without_switches} ({(1 - actual_ratio) * 100:.1f}%)")
    print(f"  Target was: {target_switch_ratio * 100:.0f}% with switches")

    if abs(actual_ratio - target_switch_ratio) > 0.05:
        print(f"  ⚠️ WARNING: Actual ratio differs from target by more than 5%")
    else:
        print(f"  ✓ Ratio is close to target")

    return balanced_segments


def create_train_val_test_split(df, train_ratio=0.7, val_ratio=0.15, test_ratio=0.15):
    """
    Create stratified split ensuring file diversity and no data leakage
    Each split MUST contain segments from multiple files
    """
    print(f"\n=== Creating Train/Val/Test Split ===")

    # Group by source file
    file_groups = {}
    for file_name in df['source_file'].unique():
        file_data = df[df['source_file'] == file_name].copy()
        file_groups[file_name] = file_data

    print(f"Total files: {len(file_groups)}")

    # Ensure we have enough files
    if len(file_groups) < 3:
        print("WARNING: Less than 3 files detected. Adding file diversity through segment splitting.")

    train_dfs = []
    val_dfs = []
    test_dfs = []

    # Strategy: Split SEGMENTS from EACH file across train/val/test
    # This ensures every split has diversity
    for file_name, file_data in file_groups.items():
        # Shuffle segments within this file
        file_data = file_data.sample(frac=1, random_state=42).reset_index(drop=True)

        n = len(file_data)
        n_train = int(n * train_ratio)
        n_val = int(n * val_ratio)

        # Split this file's segments
        train_segments = file_data.iloc[:n_train]
        val_segments = file_data.iloc[n_train:n_train + n_val]
        test_segments = file_data.iloc[n_train + n_val:]

        # Add to respective lists
        if len(train_segments) > 0:
            train_dfs.append(train_segments)
        if len(val_segments) > 0:
            val_dfs.append(val_segments)
        if len(test_segments) > 0:
            test_dfs.append(test_segments)

        print(f"  {file_name[:40]}...: {len(train_segments)} train, {len(val_segments)} val, {len(test_segments)} test")

    # Combine and shuffle
    train_df = pd.concat(train_dfs, ignore_index=True).sample(frac=1, random_state=42).reset_index(drop=True)
    val_df = pd.concat(val_dfs, ignore_index=True).sample(frac=1, random_state=42).reset_index(drop=True)
    test_df = pd.concat(test_dfs, ignore_index=True).sample(frac=1, random_state=42).reset_index(drop=True)

    # Verify file diversity
    train_files = train_df['source_file'].nunique()
    val_files = val_df['source_file'].nunique()
    test_files = test_df['source_file'].nunique()

    # Save splits
    train_df.to_csv('train_segments.csv', index=False)
    val_df.to_csv('val_segments.csv', index=False)
    test_df.to_csv('test_segments.csv', index=False)

    print(f"\nFinal split:")
    print(f"  Training: {len(train_df)} segments from {train_files} files")
    print(f"  Validation: {len(val_df)} segments from {val_files} files")
    print(f"  Test: {len(test_df)} segments from {test_files} files")

    # Verify diversity
    if train_files < 2 or val_files < 2 or test_files < 2:
        print("WARNING: Some splits have segments from fewer than 2 files!")
    else:
        print("✓ All splits have good file diversity")

    # Show file overlap
    train_file_set = set(train_df['source_file'].unique())
    val_file_set = set(val_df['source_file'].unique())
    test_file_set = set(test_df['source_file'].unique())

    common_files = train_file_set & val_file_set & test_file_set
    if common_files:
        print(f"✓ Files appearing in all splits (good for diversity): {len(common_files)}")

    return train_df, val_df, test_df


def remove_duplicate_sequences_from_train(train_df, test_df, min_duplicate_length=25):
    """
    Remove train segments that contain sequences of min_duplicate_length+ tokens
    identical to sequences in test set
    """
    print(f"\n{'=' * 80}")
    print(f"REMOVING TRAIN SEGMENTS WITH {min_duplicate_length}+ TOKEN DUPLICATES")
    print(f"{'=' * 80}")

    original_train_size = len(train_df)

    # Build index of all subsequences from test set
    test_sequences = defaultdict(list)

    for test_idx, test_row in test_df.iterrows():
        test_tokens = test_row['tokens'].split()

        # Extract all possible consecutive sequences of min_duplicate_length
        for i in range(len(test_tokens) - min_duplicate_length + 1):
            sequence = test_tokens[i:i + min_duplicate_length]
            sequence_str = ' '.join(sequence)
            sequence_hash = hashlib.md5(sequence_str.encode()).hexdigest()

            test_sequences[sequence_hash].append({
                'segment_id': test_row['segment_id'],
                'idx': test_idx,
                'position': i
            })

    print(f"Built index of test sequences (25+ tokens)")

    # Find train segments to remove
    train_segments_to_remove = set()
    duplicate_details = []

    for train_idx, train_row in train_df.iterrows():
        train_tokens = train_row['tokens'].split()

        # Check all possible consecutive sequences
        for i in range(len(train_tokens) - min_duplicate_length + 1):
            sequence = train_tokens[i:i + min_duplicate_length]
            sequence_str = ' '.join(sequence)
            sequence_hash = hashlib.md5(sequence_str.encode()).hexdigest()

            # If this sequence exists in test, mark train segment for removal
            if sequence_hash in test_sequences:
                train_segments_to_remove.add(train_idx)

                # Record details for reporting
                for test_match in test_sequences[sequence_hash]:
                    duplicate_details.append({
                        'train_idx': train_idx,
                        'train_segment': train_row['segment_id'],
                        'test_segment': test_match['segment_id'],
                        'duplicate_length': min_duplicate_length,
                        'sequence': sequence_str[:100] + '...' if len(sequence_str) > 100 else sequence_str
                    })
                break  # No need to check more sequences from this train segment

    # Remove problematic segments
    if train_segments_to_remove:
        print(f"\n❌ Found {len(train_segments_to_remove)} train segments with duplicates:")

        # Show details of what's being removed
        for detail in duplicate_details[:10]:  # Show first 10
            print(f"\n  Train segment: {detail['train_segment']}")
            print(f"  Has {detail['duplicate_length']}+ token match with test: {detail['test_segment']}")
            print(f"  Duplicate text: {detail['sequence']}")

        if len(duplicate_details) > 10:
            print(f"\n  ... and {len(duplicate_details) - 10} more duplicates")

        # Remove the segments
        clean_train_df = train_df.drop(list(train_segments_to_remove)).reset_index(drop=True)

        print(f"\n✅ Removed {len(train_segments_to_remove)} segments from training")
        print(f"  Original train size: {original_train_size}")
        print(f"  Cleaned train size: {len(clean_train_df)}")
        print(f"  Reduction: {(len(train_segments_to_remove) / original_train_size * 100):.1f}%")

    else:
        print("\n✅ No duplicate sequences found - train set is clean!")
        clean_train_df = train_df

    return clean_train_df


def validate_preprocessing(sample_text):
    """
    Validate preprocessing on a sample text
    """
    print("\n" + "="*60)
    print("VALIDATION: Testing preprocessing")
    print("="*60)
    print(f"Input: {sample_text}")

    tokens, labels = process_segment_to_tokens_labels_clean(sample_text)

    print(f"\nTokens ({len(tokens)}): {tokens}")
    print(f"Labels ({len(labels)}): {labels}")

    # Check for tag contamination
    contaminated = False
    for i, token in enumerate(tokens):
        if any(tag in token.lower() for tag in ['auto', 'allo', '<', '>']):
            print(f"⚠️ WARNING: Token {i} contains tag fragment: '{token}'")
            contaminated = True

    if not contaminated:
        print("✓ No tag contamination found in tokens")

    # Verify alignment
    if len(tokens) == len(labels):
        print("✓ Tokens and labels are aligned")
    else:
        print(f"✗ Misalignment: {len(tokens)} tokens vs {len(labels)} labels")

    return tokens, labels


# ============================================================================
# MAIN PREPROCESSING PIPELINE
# ============================================================================

def main(data_dir='dataset/annotated-data-raw',
         output_dir='dataset/preprocessed_clean',
         target_switch_ratio=0.67):
    """
    Main preprocessing pipeline - EXACT SAME as original code
    """
    print("="*80)
    print("TIBETAN CODE-SWITCHING DATA PREPROCESSING")
    print("="*80)

    os.makedirs(output_dir, exist_ok=True)

    # Step 1: Process all files (including segments without switches)
    df, all_segments = process_all_files(data_dir, require_switch=False)

    if len(df) == 0:
        print("ERROR: No segments found!")
        return

    # Step 1.5: Balance segments to achieve target ratio
    print("\nBalancing segments...")
    balanced_segments = balance_segments_by_switches(all_segments, target_switch_ratio)

    if len(balanced_segments) == 0:
        print("ERROR: No balanced segments available!")
        return

    # Convert balanced segments to DataFrame
    segments_data = []
    for segment in balanced_segments:
        segments_data.append({
            'segment_id': segment['segment_id'],
            'source_file': segment['source_file'],
            'file_type': segment['file_type'],
            'tokens': ' '.join(segment['tokens']),
            'labels': ','.join(map(str, segment['labels'])),
            'num_tokens': segment['num_tokens'],
            'num_transitions': segment['num_transitions'],
            'has_switch': segment['has_switch'],
            'original_text': segment['original_text']
        })

    df = pd.DataFrame(segments_data)

    # Save all segments
    df.to_csv(f'{output_dir}/all_segments_300_400_tokens.csv', index=False)

    # Step 2: Create train/val/test split with file diversity
    print("\nCreating train/val/test split with file diversity...")
    train_df, val_df, test_df = create_train_val_test_split(df)

    # Step 3: Check and remove duplicates
    print("\nChecking for duplicate sequences between train and test...")
    train_df_clean = remove_duplicate_sequences_from_train(train_df, test_df, min_duplicate_length=25)

    # Save the cleaned training set if needed
    if len(train_df_clean) < len(train_df):
        train_df = train_df_clean

    # Save final splits with updated names
    train_df.to_csv(f'{output_dir}/train_segments_updated_v1.csv', index=False)
    val_df.to_csv(f'{output_dir}/val_segments_updated_v1.csv', index=False)
    test_df.to_csv(f'{output_dir}/test_segments_updated_v1.csv', index=False)

    print(f"\nSaved preprocessed data to {output_dir}/")
    print(f"  - train_segments_updated_v1.csv")
    print(f"  - val_segments_updated_v1.csv")
    print(f"  - test_segments_updated_v1.csv")

    # Print final statistics
    print("\n" + "="*80)
    print("PREPROCESSING COMPLETE")
    print("="*80)

    # Analyze label distribution for each split
    for name, df_split in [('Train', train_df), ('Val', val_df), ('Test', test_df)]:
        all_labels = []
        for labels_str in df_split['labels']:
            all_labels.extend([int(l) for l in labels_str.split(',')])

        label_counts = {i: all_labels.count(i) for i in range(4)}
        total = len(all_labels)

        print(f"\n{name} Set Label Distribution:")
        if total > 0:
            print(f"  Class 0 (Auto-cont): {label_counts[0]} ({label_counts[0]/total*100:.1f}%)")
            print(f"  Class 1 (Allo-cont): {label_counts[1]} ({label_counts[1]/total*100:.1f}%)")
            print(f"  Class 2 (Switch→Auto): {label_counts[2]} ({label_counts[2]/total*100:.1f}%)")
            print(f"  Class 3 (Switch→Allo): {label_counts[3]} ({label_counts[3]/total*100:.1f}%)")

    return train_df, val_df, test_df


if __name__ == "__main__":
    # Test on sample text
    test_text = "<auto> བསྟོད་པ་ལས་ / དཔལ་ལྡན་དགྱེས་པའི་རྡོ་རྗེ་དང་ / / དམ་ཚིག་བཀོད་པའི་རྒྱལ་པོ་དང་ / / <allo> དཔའ་བོ་འཇིག་རྟེན་དབང་ཕྱུག་དང་ / / ཇོ་མོ་རྗེ་བཙུན་སྒྲོལ་མ་སོགས་ / <auto> ཞལ་གཟིགས་གནང་བ་ཐོབ་པས་ན་ /"
    tokens, labels = validate_preprocessing(test_text)

    # Run main preprocessing with exact parameters
    train_df, val_df, test_df = main(
        data_dir='dataset/annotated-data-raw',
        output_dir='dataset/preprocessed_clean',
        target_switch_ratio=0.67
    )