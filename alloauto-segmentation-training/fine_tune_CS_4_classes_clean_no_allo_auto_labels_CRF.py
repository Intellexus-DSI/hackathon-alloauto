# """
# BERT + CRF model with logical constraints
# Enforces transition rules at inference time through CRF layer
# """
#
# import torch
# import torch.nn as nn
# from torchcrf import CRF
# from transformers import (
#     AutoModel,
#     AutoTokenizer,
#     TrainingArguments,
#     Trainer,
#     DataCollatorForTokenClassification
# )
# from torch.utils.data import Dataset
# import pandas as pd
# import numpy as np
# import os
#
# os.environ["CUDA_VISIBLE_DEVICES"] = "3"
#
# print(f"Number of GPUs available: {torch.cuda.device_count()}")
# for i in range(torch.cuda.device_count()):
#     print(f"GPU {i}: {torch.cuda.get_device_name(i)}")
#
#
# class BertCrfForTokenClassification(nn.Module):
#     """
#     BERT encoder + CRF layer with impossible transitions blocked
#
#     Transition rules enforced:
#     - No consecutive switches (2→2, 2→3, 3→2, 3→3)
#     - After Switch→Auto (2), must be in Auto mode (0) or switch out (3)
#     - After Switch→Allo (3), must be in Allo mode (1) or switch out (2)
#     - Mode consistency: Auto (0) can't go to Allo (1) without switch
#     - Mode consistency: Allo (1) can't go to Auto (0) without switch
#     """
#
#     def __init__(self, model_name, num_labels=4):
#         super().__init__()
#         self.num_labels = num_labels
#         self.model_name = model_name  # Store for saving
#
#         # BERT encoder
#         self.bert = AutoModel.from_pretrained(model_name)
#         for param in self.bert.parameters():
#             param.requires_grad = True  # Ensure unfrozen
#         self.dropout = nn.Dropout(0.1)
#
#         # Emission layer
#         self.classifier = nn.Linear(self.bert.config.hidden_size, num_labels)
#
#         # CRF layer
#         self.crf = CRF(num_labels, batch_first=True)
#
#         # Apply hard constraints
#         self._init_transition_constraints()
#
#     def save_pretrained(self, save_directory):
#         """Save model to directory"""
#         import os
#         os.makedirs(save_directory, exist_ok=True)
#
#         # Save model state
#         model_path = os.path.join(save_directory, 'pytorch_model.bin')
#         torch.save({
#             'model_state_dict': self.state_dict(),
#             'model_name': self.model_name,
#             'num_labels': self.num_labels
#         }, model_path)
#
#         # Save config
#         config_path = os.path.join(save_directory, 'config.json')
#         import json
#         config = {
#             'model_type': 'bert_crf',
#             'model_name': self.model_name,
#             'num_labels': self.num_labels,
#             'hidden_size': self.bert.config.hidden_size
#         }
#         with open(config_path, 'w') as f:
#             json.dump(config, f, indent=2)
#
#         print(f"Model saved to {save_directory}")
#
#     @classmethod
#     def from_pretrained(cls, load_directory):
#         """Load model from directory"""
#         import os
#
#         model_path = os.path.join(load_directory, 'pytorch_model.bin')
#         checkpoint = torch.load(model_path, map_location='cpu')
#
#         # Create model
#         model = cls(
#             model_name=checkpoint['model_name'],
#             num_labels=checkpoint['num_labels']
#         )
#
#         # Load state
#         model.load_state_dict(checkpoint['model_state_dict'])
#
#         print(f"Model loaded from {load_directory}")
#         return model
#
#     def _init_transition_constraints(self):
#         """Initialize with stronger priors based on expected patterns"""
#         with torch.no_grad():
#             # Start with better initial values
#             IMPOSSIBLE = -10000.0
#             LIKELY = 2.0  # Encourage likely transitions
#             UNLIKELY = -2.0  # Discourage unlikely ones
#             NEUTRAL = 0.0
#
#             # Initialize all to neutral
#             self.crf.transitions.data.fill_(NEUTRAL)
#
#             # Block impossible transitions
#             self.crf.transitions.data[2, 2] = IMPOSSIBLE
#             self.crf.transitions.data[2, 3] = IMPOSSIBLE
#             self.crf.transitions.data[3, 2] = IMPOSSIBLE
#             self.crf.transitions.data[3, 3] = IMPOSSIBLE
#             self.crf.transitions.data[2, 1] = IMPOSSIBLE
#             self.crf.transitions.data[3, 0] = IMPOSSIBLE
#             self.crf.transitions.data[0, 1] = IMPOSSIBLE
#             self.crf.transitions.data[0, 2] = IMPOSSIBLE
#             self.crf.transitions.data[1, 0] = IMPOSSIBLE
#             self.crf.transitions.data[1, 3] = IMPOSSIBLE
#
#             # Encourage likely transitions
#             self.crf.transitions.data[0, 0] = LIKELY  # Auto stays Auto
#             self.crf.transitions.data[1, 1] = LIKELY  # Allo stays Allo
#             self.crf.transitions.data[2, 0] = LIKELY  # After switch to Auto, stay Auto
#             self.crf.transitions.data[3, 1] = LIKELY  # After switch to Allo, stay Allo
#
#             # Discourage (but don't block) switches - make them learnable
#             self.crf.transitions.data[0, 3] = UNLIKELY  # Auto to switch→Allo
#             self.crf.transitions.data[1, 2] = UNLIKELY  # Allo to switch→Auto
#
#             print("\n" + "=" * 80)
#             print("CRF INITIALIZED WITH STRONGER PRIORS")
#             print("=" * 80)
#     def forward(self, input_ids, attention_mask, labels=None):
#         """Forward pass with CRF"""
#         # Get BERT embeddings
#         outputs = self.bert(input_ids=input_ids, attention_mask=attention_mask)
#         sequence_output = outputs.last_hidden_state
#         sequence_output = self.dropout(sequence_output)
#
#         # Get emission scores
#         emissions = self.classifier(sequence_output)
#
#         if labels is not None:
#             # Create mask for CRF
#             mask = (labels != -100).bool()
#
#             # FIX: CRF requires first timestep to be valid (not padding)
#             # Ensure at least the first position is True
#             mask[:, 0] = True
#
#             # Prepare labels for CRF
#             labels_for_crf = labels.clone()
#             labels_for_crf[labels_for_crf == -100] = 0  # Replace -100 with 0
#
#             # CRF loss (negative log-likelihood)
#             loss = -self.crf(emissions, labels_for_crf, mask=mask, reduction='mean')
#
#             return {'loss': loss, 'logits': emissions}
#         else:
#             # Inference: Viterbi decoding
#             mask = attention_mask.bool()
#             mask[:, 0] = True  # Ensure first position is valid
#
#             predictions = self.crf.decode(emissions, mask=mask)
#             return {'logits': emissions, 'predictions': predictions}
#
# class CodeSwitchingDataset4Class(Dataset):
#     """Same dataset class as main training"""
#
#     def __init__(self, csv_file, tokenizer, max_length=512):
#         self.data = pd.read_csv(csv_file)
#         self.tokenizer = tokenizer
#         self.max_length = max_length
#
#     def __len__(self):
#         return len(self.data)
#
#     def __getitem__(self, idx):
#         row = self.data.iloc[idx]
#         tokens = row['tokens'].split()
#         labels = list(map(int, row['labels'].split(',')))
#
#         encoding = self.tokenizer(
#             tokens,
#             is_split_into_words=True,
#             padding='max_length',
#             truncation=True,
#             max_length=self.max_length
#         )
#
#         # Align labels
#         word_ids = encoding.word_ids()
#         aligned_labels = []
#         previous_word_idx = None
#
#         for word_idx in word_ids:
#             if word_idx is None:
#                 aligned_labels.append(-100)
#             elif word_idx != previous_word_idx:
#                 aligned_labels.append(labels[word_idx] if word_idx < len(labels) else -100)
#             else:
#                 aligned_labels.append(-100)
#             previous_word_idx = word_idx
#
#         return {
#             'input_ids': torch.tensor(encoding['input_ids']),
#             'attention_mask': torch.tensor(encoding['attention_mask']),
#             'labels': torch.tensor(aligned_labels)
#         }
#
#
# class CRFTrainer(Trainer):
#     """Trainer that handles CRF forward pass correctly"""
#
#     def __init__(self, *args, **kwargs):
#         if 'tokenizer' in kwargs and 'processing_class' not in kwargs:
#             kwargs['processing_class'] = kwargs.get('tokenizer')
#         super().__init__(*args, **kwargs)
#
#     def compute_loss(self, model, inputs, return_outputs=False, num_items_in_batch=None):
#         """CRF model returns loss directly"""
#         outputs = model(**inputs)
#         loss = outputs['loss']
#         return (loss, outputs) if return_outputs else loss
#
#     def prediction_step(self, model, inputs, prediction_loss_only, ignore_keys=None):
#         """Use Viterbi decoding for predictions"""
#         inputs = self._prepare_inputs(inputs)
#         labels = inputs.pop('labels')
#
#         with torch.no_grad():
#             # Get loss
#             outputs_with_loss = model(**inputs, labels=labels)
#             loss = outputs_with_loss['loss']
#
#             # Get Viterbi predictions
#             outputs = model(**inputs)
#             viterbi_predictions = outputs['predictions']  # List of variable-length sequences
#
#             # Convert to fixed-length tensor matching labels shape
#             batch_size, seq_len = labels.shape
#             pred_tensor = torch.full(
#                 (batch_size, seq_len),
#                 fill_value=-100,
#                 dtype=torch.long,
#                 device=labels.device
#             )
#
#             # Fill in predictions from Viterbi, aligning with non-padded positions
#             for i, pred_seq in enumerate(viterbi_predictions):
#                 # Find valid (non-padded) positions
#                 valid_mask = labels[i] != -100
#                 valid_positions = torch.where(valid_mask)[0]
#
#                 # Fill predictions up to the length we have
#                 pred_len = min(len(pred_seq), len(valid_positions))
#                 pred_tensor[i, valid_positions[:pred_len]] = torch.tensor(
#                     pred_seq[:pred_len],
#                     device=labels.device
#                 )
#
#             # Get logits
#             logits = outputs_with_loss['logits']
#
#         return (loss, pred_tensor, labels)
#
#
# def evaluate_with_proximity(true_labels, pred_labels, tolerance=5):
#     """Evaluate with 5-token proximity tolerance"""
#     true_labels = np.array(true_labels)
#     pred_labels = np.array(pred_labels)
#
#     true_auto = np.where(true_labels == 2)[0]
#     true_allo = np.where(true_labels == 3)[0]
#     pred_auto = np.where(pred_labels == 2)[0]
#     pred_allo = np.where(pred_labels == 3)[0]
#
#     matched_true = set()
#
#     # Match with tolerance
#     for pred_pos in pred_auto:
#         if len(true_auto) > 0:
#             distances = np.abs(true_auto - pred_pos)
#             min_dist = np.min(distances)
#             closest_true = true_auto[np.argmin(distances)]
#             if closest_true not in matched_true and min_dist <= tolerance:
#                 matched_true.add(closest_true)
#
#     for pred_pos in pred_allo:
#         if len(true_allo) > 0:
#             distances = np.abs(true_allo - pred_pos)
#             min_dist = np.min(distances)
#             closest_true = true_allo[np.argmin(distances)]
#             if closest_true not in matched_true and min_dist <= tolerance:
#                 matched_true.add(closest_true)
#
#     total_true = len(true_auto) + len(true_allo)
#     total_pred = len(pred_auto) + len(pred_allo)
#     total_matched = len(matched_true)
#
#     precision = total_matched / total_pred if total_pred > 0 else 0
#     recall = total_matched / total_true if total_true > 0 else 0
#     f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0
#
#     beta = 2
#     fbeta2 = ((1 + beta ** 2) * precision * recall /
#               (beta ** 2 * precision + recall)) if (precision + recall) > 0 else 0
#
#     return {
#         'precision': precision,
#         'recall': recall,
#         'f1': f1,
#         'fbeta2': fbeta2,
#         'true_switches': total_true,
#         'pred_switches': total_pred
#     }
#
#
# def compute_metrics(eval_pred):
#     predictions, labels = eval_pred
#
#     # Now both have same shape: (batch_size, seq_len)
#     # Flatten everything
#     predictions = predictions.flatten()
#     labels = labels.flatten()
#
#     # Remove padding
#     mask = labels != -100
#     predictions = predictions[mask]
#     labels = labels[mask]
#
#     # Accuracy
#     accuracy = (predictions == labels).mean()
#
#     # Proximity metrics
#     prox = evaluate_with_proximity(labels, predictions, tolerance=5)
#
#     return {
#         'accuracy': float(accuracy),
#         'precision': float(prox['precision']),
#         'recall': float(prox['recall']),
#         'f1': float(prox['f1']),
#         'fbeta2': float(prox['fbeta2']),
#         'true_switches': prox['true_switches'],
#         'pred_switches': prox['pred_switches']
#     }
#
# def train_crf_model():
#     print("=" * 80)
#     print("TRAINING BERT + CRF MODEL")
#     print("Using EXACT SAME data splits as main ALTO model")
#     print("=" * 80)
#
#     # Use the SAME files as your main training
#     if os.path.exists('train_segments_stratified.csv'):
#         print("\n✓ Found stratified splits")
#         TRAIN_FILE = 'train_segments_stratified.csv'
#         VAL_FILE = 'val_segments_stratified.csv'
#         TEST_FILE = 'test_segments_stratified.csv'
#     elif os.path.exists('train_segments.csv'):
#         print("\n✓ Found regular splits")
#         TRAIN_FILE = 'train_segments.csv'
#         VAL_FILE = 'val_segments.csv'
#         TEST_FILE = 'test_segments.csv'
#     else:
#         print("\n❌ ERROR: No data splits found!")
#         print("Please run your main training script first to create the splits.")
#         return None, None, None
#
#     OUTPUT_DIR = './alloauto-segmentation-training/fine_tuned_ALTO_models/crf_model'
#
#     # Use SAME base model as main training
#     model_name = 'OMRIDRORI/mbert-tibetan-continual-wylie-final'
#     print(f"\nLoading tokenizer from: {model_name}")
#     tokenizer = AutoTokenizer.from_pretrained(model_name)
#
#     print("Creating BERT + CRF model with constraints...")
#     model = BertCrfForTokenClassification(model_name, num_labels=4)
#
#     device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
#     model = model.to(device)
#     print(f"Using device: {device}")
#
#     # Load datasets
#     print(f"\nLoading datasets...")
#     train_dataset = CodeSwitchingDataset4Class(TRAIN_FILE, tokenizer)
#     val_dataset = CodeSwitchingDataset4Class(VAL_FILE, tokenizer)
#     test_dataset = CodeSwitchingDataset4Class(TEST_FILE, tokenizer)
#
#     print(f"  Train: {len(train_dataset)} segments")
#     print(f"  Val: {len(val_dataset)} segments")
#     print(f"  Test: {len(test_dataset)} segments")
#
#     data_collator = DataCollatorForTokenClassification(tokenizer)
#     optimizer_grouped_parameters = [
#         {
#             'params': [p for n, p in model.named_parameters() if 'crf' not in n],
#             'lr': 2e-5  # Lower for BERT
#         },
#         {
#             'params': [p for n, p in model.named_parameters() if 'crf' in n],
#             'lr': 1e-3  # Much higher for CRF
#         }
#     ]
#
#     from torch.optim import AdamW
#     optimizer = AdamW(optimizer_grouped_parameters, weight_decay=0.01)
#
#     # Training arguments
#     training_args = TrainingArguments(
#         output_dir=OUTPUT_DIR,
#         eval_strategy="epoch",  # Changed from steps
#         save_strategy="no",
#         learning_rate=5e-5,  # Higher
#         per_device_train_batch_size=8,
#         per_device_eval_batch_size=8,
#         num_train_epochs=40,  # More epochs
#         weight_decay=0.01,
#         warmup_ratio=0.15,  # More warmup
#         metric_for_best_model='fbeta2',
#         fp16=torch.cuda.is_available(),
#         report_to=[]
#     )
#
#     # Trainer
#     trainer = CRFTrainer(
#         model=model,
#         args=training_args,
#         train_dataset=train_dataset,
#         eval_dataset=val_dataset,
#         data_collator=data_collator,
#         processing_class=tokenizer,
#         compute_metrics=compute_metrics,
#         optimizers=(optimizer, None)  # Custom optimizer
#     )
#
#     # Train
#     print("\nStarting training with CRF constraints...")
#     print("Viterbi decoding will enforce logical transitions during inference")
#     trainer.train()
#
#     # Save
#     print(f"\nSaving model to {OUTPUT_DIR}/final_model...")
#     model.save_pretrained(f'{OUTPUT_DIR}/final_model')
#     tokenizer.save_pretrained(f'{OUTPUT_DIR}/final_model')
#
#     print(f"✅ Model saved successfully")
#
#     # Evaluate
#     print("\nEvaluating on test set...")
#     test_results = trainer.evaluate(eval_dataset=test_dataset)
#
#     print(f"\n=== CRF Model Test Results ===")
#     print(f"Accuracy: {test_results['eval_accuracy']:.3f}")
#     print(f"Precision: {test_results['eval_precision']:.3f}")
#     print(f"Recall: {test_results['eval_recall']:.3f}")
#     print(f"F1: {test_results['eval_f1']:.3f}")
#     print(f"F-beta(2): {test_results['eval_fbeta2']:.3f}")
#
#     # Show learned transitions
#     print("\n" + "=" * 80)
#     print("LEARNED CRF TRANSITION MATRIX")
#     print("=" * 80)
#     transitions = model.crf.transitions.data.cpu().numpy()
#
#     labels = ['Auto(0)', 'Allo(1)', '→Auto(2)', '→Allo(3)']
#     print(f"\n{'From/To':<12}", end='')
#     for to_label in labels:
#         print(f"{to_label:>12}", end='')
#     print()
#     print("-" * 60)
#
#     for i, from_label in enumerate(labels):
#         print(f"{from_label:<12}", end='')
#         for j in range(4):
#             score = transitions[i, j]
#             if score < -1000:
#                 print(f"{'BLOCKED':>12}", end='')
#             else:
#                 print(f"{score:>12.2f}", end='')
#         print()
#
#     print("\nBLOCKED transitions remain impossible due to hard constraints")
#
#     return trainer, model, tokenizer
#
#
# if __name__ == "__main__":
#     trainer, model, tokenizer = train_crf_model()

"""
4-Class Code-Switching Detection System for Tibetan Text
Classes:
  0: Non-switching Auto (continuing in Auto mode)
  1: Non-switching Allo (continuing in Allo mode)
  2: Switch TO Auto
  3: Switch TO Allo

Focus: Proximity-aware loss and evaluation with 5-token tolerance
"""
#MOST updated model! 18/9/2025"
import torch
import torch.nn as nn
import torch.nn.functional as F
import numpy as np
import pandas as pd
import json
import re
import docx
import os
import ssl
from typing import List, Tuple, Dict
from collections import defaultdict
from sklearn.metrics import classification_report
from transformers import (
    AutoTokenizer,
    AutoModelForTokenClassification,
    TrainingArguments,
    DataCollatorForTokenClassification,
    Trainer
)
from torch.utils.data import Dataset
from pathlib import Path

# Environment setup
os.environ["CUDA_VISIBLE_DEVICES"] = "0"
ssl._create_default_https_context = ssl._create_unverified_context

print(f"Number of GPUs available: {torch.cuda.device_count()}")
for i in range(torch.cuda.device_count()):
    print(f"GPU {i}: {torch.cuda.get_device_name(i)}")


# ============================================================================
# PART 1: DATA PROCESSING
# ============================================================================

def aggressive_tag_removal(text):
    """
    Aggressively remove ALL forms of tags and partial tags from text
    """
    tag_patterns = [
        # Full tags with any case and spacing
        r'<\s*[Aa][Uu][Tt][Oo]\s*>',
        r'<\s*[Aa][Ll][Ll][Oo]\s*>',
        # Partial tags (opening)
        r'<\s*[Aa][Uu][Tt][Oo]\b',
        r'<\s*[Aa][Ll][Ll][Oo]\b',
        r'<\s*[Aa][Uu][Tt]\b',
        r'<\s*[Aa][Ll][Ll]\b',
        r'<\s*[Aa][Uu]\b',
        r'<\s*[Aa][Ll]\b',
        # Partial tags (closing)
        r'\b[Aa][Uu][Tt][Oo]\s*>',
        r'\b[Aa][Ll][Ll][Oo]\s*>',
        r'\b[Uu][Tt][Oo]\s*>',
        r'\b[Ll][Ll][Oo]\s*>',
        # Any word containing tag patterns
        r'\S*<[Aa][Uu][Tt][Oo]\S*',
        r'\S*<[Aa][Ll][Ll][Oo]\S*',
        r'\S*[Aa][Uu][Tt][Oo]>\S*',
        r'\S*[Aa][Ll][Ll][Oo]>\S*',
    ]

    cleaned_text = text
    for pattern in tag_patterns:
        cleaned_text = re.sub(pattern, ' ', cleaned_text)

    cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
    return cleaned_text.strip()


def process_segment_to_tokens_labels_clean(segment_text):
    """
    Convert segment to tokens and labels, ensuring NO tags in tokens
    """
    # Normalize tags for consistent splitting
    normalized_text = re.sub(r'<\s*AUTO\s*>', '<auto>', segment_text, flags=re.IGNORECASE)
    normalized_text = re.sub(r'<\s*ALLO\s*>', '<allo>', normalized_text, flags=re.IGNORECASE)

    # Split by tags
    parts = re.split(r'(<auto>|<allo>)', normalized_text.lower())

    tokens = []
    labels = []
    current_mode = None

    for i, part in enumerate(parts):
        if part == '<auto>':
            continue  # Skip tag
        elif part == '<allo>':
            continue  # Skip tag
        elif part.strip():
            # Clean any remaining tag fragments
            clean_part = aggressive_tag_removal(part)
            words = clean_part.strip().split()

            # Determine mode from previous tags
            segment_mode = None
            for j in range(i - 1, -1, -1):
                if parts[j] == '<auto>':
                    segment_mode = 'auto'
                    break
                elif parts[j] == '<allo>':
                    segment_mode = 'allo'
                    break

            if segment_mode is None:
                segment_mode = 'auto'

            for word_idx, word in enumerate(words):
                # Double-check word doesn't contain tags
                if contains_any_tag_pattern(word):
                    continue  # Skip tag fragments

                tokens.append(word)

                # Assign label
                if word_idx == 0 and current_mode is not None and current_mode != segment_mode:
                    label = 2 if segment_mode == 'auto' else 3
                else:
                    label = 0 if segment_mode == 'auto' else 1

                labels.append(label)

            current_mode = segment_mode

    return tokens, labels

def contains_any_tag_pattern(text):
    """
    Check if text contains ANY form of tag or partial tag
    Returns True if ANY tag pattern is found
    """
    # Comprehensive list of patterns to check
    forbidden_patterns = [
        # Full tags
        r'<auto>', r'<AUTO>', r'<allo>', r'<ALLO>',
        r'<Auto>', r'<Allo>',
        # With spaces
        r'<\s+auto\s+>', r'<\s+AUTO\s+>', r'<\s+allo\s+>', r'<\s+ALLO\s+>',
        # Partial tags (missing closing)
        r'<auto\b', r'<AUTO\b', r'<allo\b', r'<ALLO\b',
        r'<aut\b', r'<AUT\b', r'<all\b', r'<ALL\b',
        r'<au\b', r'<AU\b', r'<al\b', r'<AL\b',
        # Partial tags (missing opening)
        r'\bauto>', r'\bAUTO>', r'\ballo>', r'\bALLO>',
        r'\buto>', r'\bUTO>', r'\bllo>', r'\bLLO>',
        r'\bto>', r'\bTO>', r'\blo>', r'\bLO>',
        # Check for fragments within words
        r'<auto', r'<AUTO', r'<allo', r'<ALLO',
        r'auto>', r'AUTO>', r'allo>', r'ALLO>',
    ]

    text_lower = text.lower()

    for pattern in forbidden_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return True

    # Also check for simple substring presence
    forbidden_substrings = [
        '<auto', '<AUTO', '<allo', '<ALLO',
        'auto>', 'AUTO>', 'allo>', 'ALLO>',
        '<Auto', '<Allo', 'Auto>', 'Allo>'
    ]

    for substring in forbidden_substrings:
        if substring in text or substring.lower() in text_lower:
            return True

    return False
def process_segment_to_tokens_labels_v2(segment_text):
    """
    Convert a segment to tokens and 4-class labels
    MODIFIED: Tags are used for labeling but NOT included in tokens
    """
    # Split by tags while keeping them
    parts = re.split(r'(<auto>|<allo>)', segment_text)

    tokens = []
    labels = []
    current_mode = None

    for i, part in enumerate(parts):
        if part == '<auto>':
            # Don't add tag to tokens, just note mode change
            continue
        elif part == '<allo>':
            # Don't add tag to tokens, just note mode change
            continue
        elif part.strip():
            words = part.strip().split()

            # Determine what mode this segment should be in
            segment_mode = None
            for j in range(i - 1, -1, -1):
                if parts[j] == '<auto>':
                    segment_mode = 'auto'
                    break
                elif parts[j] == '<allo>':
                    segment_mode = 'allo'
                    break

            if segment_mode is None:
                segment_mode = 'auto'  # Default to auto if no tag found before

            for word_idx, word in enumerate(words):
                # Add the actual word (no tags)
                tokens.append(word)

                # Determine label based on position and mode change
                if word_idx == 0 and current_mode is not None and current_mode != segment_mode:
                    # This is the FIRST word after a mode change - it's a switch point
                    label = 2 if segment_mode == 'auto' else 3  # Switch TO auto or Switch TO allo
                else:
                    # Continuation in current mode
                    label = 0 if segment_mode == 'auto' else 1  # Non-switch auto or Non-switch allo

                labels.append(label)

            current_mode = segment_mode

    return tokens, labels


def verify_segment_has_switch(segment_text):
    """
    Verify that a segment contains at least one actual switch point
    Returns True if segment has at least one transition between auto and allo
    """
    # Must have both tags
    if '<auto>' not in segment_text or '<allo>' not in segment_text:
        return False

    # Check for actual switches by looking at tag sequences
    parts = re.split(r'(<auto>|<allo>)', segment_text)

    last_tag = None
    switches_found = 0

    for part in parts:
        if part == '<auto>':
            if last_tag == '<allo>':
                switches_found += 1
            last_tag = '<auto>'
        elif part == '<allo>':
            if last_tag == '<auto>':
                switches_found += 1
            last_tag = '<allo>'

    return switches_found > 0


def process_file_to_segments(file_path, file_type):
    """
    Process a single file and extract segments with guaranteed switches
    MODIFIED: Ensures tokens don't contain ANY form of tags
    """
    print(f"Processing {file_type} file: {file_path.name}")

    # Read file content
    if file_type == 'docx':
        doc = docx.Document(str(file_path))
        text_content = ' '.join([para.text for para in doc.paragraphs])
    else:  # txt file
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        text_content = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text_content = f.read()
                break
            except UnicodeDecodeError:
                continue

        if text_content is None:
            raise ValueError(f"Could not read {file_path.name} with any encoding")

    # Clean and normalize
    text_content = clean_and_normalize_text(text_content)

    # Split into sentences
    sentences = split_into_sentences(text_content)
    print(f"  Found {len(sentences)} sentences")

    # Extract segments with token limit and switch verification
    segments = extract_segments_with_token_limit(sentences, min_tokens=300, max_tokens=400)
    print(f"  Found {len(segments)} segments with verified switches (300-400 tokens each)")

    # Convert segments to token-label pairs
    processed_segments = []
    segments_without_switches = 0
    segments_with_tags = 0

    for seg_idx, segment in enumerate(segments):
        # Use the clean processing function
        tokens, labels = process_segment_to_tokens_labels_clean(segment['text'])

        if len(tokens) > 0:
            # Count actual transitions in labels
            num_transitions = sum(1 for l in labels if l in [2, 3])

            # Double-check that we have switches
            if num_transitions == 0:
                segments_without_switches += 1
                continue  # Skip segments without actual switches

            # CRITICAL: Verify no tags in tokens
            has_tags = False
            for token in tokens:
                if contains_any_tag_pattern(token):
                    has_tags = True
                    segments_with_tags += 1
                    print(f"    WARNING: Found tag in token: '{token}'")
                    break

            if has_tags:
                continue  # Skip this segment entirely

            # Create text WITHOUT tags for storage
            # Join tokens directly (they're already clean)
            text_without_tags = ' '.join(tokens)

            # Final verification
            if contains_any_tag_pattern(text_without_tags):
                print(f"    WARNING: Tags still found after cleaning in segment {seg_idx}")
                segments_with_tags += 1
                continue

            processed_segments.append({
                'segment_id': f"{file_path.stem}_{seg_idx}",
                'source_file': file_path.name,
                'file_type': file_type,
                'tokens': tokens,  # These are now completely tag-free tokens
                'labels': labels,
                'num_tokens': len(tokens),
                'num_transitions': num_transitions,
                'original_text': text_without_tags  # Store clean text
            })

    if segments_without_switches > 0:
        print(f"  Warning: Filtered out {segments_without_switches} segments without switches")

    if segments_with_tags > 0:
        print(f"  Warning: Filtered out {segments_with_tags} segments that still contained tags")

    print(f"  Processed {len(processed_segments)} valid, clean segments (all have switches, no tags)")
    return processed_segments


def process_file_to_segments_v1(file_path, file_type):
    """
    Process a single file and extract segments with guaranteed switches
    MODIFIED: Ensures tokens don't contain tags
    """
    print(f"Processing {file_type} file: {file_path.name}")

    # Read file content
    if file_type == 'docx':
        doc = docx.Document(str(file_path))
        text_content = ' '.join([para.text for para in doc.paragraphs])
    else:  # txt file
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        text_content = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text_content = f.read()
                break
            except UnicodeDecodeError:
                continue

        if text_content is None:
            raise ValueError(f"Could not read {file_path.name} with any encoding")

    # Clean and normalize
    text_content = clean_and_normalize_text(text_content)

    # Split into sentences
    sentences = split_into_sentences(text_content)
    print(f"  Found {len(sentences)} sentences")

    # Extract segments with token limit and switch verification
    segments = extract_segments_with_token_limit(sentences, min_tokens=300, max_tokens=400)
    print(f"  Found {len(segments)} segments with verified switches (300-400 tokens each)")

    # Convert segments to token-label pairs
    processed_segments = []
    segments_without_switches = 0

    for seg_idx, segment in enumerate(segments):
        tokens, labels = process_segment_to_tokens_labels_v2(segment['text'])

        if len(tokens) > 0:
            # Count actual transitions in labels
            num_transitions = sum(1 for l in labels if l in [2, 3])

            # Double-check that we have switches
            if num_transitions == 0:
                segments_without_switches += 1
                continue  # Skip segments without actual switches

            # Create text WITHOUT tags for storage
            text_without_tags = segment['text']
            text_without_tags = re.sub(r'<auto>', '', text_without_tags)
            text_without_tags = re.sub(r'<allo>', '', text_without_tags)
            text_without_tags = re.sub(r'\s+', ' ', text_without_tags).strip()

            processed_segments.append({
                'segment_id': f"{file_path.stem}_{seg_idx}",
                'source_file': file_path.name,
                'file_type': file_type,
                'tokens': tokens,  # These are now tag-free tokens
                'labels': labels,
                'num_tokens': len(tokens),
                'num_transitions': num_transitions,
                'original_text': text_without_tags  # Store text without tags
            })

    if segments_without_switches > 0:
        print(f"  Warning: Filtered out {segments_without_switches} segments without switches")

    print(f"  Processed {len(processed_segments)} valid segments (all have switches)")
    return processed_segments


def validate_preprocessing(sample_segment_text):
    """
    Helper function to validate that preprocessing works correctly
    Shows how tags are used for labeling but not included in tokens
    """
    print("\n=== PREPROCESSING VALIDATION ===")
    print(f"Original text with tags:\n{sample_segment_text}\n")

    tokens, labels = process_segment_to_tokens_labels_clean(sample_segment_text)

    print("Tokens (no tags):")
    for i, (token, label) in enumerate(zip(tokens[:20], labels[:20])):  # Show first 20
        label_names = {
            0: 'NonSwitch-Auto',
            1: 'NonSwitch-Allo',
            2: 'SWITCH→Auto',
            3: 'SWITCH→Allo'
        }
        print(f"  {i:3d}: '{token:15s}' -> {label_names[label]}")

    # Verify no tags in tokens
    has_tags = any('<auto>' in t.lower() or '<allo>' in t.lower() for t in tokens)
    print(f"\nTokens contain tags: {has_tags} (should be False)")

    # Count switches
    num_switches = sum(1 for l in labels if l in [2, 3])
    print(f"Number of switch points: {num_switches}")

    return tokens, labels



def analyze_and_balance_switch_distribution(train_df, val_df, test_df):
    """
    Analyze switch type distribution and check balance
    """
    print("\n" + "=" * 80)
    print("SWITCH TYPE DISTRIBUTION ANALYSIS")
    print("=" * 80)

    def get_switch_stats(df, split_name):
        """Get detailed switch statistics for a split"""
        total_to_auto = 0
        total_to_allo = 0
        segments_with_auto = 0
        segments_with_allo = 0

        for idx in range(len(df)):
            labels = [int(l) for l in df.iloc[idx]['labels'].split(',')]

            # Count switches in this segment
            to_auto_count = labels.count(2)
            to_allo_count = labels.count(3)

            total_to_auto += to_auto_count
            total_to_allo += to_allo_count

            if to_auto_count > 0:
                segments_with_auto += 1
            if to_allo_count > 0:
                segments_with_allo += 1

        total_switches = total_to_auto + total_to_allo

        stats = {
            'split': split_name,
            'n_segments': len(df),
            'total_switches': total_switches,
            'to_auto': total_to_auto,
            'to_allo': total_to_allo,
            'auto_pct': (total_to_auto / total_switches * 100) if total_switches > 0 else 0,
            'allo_pct': (total_to_allo / total_switches * 100) if total_switches > 0 else 0,
            'avg_switches_per_seg': total_switches / len(df) if len(df) > 0 else 0,
            'avg_auto_per_seg': total_to_auto / len(df) if len(df) > 0 else 0,
            'avg_allo_per_seg': total_to_allo / len(df) if len(df) > 0 else 0,
            'segs_with_auto': segments_with_auto,
            'segs_with_allo': segments_with_allo
        }
        return stats

    # Get stats for each split
    train_stats = get_switch_stats(train_df, 'Train')
    val_stats = get_switch_stats(val_df, 'Val')
    test_stats = get_switch_stats(test_df, 'Test')

    # Print detailed statistics
    print("\n📊 DETAILED STATISTICS BY SPLIT:\n")
    print(f"{'Split':<8} {'Segments':<10} {'Total SW':<10} {'→Auto':<10} {'→Allo':<10} {'Auto%':<10} {'Allo%':<10}")
    print("-" * 70)

    for stats in [train_stats, val_stats, test_stats]:
        print(f"{stats['split']:<8} {stats['n_segments']:<10} {stats['total_switches']:<10} "
              f"{stats['to_auto']:<10} {stats['to_allo']:<10} "
              f"{stats['auto_pct']:<10.1f} {stats['allo_pct']:<10.1f}")

    print("\n📈 AVERAGE SWITCHES PER SEGMENT:\n")
    print(f"{'Split':<8} {'Avg Total':<12} {'Avg →Auto':<12} {'Avg →Allo':<12}")
    print("-" * 45)

    for stats in [train_stats, val_stats, test_stats]:
        print(f"{stats['split']:<8} {stats['avg_switches_per_seg']:<12.2f} "
              f"{stats['avg_auto_per_seg']:<12.2f} {stats['avg_allo_per_seg']:<12.2f}")

    print("\n📑 SEGMENT COVERAGE:\n")
    print(f"{'Split':<8} {'Has →Auto':<15} {'Has →Allo':<15}")
    print("-" * 40)

    for stats in [train_stats, val_stats, test_stats]:
        print(f"{stats['split']:<8} {stats['segs_with_auto']}/{stats['n_segments']:<10} "
              f"{stats['segs_with_allo']}/{stats['n_segments']}")

    # Check for distribution imbalance
    auto_pcts = [train_stats['auto_pct'], val_stats['auto_pct'], test_stats['auto_pct']]
    allo_pcts = [train_stats['allo_pct'], val_stats['allo_pct'], test_stats['allo_pct']]

    max_auto_diff = max(auto_pcts) - min(auto_pcts)
    max_allo_diff = max(allo_pcts) - min(allo_pcts)

    print("\n⚖️ DISTRIBUTION BALANCE CHECK:")
    print(f"  Max difference in Auto%: {max_auto_diff:.1f}%")
    print(f"  Max difference in Allo%: {max_allo_diff:.1f}%")
    print(f"  Train vs Test Auto% difference: {abs(train_stats['auto_pct'] - test_stats['auto_pct']):.1f}%")
    print(f"  Train vs Test Allo% difference: {abs(train_stats['allo_pct'] - test_stats['allo_pct']):.1f}%")

    if max_auto_diff > 5 or max_allo_diff > 5:
        print("  ⚠️ WARNING: Imbalanced distribution detected (>5% difference)")
        print("  Consider using stratified splitting based on switch types")
    else:
        print("  ✅ Distribution is well balanced across splits")

    return train_stats, val_stats, test_stats


def create_stratified_switch_split(df, train_ratio=0.7, val_ratio=0.15, test_ratio=0.15):
    """
    Create stratified split that maintains switch type distribution
    """
    print("\n" + "=" * 80)
    print("CREATING STRATIFIED SPLIT BY SWITCH TYPE")
    print("=" * 80)

    # Calculate switch type ratio for each segment
    segment_features = []

    for idx in range(len(df)):
        row = df.iloc[idx]
        labels = [int(l) for l in row['labels'].split(',')]

        to_auto = labels.count(2)
        to_allo = labels.count(3)
        total_switches = to_auto + to_allo

        # Categorize segments by dominant switch type
        if total_switches == 0:
            category = 'no_switch'  # Shouldn't happen with your filtering
        elif to_auto > 0 and to_allo == 0:
            category = 'auto_only'
        elif to_allo > 0 and to_auto == 0:
            category = 'allo_only'
        elif to_auto > to_allo:
            category = 'auto_dominant'
        elif to_allo > to_auto:
            category = 'allo_dominant'
        else:
            category = 'balanced'

        segment_features.append({
            'idx': idx,
            'category': category,
            'to_auto': to_auto,
            'to_allo': to_allo,
            'auto_ratio': to_auto / total_switches if total_switches > 0 else 0
        })

    # Add category to dataframe
    df['switch_category'] = [sf['category'] for sf in segment_features]

    # Print category distribution
    print("\nSegment categories:")
    category_counts = df['switch_category'].value_counts()
    for cat, count in category_counts.items():
        print(f"  {cat}: {count} segments ({count / len(df) * 100:.1f}%)")

    # Stratified split by category
    # First split: train+val vs test
    train_val_df, test_df = train_test_split(
        df,
        test_size=test_ratio,
        stratify=df['switch_category'],
        random_state=42
    )

    # Second split: train vs val
    relative_val_size = val_ratio / (train_ratio + val_ratio)
    train_df, val_df = train_test_split(
        train_val_df,
        test_size=relative_val_size,
        stratify=train_val_df['switch_category'],
        random_state=42
    )

    # Remove temporary column
    train_df = train_df.drop('switch_category', axis=1)
    val_df = val_df.drop('switch_category', axis=1)
    test_df = test_df.drop('switch_category', axis=1)

    # Save splits
    train_df.to_csv('train_segments_stratified.csv', index=False)
    val_df.to_csv('val_segments_stratified.csv', index=False)
    test_df.to_csv('test_segments_stratified.csv', index=False)

    print(f"\n✅ Stratified splits created:")
    print(f"  Train: {len(train_df)} segments")
    print(f"  Val: {len(val_df)} segments")
    print(f"  Test: {len(test_df)} segments")

    return train_df, val_df, test_df

def clean_and_normalize_text(text_content):
    """
    Clean text and normalize tags, remove ALL newlines and extra spaces
    """
    # First normalize all tag variations
    text_content = re.sub(r'<\s*AUTO\s*>', '<auto>', text_content, flags=re.IGNORECASE)
    text_content = re.sub(r'<\s*ALLO\s*>', '<allo>', text_content, flags=re.IGNORECASE)

    # Remove ALL newlines and replace with spaces
    text_content = text_content.replace('\n', ' ')
    text_content = text_content.replace('\r', ' ')
    text_content = text_content.replace('\t', ' ')

    # Remove any separator lines (multiple dashes, underscores, etc)
    text_content = re.sub(r'[-_=]{3,}', ' ', text_content)

    # Clean up multiple spaces
    text_content = re.sub(r'\s+', ' ', text_content)

    # Ensure proper spacing around tags for splitting
    text_content = re.sub(r'<auto>', ' <auto> ', text_content)
    text_content = re.sub(r'<allo>', ' <allo> ', text_content)

    # Final cleanup of multiple spaces
    text_content = re.sub(r'\s+', ' ', text_content)

    return text_content.strip()


def split_into_sentences(text_content):
    """
    Split text into sentences based on . / // boundaries
    """
    sentences = []
    current_sentence = ""

    # First split by sentence endings, keeping the delimiters
    parts = re.split(r'(\.|//?)', text_content)

    for i, part in enumerate(parts):
        if part.strip() in ['.', '/', '//']:
            # End of sentence
            if current_sentence.strip():
                current_sentence += part
                sentences.append(current_sentence.strip())
                current_sentence = ""
        else:
            current_sentence += part

    # Add any remaining content
    if current_sentence.strip():
        sentences.append(current_sentence.strip())

    return [s for s in sentences if s.strip()]


def extract_segments_with_token_limit(sentences, min_tokens=300, max_tokens=400):
    """
    Extract segments of 300-400 tokens that MUST contain at least one switch
    """
    segments = []
    current_segment = []
    current_token_count = 0

    for sentence in sentences:
        # Count tokens in this sentence (simple whitespace tokenization for now)
        sentence_tokens = sentence.split()
        sentence_token_count = len(sentence_tokens)

        # If adding this sentence would exceed max_tokens
        if current_token_count + sentence_token_count > max_tokens:
            # Save current segment if it meets requirements
            if current_token_count >= min_tokens:
                segment_text = ' '.join(current_segment)

                # Verify this segment has at least one actual switch
                if verify_segment_has_switch(segment_text):
                    segments.append({
                        'text': segment_text,
                        'token_count': current_token_count,
                        'has_transition': True
                    })
                else:
                    # If no switch found, try to extend segment slightly to capture one
                    extended = extend_to_find_switch(current_segment, sentences,
                                                    sentences.index(sentence), max_tokens)
                    if extended:
                        segments.append(extended)

            # Start new segment with current sentence
            current_segment = [sentence]
            current_token_count = sentence_token_count
        else:
            # Add sentence to current segment
            current_segment.append(sentence)
            current_token_count += sentence_token_count

            # Check if we've reached the target range AND have a switch
            if current_token_count >= min_tokens:
                segment_text = ' '.join(current_segment)

                if verify_segment_has_switch(segment_text):
                    segments.append({
                        'text': segment_text,
                        'token_count': current_token_count,
                        'has_transition': True
                    })
                    # Start new segment
                    current_segment = []
                    current_token_count = 0

    # Handle remaining sentences
    if current_segment and current_token_count >= min_tokens // 2:  # Be lenient for last segment
        segment_text = ' '.join(current_segment)
        if verify_segment_has_switch(segment_text):
            segments.append({
                'text': segment_text,
                'token_count': current_token_count,
                'has_transition': True
            })

    return segments


def extend_to_find_switch(current_segment, all_sentences, current_idx, max_tokens):
    """
    Try to extend segment to include at least one switch
    """
    extended_segment = current_segment.copy()
    token_count = sum(len(s.split()) for s in extended_segment)

    # Look ahead up to 5 sentences
    for i in range(current_idx, min(current_idx + 5, len(all_sentences))):
        if i >= len(all_sentences):
            break

        sentence = all_sentences[i]
        sentence_tokens = len(sentence.split())

        # Don't exceed max tokens too much
        if token_count + sentence_tokens > max_tokens + 50:  # Allow slight overflow
            break

        extended_segment.append(sentence)
        token_count += sentence_tokens

        # Check if we now have a switch
        segment_text = ' '.join(extended_segment)
        if verify_segment_has_switch(segment_text):
            return {
                'text': segment_text,
                'token_count': token_count,
                'has_transition': True
            }

    return None


def process_segment_to_tokens_labels_old(segment_text):
    """
    Convert a segment to tokens and 4-class labels
    """
    # Split by tags while keeping them
    parts = re.split(r'(<auto>|<allo>)', segment_text)

    tokens = []
    labels = []
    current_mode = None

    for i, part in enumerate(parts):
        if part == '<auto>':
            continue
        elif part == '<allo>':
            continue
        elif part.strip():
            words = part.strip().split()

            # Determine what mode this segment should be in
            segment_mode = None
            for j in range(i - 1, -1, -1):
                if parts[j] == '<auto>':
                    segment_mode = 'auto'
                    break
                elif parts[j] == '<allo>':
                    segment_mode = 'allo'
                    break

            if segment_mode is None:
                segment_mode = 'auto'  # Default

            for word_idx, word in enumerate(words):
                tokens.append(word)

                # Determine label based on position and mode change
                if word_idx == 0 and current_mode is not None and current_mode != segment_mode:
                    # This is a switch point
                    label = 2 if segment_mode == 'auto' else 3
                else:
                    # Continuation
                    label = 0 if segment_mode == 'auto' else 1

                labels.append(label)

            current_mode = segment_mode

    return tokens, labels


def process_file_to_segments_old(file_path, file_type):
    """
    Process a single file and extract segments with guaranteed switches
    """
    print(f"Processing {file_type} file: {file_path.name}")

    # Read file content
    if file_type == 'docx':
        doc = docx.Document(str(file_path))
        text_content = ' '.join([para.text for para in doc.paragraphs])
    else:  # txt file
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        text_content = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text_content = f.read()
                break
            except UnicodeDecodeError:
                continue

        if text_content is None:
            raise ValueError(f"Could not read {file_path.name} with any encoding")

    # Clean and normalize
    text_content = clean_and_normalize_text(text_content)

    # Split into sentences
    sentences = split_into_sentences(text_content)
    print(f"  Found {len(sentences)} sentences")

    # Extract segments with token limit and switch verification
    segments = extract_segments_with_token_limit(sentences, min_tokens=300, max_tokens=400)
    print(f"  Found {len(segments)} segments with verified switches (300-400 tokens each)")

    # Convert segments to token-label pairs
    processed_segments = []
    segments_without_switches = 0

    for seg_idx, segment in enumerate(segments):
        tokens, labels = process_segment_to_tokens_labels_v2(segment['text'])

        if len(tokens) > 0:
            # Count actual transitions in labels
            num_transitions = sum(1 for l in labels if l in [2, 3])

            # Double-check that we have switches
            if num_transitions == 0:
                segments_without_switches += 1
                continue  # Skip segments without actual switches

            processed_segments.append({
                'segment_id': f"{file_path.stem}_{seg_idx}",
                'source_file': file_path.name,
                'file_type': file_type,
                'tokens': tokens,
                'labels': labels,
                'num_tokens': len(tokens),
                'num_transitions': num_transitions,
                'original_text': segment['text']
            })

    if segments_without_switches > 0:
        print(f"  Warning: Filtered out {segments_without_switches} segments without switches")

    print(f"  Processed {len(processed_segments)} valid segments (all have switches)")
    return processed_segments


def has_code_switching_tags(file_path, file_type):
    """Check if file contains AUTO/ALLO tags before processing"""
    try:
        if file_type == 'docx':
            doc = docx.Document(str(file_path))
            text_content = ' '.join([para.text for para in doc.paragraphs])
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()

        # Check for any variation of tags
        text_lower = text_content.lower()
        return ('<auto>' in text_lower or '<allo>' in text_lower)
    except:
        return False
def process_all_files(data_dir):
    """
    Process all .txt and .docx files in the data directory
    """
    data_path = Path(data_dir)

    # Find all files
    txt_files = list(data_path.glob("*.txt"))
    docx_files = list(data_path.glob("*.docx"))

    # all_files = [(f, 'txt') for f in txt_files] + [(f, 'docx') for f in docx_files]
    #

    all_files = [(f, 'txt') for f in txt_files] + [(f, 'docx') for f in docx_files]
    # Filter files with tags
    all_files = [(f, t) for f, t in all_files if has_code_switching_tags(f, t)]
    print(f"Files with code-switching tags: {len(all_files)}")
    print(f"\nFound {len(txt_files)} .txt files and {len(docx_files)} .docx files")
    print(f"Total files to process: {len(all_files)}")

    all_segments = []
    total_segments_without_switches = 0

    for file_path, file_type in all_files:
        try:
            file_segments = process_file_to_segments(file_path, file_type)
            all_segments.extend(file_segments)
        except Exception as e:
            print(f"  ERROR processing {file_path.name}: {e}")
            continue

    print(f"\n=== Processing Complete ===")
    print(f"Total segments with verified switches: {len(all_segments)}")

    if len(all_segments) == 0:
        print("ERROR: No valid segments with switches found!")
        return pd.DataFrame(), []

    # Create DataFrame
    segments_data = []
    min_transitions = float('inf')
    max_transitions = 0

    for segment in all_segments:
        min_transitions = min(min_transitions, segment['num_transitions'])
        max_transitions = max(max_transitions, segment['num_transitions'])

        segments_data.append({
            'segment_id': segment['segment_id'],
            'source_file': segment['source_file'],
            'file_type': segment['file_type'],
            'tokens': ' '.join(segment['tokens']),
            'labels': ','.join(map(str, segment['labels'])),
            'num_tokens': segment['num_tokens'],
            'num_transitions': segment['num_transitions'],
            'original_text': segment['original_text']
        })

    df = pd.DataFrame(segments_data)

    # Print statistics
    print(f"\n=== Dataset Statistics ===")
    print(f"Total segments: {len(df)} (ALL have at least 1 switch)")
    print(f"Total tokens: {df['num_tokens'].sum()}")
    print(f"Average tokens per segment: {df['num_tokens'].mean():.1f}")
    print(f"Total transitions: {df['num_transitions'].sum()}")
    print(f"Transitions per segment: min={min_transitions}, max={max_transitions}, avg={df['num_transitions'].mean():.1f}")
    print(f"Files represented: {df['source_file'].nunique()}")

    # Verify every segment has switches
    segments_with_no_switches = df[df['num_transitions'] == 0]
    if len(segments_with_no_switches) > 0:
        print(f"⚠️ WARNING: Found {len(segments_with_no_switches)} segments without switches (will be removed)")
        df = df[df['num_transitions'] > 0]
        print(f"After filtering: {len(df)} segments")

    # Label distribution
    all_labels = []
    for labels_str in df['labels']:
        all_labels.extend([int(l) for l in labels_str.split(',')])

    label_names = ['Non-switch Auto', 'Non-switch Allo', 'Switch to Auto', 'Switch to Allo']
    print(f"\nLabel distribution:")
    switch_count = 0
    for i in range(4):
        count = sum(1 for l in all_labels if l == i)
        percentage = count / len(all_labels) * 100
        print(f"  Class {i} ({label_names[i]}): {count} ({percentage:.2f}%)")
        if i >= 2:
            switch_count += count

    print(f"\nSwitch statistics:")
    print(f"  Total switch labels: {switch_count}")
    print(f"  Avg switches per segment: {switch_count / len(df):.1f}")
    print(f"  Switch density: {switch_count / len(all_labels) * 100:.2f}% of all tokens")

    return df, all_segments


def create_train_val_test_split(df, train_ratio=0.7, val_ratio=0.15, test_ratio=0.15):
    """
    Create stratified split ensuring file diversity and no data leakage
    Each split MUST contain segments from multiple files
    """
    print(f"\n=== Creating Train/Val/Test Split ===")

    # Group by source file
    file_groups = {}
    for file_name in df['source_file'].unique():
        file_data = df[df['source_file'] == file_name].copy()
        file_groups[file_name] = file_data

    print(f"Total files: {len(file_groups)}")

    # Ensure we have enough files
    if len(file_groups) < 3:
        print("WARNING: Less than 3 files detected. Adding file diversity through segment splitting.")

    train_dfs = []
    val_dfs = []
    test_dfs = []

    # Strategy: Split SEGMENTS from EACH file across train/val/test
    # This ensures every split has diversity
    for file_name, file_data in file_groups.items():
        # Shuffle segments within this file
        file_data = file_data.sample(frac=1, random_state=42).reset_index(drop=True)

        n = len(file_data)
        n_train = int(n * train_ratio)
        n_val = int(n * val_ratio)

        # Split this file's segments
        train_segments = file_data.iloc[:n_train]
        val_segments = file_data.iloc[n_train:n_train + n_val]
        test_segments = file_data.iloc[n_train + n_val:]

        # Add to respective lists
        if len(train_segments) > 0:
            train_dfs.append(train_segments)
        if len(val_segments) > 0:
            val_dfs.append(val_segments)
        if len(test_segments) > 0:
            test_dfs.append(test_segments)

        print(f"  {file_name[:40]}...: {len(train_segments)} train, {len(val_segments)} val, {len(test_segments)} test")

    # Combine and shuffle
    train_df = pd.concat(train_dfs, ignore_index=True).sample(frac=1, random_state=42).reset_index(drop=True)
    val_df = pd.concat(val_dfs, ignore_index=True).sample(frac=1, random_state=42).reset_index(drop=True)
    test_df = pd.concat(test_dfs, ignore_index=True).sample(frac=1, random_state=42).reset_index(drop=True)

    # Verify file diversity
    train_files = train_df['source_file'].nunique()
    val_files = val_df['source_file'].nunique()
    test_files = test_df['source_file'].nunique()

    # Save splits
    train_df.to_csv('train_segments.csv', index=False)
    val_df.to_csv('val_segments.csv', index=False)
    test_df.to_csv('test_segments.csv', index=False)

    print(f"\nFinal split:")
    print(f"  Training: {len(train_df)} segments from {train_files} files")
    print(f"  Validation: {len(val_df)} segments from {val_files} files")
    print(f"  Test: {len(test_df)} segments from {test_files} files")

    # Verify diversity
    if train_files < 2 or val_files < 2 or test_files < 2:
        print("WARNING: Some splits have segments from fewer than 2 files!")
    else:
        print("✓ All splits have good file diversity")

    # Show file overlap
    train_file_set = set(train_df['source_file'].unique())
    val_file_set = set(val_df['source_file'].unique())
    test_file_set = set(test_df['source_file'].unique())

    common_files = train_file_set & val_file_set & test_file_set
    if common_files:
        print(f"✓ Files appearing in all splits (good for diversity): {len(common_files)}")

    return train_df, val_df, test_df


# ============================================================================
# PART 2: DATASET CLASS
# ============================================================================

class CodeSwitchingDataset4Class(Dataset):
    """Dataset for 4-class token-level code-switching."""

    def __init__(self, csv_file, tokenizer, max_length=512):
        self.data = pd.read_csv(csv_file)
        self.tokenizer = tokenizer
        self.max_length = max_length

    def __len__(self):
        return len(self.data)

    def __getitem__(self, idx):
        row = self.data.iloc[idx]
        tokens = row['tokens'].split()
        labels = list(map(int, row['labels'].split(',')))

        encoding = self.tokenize_and_align_labels(tokens, labels)

        return {
            'input_ids': torch.tensor(encoding['input_ids']),
            'attention_mask': torch.tensor(encoding['attention_mask']),
            'labels': torch.tensor(encoding['labels'])
        }

    def tokenize_and_align_labels(self, tokens, labels):
        tokenized_inputs = self.tokenizer(
            tokens,
            is_split_into_words=True,
            padding='max_length',
            truncation=True,
            max_length=self.max_length
        )

        word_ids = tokenized_inputs.word_ids()
        aligned_labels = []
        previous_word_idx = None

        for word_idx in word_ids:
            if word_idx is None:
                aligned_labels.append(-100)
            elif word_idx != previous_word_idx:
                aligned_labels.append(labels[word_idx] if word_idx < len(labels) else -100)
            else:
                aligned_labels.append(-100)
            previous_word_idx = word_idx

        tokenized_inputs['labels'] = aligned_labels
        return tokenized_inputs


# ============================================================================
# PART 3: PROXIMITY-AWARE LOSS FUNCTION
# ============================================================================

def apply_transition_constraints(predictions, logits=None):
    """
    Apply logical constraints to predictions:
    - If in Auto mode (0), can only switch to Allo (3)
    - If in Allo mode (1), can only switch to Auto (2)
    """
    import torch

    if isinstance(predictions, torch.Tensor):
        predictions = predictions.cpu().numpy()

    corrected_predictions = predictions.copy()
    current_mode = 0  # Start in Auto by default

    for i in range(len(predictions)):
        pred = predictions[i]

        if pred == -100:  # Skip padding
            continue

        # Check for invalid transitions
        if pred == 2:  # Switch to Auto
            if current_mode == 0:  # Already in Auto - INVALID
                # Change to continuation in Auto
                corrected_predictions[i] = 0
            else:  # Was in Allo - VALID
                current_mode = 0

        elif pred == 3:  # Switch to Allo
            if current_mode == 1:  # Already in Allo - INVALID
                # Change to continuation in Allo
                corrected_predictions[i] = 1
            else:  # Was in Auto - VALID
                current_mode = 1

        elif pred == 0:  # Continue in Auto
            current_mode = 0

        elif pred == 1:  # Continue in Allo
            current_mode = 1

    return corrected_predictions


class ProximityAwareLoss4Class(nn.Module):
    """
    Loss function with logical transition constraints
    """

    def __init__(self, switch_loss_weight=30.0, proximity_tolerance=5,
                 distance_decay=0.7, false_positive_penalty=10.0,
                 invalid_transition_penalty=100.0):
        super().__init__()
        self.switch_loss_weight = switch_loss_weight
        self.proximity_tolerance = proximity_tolerance
        self.distance_decay = distance_decay
        self.false_positive_penalty = false_positive_penalty
        self.invalid_transition_penalty = invalid_transition_penalty

        # Much higher weights for switch classes to combat class imbalance
        self.class_weights = torch.tensor([
            1.0,  # Class 0: Non-switch auto
            1.0,  # Class 1: Non-switch allo
            switch_loss_weight,  # Class 2: Switch to auto (increased)
            switch_loss_weight   # Class 3: Switch to allo (increased)
        ])

    def forward(self, logits, labels):
        batch_size, seq_len, num_classes = logits.shape
        device = logits.device
        self.class_weights = self.class_weights.to(device)

        # Standard cross-entropy loss with class weights
        ce_loss = F.cross_entropy(
            logits.view(-1, num_classes),
            labels.view(-1),
            weight=self.class_weights,
            reduction='none'
        ).view(batch_size, seq_len)

        valid_mask = (labels != -100).float()

        # Get predictions
        predictions = torch.argmax(logits, dim=-1)

        # Apply proximity-aware adjustments
        proximity_adjusted_loss = ce_loss.clone()

        for b in range(batch_size):
            # Track current mode to detect invalid transitions
            current_true_mode = 0  # Start in Auto
            current_pred_mode = 0

            for t in range(seq_len):
                if labels[b, t] == -100:
                    continue

                true_label = labels[b, t].item()
                pred_label = predictions[b, t].item()

                # PENALIZE INVALID TRANSITIONS HEAVILY
                # Check predicted transitions
                if pred_label == 2:  # Predicted Switch to Auto
                    if current_pred_mode == 0:  # Already in Auto - INVALID!
                        proximity_adjusted_loss[b, t] *= self.invalid_transition_penalty
                    current_pred_mode = 0
                elif pred_label == 3:  # Predicted Switch to Allo
                    if current_pred_mode == 1:  # Already in Allo - INVALID!
                        proximity_adjusted_loss[b, t] *= self.invalid_transition_penalty
                    current_pred_mode = 1
                elif pred_label == 0:
                    current_pred_mode = 0
                elif pred_label == 1:
                    current_pred_mode = 1

                # REWARD CORRECT LOGICAL TRANSITIONS
                # Check if this is a true switch point
                if true_label == 2:  # True Switch to Auto
                    if current_true_mode == 1 and pred_label == 2:  # Correct logic
                        proximity_adjusted_loss[b, t] *= 0.1  # Strong reward
                    current_true_mode = 0
                elif true_label == 3:  # True Switch to Allo
                    if current_true_mode == 0 and pred_label == 3:  # Correct logic
                        proximity_adjusted_loss[b, t] *= 0.1  # Strong reward
                    current_true_mode = 1
                elif true_label == 0:
                    current_true_mode = 0
                elif true_label == 1:
                    current_true_mode = 1

            # Original proximity logic (but with type matching)
            true_switches_to_auto = torch.where(labels[b] == 2)[0]
            true_switches_to_allo = torch.where(labels[b] == 3)[0]
            pred_switches_to_auto = torch.where(predictions[b] == 2)[0]
            pred_switches_to_allo = torch.where(predictions[b] == 3)[0]

            # For predicted "switch to auto"
            for pred_pos in pred_switches_to_auto:
                if len(true_switches_to_auto) > 0:
                    distances = torch.abs(true_switches_to_auto - pred_pos)
                    min_distance = torch.min(distances).item()

                    if min_distance == 0:
                        proximity_adjusted_loss[b, pred_pos] *= 0.1
                    elif min_distance <= self.proximity_tolerance:
                        decay_factor = self.distance_decay ** min_distance
                        proximity_adjusted_loss[b, pred_pos] *= decay_factor
                    else:
                        proximity_adjusted_loss[b, pred_pos] *= self.false_positive_penalty
                else:
                    proximity_adjusted_loss[b, pred_pos] *= self.false_positive_penalty * 2

            # For predicted "switch to allo"
            for pred_pos in pred_switches_to_allo:
                if len(true_switches_to_allo) > 0:
                    distances = torch.abs(true_switches_to_allo - pred_pos)
                    min_distance = torch.min(distances).item()

                    if min_distance == 0:
                        proximity_adjusted_loss[b, pred_pos] *= 0.1
                    elif min_distance <= self.proximity_tolerance:
                        decay_factor = self.distance_decay ** min_distance
                        proximity_adjusted_loss[b, pred_pos] *= decay_factor
                    else:
                        proximity_adjusted_loss[b, pred_pos] *= self.false_positive_penalty
                else:
                    proximity_adjusted_loss[b, pred_pos] *= self.false_positive_penalty * 2
            # Penalize missed true switches with distance-aware scaling
            for true_pos in true_switches_to_auto:
                if len(pred_switches_to_auto) == 0:
                    # No predictions at all - heavy penalty
                    proximity_adjusted_loss[b, true_pos] *= 5.0
                else:
                    distances = torch.abs(pred_switches_to_auto - true_pos)
                    min_distance = torch.min(distances).item()

                    if min_distance > self.proximity_tolerance:
                        # Scale penalty based on how far the nearest prediction is
                        # Further away = higher penalty
                        distance_penalty = 2.0 + (min_distance - self.proximity_tolerance) * 0.3
                        distance_penalty = min(distance_penalty, 8.0)  # Cap at 8x
                        proximity_adjusted_loss[b, true_pos] *= distance_penalty
                    # If within tolerance, the positive proximity logic already handled it

            # Same for allo switches
            for true_pos in true_switches_to_allo:
                if len(pred_switches_to_allo) == 0:
                    proximity_adjusted_loss[b, true_pos] *= 5.0
                else:
                    distances = torch.abs(pred_switches_to_allo - true_pos)
                    min_distance = torch.min(distances).item()

                    if min_distance > self.proximity_tolerance:
                        distance_penalty = 2.0 + (min_distance - self.proximity_tolerance) * 0.3
                        distance_penalty = min(distance_penalty, 8.0)
                        proximity_adjusted_loss[b, true_pos] *= distance_penalty

            # Penalize missed true switches
            # for true_pos in true_switches_to_auto:
            #     if len(pred_switches_to_auto) == 0:
            #         proximity_adjusted_loss[b, true_pos] *= 3.0
            #     else:
            #         distances = torch.abs(pred_switches_to_auto - true_pos)
            #         min_distance = torch.min(distances).item()
            #         if min_distance > self.proximity_tolerance:
            #             proximity_adjusted_loss[b, true_pos] *= 2.0

            # for true_pos in true_switches_to_allo:
            #     if len(pred_switches_to_allo) == 0:
            #         proximity_adjusted_loss[b, true_pos] *= 3.0
            #     else:
            #         distances = torch.abs(pred_switches_to_allo - true_pos)
            #         min_distance = torch.min(distances).item()
            #         if min_distance > self.proximity_tolerance:
            #             proximity_adjusted_loss[b, true_pos] *= 2.0

        # Apply valid mask and compute mean
        total_loss = proximity_adjusted_loss * valid_mask
        loss = total_loss.sum() / valid_mask.sum().clamp(min=1)

        return loss


# ============================================================================
# PART 4: PROXIMITY-AWARE METRICS
# ============================================================================

def evaluate_switch_detection_with_proximity(true_labels, pred_labels, tolerance=5):
    """
    Evaluate switch detection with proximity tolerance and TYPE matching
    Switch types must match: class 2 with class 2, class 3 with class 3
    """
    true_labels = np.array(true_labels)
    pred_labels = np.array(pred_labels)

    # Find switch positions BY TYPE
    true_switches_to_auto = np.where(true_labels == 2)[0]  # Switch to Auto
    true_switches_to_allo = np.where(true_labels == 3)[0]  # Switch to Allo
    pred_switches_to_auto = np.where(pred_labels == 2)[0]  # Predicted Switch to Auto
    pred_switches_to_allo = np.where(pred_labels == 3)[0]  # Predicted Switch to Allo

    # Track matches separately by type
    matched_true_to_auto = set()
    matched_pred_to_auto = set()
    matched_true_to_allo = set()
    matched_pred_to_allo = set()

    exact_matches = 0
    proximity_matches = 0

    # Match "Switch to Auto" predictions with "Switch to Auto" ground truth
    for pred_pos in pred_switches_to_auto:
        if len(true_switches_to_auto) > 0:
            distances = np.abs(true_switches_to_auto - pred_pos)
            min_distance = np.min(distances)
            closest_true_idx = np.argmin(distances)
            closest_true_pos = true_switches_to_auto[closest_true_idx]

            # Only count if not already matched
            if closest_true_pos not in matched_true_to_auto:
                if min_distance == 0:
                    exact_matches += 1
                    matched_true_to_auto.add(closest_true_pos)
                    matched_pred_to_auto.add(pred_pos)
                elif min_distance <= tolerance:
                    proximity_matches += 1
                    matched_true_to_auto.add(closest_true_pos)
                    matched_pred_to_auto.add(pred_pos)

    # Match "Switch to Allo" predictions with "Switch to Allo" ground truth
    for pred_pos in pred_switches_to_allo:
        if len(true_switches_to_allo) > 0:
            distances = np.abs(true_switches_to_allo - pred_pos)
            min_distance = np.min(distances)
            closest_true_idx = np.argmin(distances)
            closest_true_pos = true_switches_to_allo[closest_true_idx]

            # Only count if not already matched
            if closest_true_pos not in matched_true_to_allo:
                if min_distance == 0:
                    exact_matches += 1
                    matched_true_to_allo.add(closest_true_pos)
                    matched_pred_to_allo.add(pred_pos)
                elif min_distance <= tolerance:
                    proximity_matches += 1
                    matched_true_to_allo.add(closest_true_pos)
                    matched_pred_to_allo.add(pred_pos)

    # Total counts
    total_true_switches = len(true_switches_to_auto) + len(true_switches_to_allo)
    total_pred_switches = len(pred_switches_to_auto) + len(pred_switches_to_allo)
    total_matched_true = len(matched_true_to_auto) + len(matched_true_to_allo)
    total_matched_pred = len(matched_pred_to_auto) + len(matched_pred_to_allo)

    total_matches = exact_matches + proximity_matches
    missed_switches = total_true_switches - total_matched_true
    false_switches = total_pred_switches - total_matched_pred

    # Calculate metrics
    precision = total_matches / total_pred_switches if total_pred_switches > 0 else 0
    recall = total_matches / total_true_switches if total_true_switches > 0 else 0
    f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0

    # Also calculate per-type metrics
    to_auto_precision = len(matched_pred_to_auto) / len(pred_switches_to_auto) if len(pred_switches_to_auto) > 0 else 0
    to_auto_recall = len(matched_true_to_auto) / len(true_switches_to_auto) if len(true_switches_to_auto) > 0 else 0
    to_allo_precision = len(matched_pred_to_allo) / len(pred_switches_to_allo) if len(pred_switches_to_allo) > 0 else 0
    to_allo_recall = len(matched_true_to_allo) / len(true_switches_to_allo) if len(true_switches_to_allo) > 0 else 0

    return {
        'exact_matches': exact_matches,
        'proximity_matches': proximity_matches,
        'total_matches': total_matches,
        'true_switches': total_true_switches,
        'pred_switches': total_pred_switches,
        'missed_switches': missed_switches,
        'false_switches': false_switches,
        'precision': precision,
        'recall': recall,
        'f1': f1,
        # Per-type metrics
        'true_to_auto': len(true_switches_to_auto),
        'true_to_allo': len(true_switches_to_allo),
        'pred_to_auto': len(pred_switches_to_auto),
        'pred_to_allo': len(pred_switches_to_allo),
        'matched_to_auto': len(matched_true_to_auto),
        'matched_to_allo': len(matched_true_to_allo),
        'to_auto_precision': to_auto_precision,
        'to_auto_recall': to_auto_recall,
        'to_allo_precision': to_allo_precision,
        'to_allo_recall': to_allo_recall
    }


def compute_metrics_for_trainer(eval_pred, tolerance=5):
    """
    Compute metrics for the trainer with proximity awareness and TYPE matching
    """
    predictions, labels = eval_pred
    predictions = np.argmax(predictions, axis=2)

    # Flatten all sequences
    all_predictions = predictions.flatten()
    all_labels = labels.flatten()

    # Remove padding
    mask = all_labels != -100
    all_predictions = all_predictions[mask]
    all_labels = all_labels[mask]

    # Basic accuracy
    accuracy = (all_predictions == all_labels).mean()

    # Proximity-aware switch metrics with TYPE matching
    switch_metrics = evaluate_switch_detection_with_proximity(
        all_labels, all_predictions, tolerance=tolerance
    )

    return {
        'accuracy': float(accuracy),
        'switch_precision': float(switch_metrics['precision']),
        'switch_recall': float(switch_metrics['recall']),
        'switch_f1': float(switch_metrics['f1']),
        'true_switches': switch_metrics['true_switches'],
        'pred_switches': switch_metrics['pred_switches'],
        'exact_matches': switch_metrics['exact_matches'],
        'proximity_matches': switch_metrics['proximity_matches'],
        # Per-type metrics
        'to_auto_precision': float(switch_metrics['to_auto_precision']),
        'to_auto_recall': float(switch_metrics['to_auto_recall']),
        'to_allo_precision': float(switch_metrics['to_allo_precision']),
        'to_allo_recall': float(switch_metrics['to_allo_recall']),
        'true_to_auto': switch_metrics['true_to_auto'],
        'true_to_allo': switch_metrics['true_to_allo'],
        'matched_to_auto': switch_metrics['matched_to_auto'],
        'matched_to_allo': switch_metrics['matched_to_allo']
    }

import torch
import torch.nn as nn
from torchcrf import CRF
from transformers import AutoModelForTokenClassification

# ============================================================================
# PART 5: CUSTOM TRAINER
# ============================================================================
class ProximityTrainedModelWithCRF(nn.Module):
    """
    Wraps your existing trained model and adds CRF for constrained inference
    Training: Uses proximity loss (keeps all your existing learning)
    Inference: Uses CRF Viterbi to enforce logical constraints
    """

    def __init__(self, pretrained_model_path):
        super().__init__()

        # Load your EXISTING proximity-trained model
        self.bert_model = AutoModelForTokenClassification.from_pretrained(pretrained_model_path)

        # Add CRF layer on top (only for inference)
        self.crf = CRF(num_tags=4, batch_first=True)

        # Initialize CRF transitions with constraints
        self._init_crf_constraints()

        # Flag to control whether to use CRF
        self.use_crf = False  # False during training, True during inference

    def _init_crf_constraints(self):
        """Set impossible transitions to very negative values"""
        with torch.no_grad():
            IMPOSSIBLE = -10000.0

            # Block invalid transitions (same as before)
            self.crf.transitions[2, 2] = IMPOSSIBLE  # Auto→Auto switch
            self.crf.transitions[2, 3] = IMPOSSIBLE
            self.crf.transitions[3, 2] = IMPOSSIBLE
            self.crf.transitions[3, 3] = IMPOSSIBLE
            self.crf.transitions[2, 1] = IMPOSSIBLE
            self.crf.transitions[3, 0] = IMPOSSIBLE
            self.crf.transitions[0, 1] = IMPOSSIBLE
            self.crf.transitions[0, 2] = IMPOSSIBLE
            self.crf.transitions[1, 0] = IMPOSSIBLE
            self.crf.transitions[1, 3] = IMPOSSIBLE

    def forward(self, input_ids, attention_mask, **kwargs):
        # Accept but ignore token_type_ids and other unexpected args
        # Get BERT predictions
        outputs = self.bert(input_ids=input_ids, attention_mask=attention_mask)
        logits = outputs.logits

        # Apply CRF Viterbi for constrained decoding
        mask = attention_mask.bool()
        predictions = self.crf.decode(logits, mask=mask)

        return {
            'logits': logits,
            'predictions': predictions
        }
class ProximityAwareTrainer4Class(Trainer):
    """Custom trainer with proximity-aware loss and transition constraints"""

    def __init__(self, switch_loss_weight=30.0, proximity_tolerance=5,
                 distance_decay=0.7, false_positive_penalty=10.0,
                 invalid_transition_penalty=100.0, *args, **kwargs):
        # Handle tokenizer/processing_class
        if 'tokenizer' in kwargs and 'processing_class' not in kwargs:
            kwargs['processing_class'] = kwargs.get('tokenizer')

        super().__init__(*args, **kwargs)

        # Initialize custom loss with all parameters including invalid transition penalty
        self.proximity_loss = ProximityAwareLoss4Class(
            switch_loss_weight=switch_loss_weight,
            proximity_tolerance=proximity_tolerance,
            distance_decay=distance_decay,
            false_positive_penalty=false_positive_penalty,
            invalid_transition_penalty=invalid_transition_penalty
        )
        self.proximity_tolerance = proximity_tolerance

    def compute_loss(self, model, inputs, return_outputs=False, num_items_in_batch=None):
        labels = inputs.pop("labels")
        outputs = model(**inputs)
        logits = outputs.get('logits')
        loss = self.proximity_loss(logits, labels)
        return (loss, outputs) if return_outputs else loss

    def prediction_step(self, model, inputs, prediction_loss_only, ignore_keys=None):
        """Override to apply transition constraints during evaluation"""
        outputs = super().prediction_step(model, inputs, prediction_loss_only, ignore_keys)

        if not prediction_loss_only:
            loss, logits, labels = outputs

            # Get raw predictions
            predictions = torch.argmax(logits, dim=-1)

            # Apply transition constraints to each sequence in the batch
            batch_size = predictions.shape[0]
            for b in range(batch_size):
                predictions[b] = torch.tensor(
                    apply_transition_constraints(predictions[b]),
                    device=predictions.device
                )

            # Return with corrected predictions
            return (loss, logits, labels)

        return outputs

    def evaluate(self, eval_dataset=None, ignore_keys=None, metric_key_prefix="eval"):
        # Override to use proximity-aware metrics
        self.compute_metrics = lambda eval_pred: compute_metrics_for_trainer(
            eval_pred, tolerance=self.proximity_tolerance
        )
        return super().evaluate(eval_dataset, ignore_keys, metric_key_prefix)


# ============================================================================
# PART 6: EVALUATION AND VISUALIZATION
# ============================================================================

def print_test_examples_with_constraints(model, tokenizer, test_csv='test_segments.csv',
                                        num_examples=5, tolerance=5):
    """
    Print test examples showing effect of logical constraints
    """
    print("\n" + "=" * 80)
    print(f"SWITCH DETECTION WITH LOGICAL CONSTRAINTS")
    print("Rules: Can only switch FROM current mode TO different mode")
    print("=" * 80)

    test_df = pd.read_csv(test_csv)
    device = next(model.parameters()).device
    model.eval()

    label_names = {
        0: 'NonSwitch-Auto',
        1: 'NonSwitch-Allo',
        2: 'SWITCH→Auto',
        3: 'SWITCH→Allo'
    }

    sample_indices = np.random.choice(len(test_df), min(num_examples, len(test_df)), replace=False)

    for i, idx in enumerate(sample_indices):
        row = test_df.iloc[idx]
        tokens = row['tokens'].split()
        true_labels = list(map(int, row['labels'].split(',')))

        print(f"\n--- Example {i+1} (from {row['source_file'][:40]}...) ---")
        print(f"Segment length: {row['num_tokens']} tokens")

        # Get predictions
        tokenizer_output = tokenizer(
            tokens,
            is_split_into_words=True,
            return_tensors="pt",
            padding=True,
            truncation=True,
            max_length=512
        )

        inputs = {k: v.to(device) for k, v in tokenizer_output.items()}

        with torch.no_grad():
            outputs = model(**inputs)
            raw_predictions = torch.argmax(outputs.logits, dim=2)

        # Align predictions
        word_ids = tokenizer_output.word_ids()
        aligned_raw = []
        previous_word_idx = None

        for j, word_idx in enumerate(word_ids):
            if word_idx is not None and word_idx != previous_word_idx:
                pred = raw_predictions[0][j].item()
                aligned_raw.append(pred)
            previous_word_idx = word_idx

        # Apply constraints
        final_len = min(len(aligned_raw), len(true_labels), len(tokens))
        aligned_raw = aligned_raw[:final_len]
        aligned_constrained = apply_transition_constraints(aligned_raw)
        true_labels = true_labels[:final_len]
        tokens = tokens[:final_len]

        # Find constraint violations that were corrected
        violations_corrected = []
        current_mode = 0

        for j in range(len(aligned_raw)):
            raw = aligned_raw[j]
            constrained = aligned_constrained[j]

            if raw != constrained:
                violations_corrected.append((j, raw, constrained, current_mode))

            # Update mode based on constrained prediction
            if constrained == 2:
                current_mode = 0
            elif constrained == 3:
                current_mode = 1
            elif constrained == 0:
                current_mode = 0
            elif constrained == 1:
                current_mode = 1

        print(f"\nConstraint violations corrected: {len(violations_corrected)}")
        if violations_corrected:
            for pos, raw, fixed, mode in violations_corrected[:5]:  # Show first 5
                mode_str = "Auto" if mode == 0 else "Allo"
                print(f"  Pos {pos}: {label_names[raw]} → {label_names[fixed]} (was in {mode_str} mode)")

        # Evaluate both versions
        raw_eval = evaluate_switch_detection_with_proximity(true_labels, aligned_raw, tolerance)
        const_eval = evaluate_switch_detection_with_proximity(true_labels, aligned_constrained, tolerance)

        print(f"\nPerformance comparison:")
        print(f"  {'Metric':<20} {'Raw':<15} {'Constrained':<15}")
        print(f"  {'-'*50}")
        print(f"  {'Precision':<20} {raw_eval['precision']:<15.3f} {const_eval['precision']:<15.3f}")
        print(f"  {'Recall':<20} {raw_eval['recall']:<15.3f} {const_eval['recall']:<15.3f}")
        print(f"  {'F1':<20} {raw_eval['f1']:<15.3f} {const_eval['f1']:<15.3f}")
        print(f"  {'False Positives':<20} {raw_eval['false_switches']:<15} {const_eval['false_switches']:<15}")

        # Show switch regions
        true_switches = [(j, l) for j, l in enumerate(true_labels) if l in [2, 3]]
        const_switches = [(j, l) for j, l in enumerate(aligned_constrained) if l in [2, 3]]

        if true_switches or const_switches:
            print(f"\nSwitch regions (showing logical flow):")
            all_switch_pos = set([p for p, _ in true_switches + const_switches])

            for switch_pos in sorted(all_switch_pos):
                start = max(0, switch_pos - 2)
                end = min(len(tokens), switch_pos + 3)

                print(f"\n  Around position {switch_pos}:")
                current_true_mode = 0
                current_pred_mode = 0

                for pos in range(start, end):
                    if pos < len(tokens):
                        token = tokens[pos]
                        true_label = true_labels[pos]
                        const_pred = aligned_constrained[pos]

                        # Determine modes
                        if true_label in [0, 2]:
                            current_true_mode = 0
                        elif true_label in [1, 3]:
                            current_true_mode = 1

                        if const_pred in [0, 2]:
                            current_pred_mode = 0
                        elif const_pred in [1, 3]:
                            current_pred_mode = 1

                        marker = "→" if pos == switch_pos else " "

                        # Check logical validity
                        logic_check = ""
                        if const_pred == 2 and current_pred_mode == 0:
                            logic_check = "⚠️ INVALID"
                        elif const_pred == 3 and current_pred_mode == 1:
                            logic_check = "⚠️ INVALID"
                        elif const_pred in [2, 3]:
                            logic_check = "✓ VALID"

                        print(f"    {marker} [{pos:3d}] {token:12s} | True: {label_names[true_label]:15s} | Pred: {label_names[const_pred]:15s} | {logic_check}")

    return sample_indices


# Add this after print_test_examples_with_constraints function:
def print_test_examples_proximity(model, tokenizer, test_csv='test_segments.csv',
                                 num_examples=5, tolerance=5):
    """
    Print detailed test examples with proximity-aware evaluation
    Shows TYPE-specific switch matching
    """
    # Fallback to the constraint version if needed
    return print_test_examples_with_constraints(model, tokenizer, test_csv, num_examples, tolerance)


# ============================================================================
# PART 7: MAIN TRAINING PIPELINE
# ============================================================================
import torch
import torch.nn as nn
from torchcrf import CRF
from transformers import AutoModelForTokenClassification, AutoTokenizer
import os


class BERTWithCRFWrapper(nn.Module):
    """
    Wraps your trained BERT model with CRF for constrained inference
    """

    def __init__(self, bert_model):
        super().__init__()
        self.bert = bert_model
        self.crf = CRF(num_tags=4, batch_first=True)
        self._init_crf_constraints()

    def _init_crf_constraints(self):
        """Set impossible transitions"""
        with torch.no_grad():
            IMPOSSIBLE = -10000.0

            # All 10 impossible transitions
            self.crf.transitions[2, 2] = IMPOSSIBLE  # No consecutive switches
            self.crf.transitions[2, 3] = IMPOSSIBLE
            self.crf.transitions[3, 2] = IMPOSSIBLE
            self.crf.transitions[3, 3] = IMPOSSIBLE
            self.crf.transitions[2, 1] = IMPOSSIBLE  # After switch, stay in mode
            self.crf.transitions[3, 0] = IMPOSSIBLE
            self.crf.transitions[0, 1] = IMPOSSIBLE  # No mode change without switch
            self.crf.transitions[0, 2] = IMPOSSIBLE  # Can't switch to same mode
            self.crf.transitions[1, 0] = IMPOSSIBLE
            self.crf.transitions[1, 3] = IMPOSSIBLE

    def forward(self, input_ids, attention_mask, **kwargs):  # ← ADD **kwargs HERE
        # Accept but ignore token_type_ids and other unexpected args
        # Get BERT predictions
        outputs = self.bert(input_ids=input_ids, attention_mask=attention_mask)
        logits = outputs.logits

        # Apply CRF Viterbi for constrained decoding
        mask = attention_mask.bool()
        predictions = self.crf.decode(logits, mask=mask)

        return {
            'logits': logits,
            'predictions': predictions
        }

    def save_pretrained(self, save_path):
        """Save both BERT and CRF"""
        os.makedirs(save_path, exist_ok=True)

        # Save BERT
        self.bert.save_pretrained(save_path)

        # Save CRF transitions
        torch.save({
            'crf_state_dict': self.crf.state_dict(),
            'transitions': self.crf.transitions.data
        }, os.path.join(save_path, 'crf_layer.pt'))

        print(f"Model saved to {save_path}")

    @classmethod
    def from_pretrained(cls, load_path):
        """Load both BERT and CRF"""
        # Load BERT
        bert_model = AutoModelForTokenClassification.from_pretrained(load_path)

        # Create wrapper
        model = cls(bert_model)

        # Load CRF if exists
        crf_path = os.path.join(load_path, 'crf_layer.pt')
        if os.path.exists(crf_path):
            crf_checkpoint = torch.load(crf_path, map_location='cpu')
            model.crf.load_state_dict(crf_checkpoint['crf_state_dict'])
            print(f"CRF layer loaded from {crf_path}")
        else:
            print("No saved CRF found - using initialized constraints")

        return model


# ============================================================================
# USAGE: Create and Save
# ============================================================================

def create_and_save_crf_model():
    """One-time setup: wrap your model and save it"""

    # Load your existing trained model
    pretrained_path = "./tibetan_code_switching_constrained_model_wylie-final_all_data_no_labels_no_prtial_v2_2_10/final_model"
    print(f"Loading trained model from {pretrained_path}")

    bert_model = AutoModelForTokenClassification.from_pretrained(pretrained_path)
    tokenizer = AutoTokenizer.from_pretrained(pretrained_path)

    # Wrap with CRF
    print("Adding CRF layer with constraints...")
    crf_model = BERTWithCRFWrapper(bert_model)

    # Save the wrapped model
    save_path = "./alloauto-segmentation-training/fine_tuned_ALTO_models/crf_enhanced_model"
    crf_model.save_pretrained(save_path)
    tokenizer.save_pretrained(save_path)

    print(f"\nModel with CRF saved to: {save_path}")
    print("You can now load this model anytime for inference")

    return crf_model, tokenizer


# ============================================================================
# USAGE: Load and Use
# ============================================================================

def load_and_use_crf_model():
    """Load the saved CRF model for inference"""

    model_path = "./alloauto-segmentation-training/fine_tuned_ALTO_models/crf_enhanced_model"

    print(f"Loading CRF-enhanced model from {model_path}")
    model = BERTWithCRFWrapper.from_pretrained(model_path)
    tokenizer = AutoTokenizer.from_pretrained(model_path)

    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    model = model.to(device)
    model.eval()

    print("Model loaded and ready for inference")
    return model, tokenizer


# ============================================================================
# Example Usage
# ============================================================================

if __name__ == "__main__":
    # Step 1: Create and save (do this ONCE)
    print("=" * 80)
    print("STEP 1: Creating CRF-enhanced model")
    print("=" * 80)
    crf_model, tokenizer = create_and_save_crf_model()

    # Step 2: Load and use (do this anytime you need inference)
    print("\n" + "=" * 80)
    print("STEP 2: Loading saved model for inference")
    print("=" * 80)
    model, tokenizer = load_and_use_crf_model()

    # Step 3: Test on a sample
    tokens = ["བོད་ཡིག", "Tibetan", "language", "སྐད་ཡིག"]
    inputs = tokenizer(
        tokens,
        is_split_into_words=True,
        return_tensors="pt",
        padding=True,
        truncation=True
    )

    # Move inputs to the same device as the model
    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    inputs = {k: v.to(device) for k, v in inputs.items()}  # ← ADD THIS LINE

    with torch.no_grad():
        outputs = model(**inputs)
        predictions = outputs['predictions'][0]

    print("\nSample prediction (guaranteed valid sequence):")
    label_names = {0: 'Auto', 1: 'Allo', 2: '→Auto', 3: '→Allo'}
    for token, pred in zip(tokens, predictions):
        print(f"  {token}: {label_names[pred]}")



